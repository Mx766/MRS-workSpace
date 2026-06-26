"""
rhBMP-2 专用翻译：提取PDF结构 → DeepSeek分段翻译 → 生成格式化DOCX
"""
import sys, io, os, json, hashlib
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'app'))
import pymupdf
import httpx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from config import load_config

cfg = load_config()
API_KEY = cfg['api_key']
MODEL = cfg['model']
BASE_URL = cfg.get('base_url_override', '') or 'https://api.deepseek.com'

ROOT_DIR = os.path.dirname(__file__)
PDF = os.path.join(ROOT_DIR, 'original', 'rhBMP-2 (DWP431) Repeat-Dose Toxicity_Rat_2 weeks.pdf')
OUTPUT = os.path.join(ROOT_DIR, 'output', 'toxicity', 'rhBMP-2_中文版.docx')
CACHE_FILE = os.path.join(ROOT_DIR, 'output', 'toxicity', 'rhbmp2_trans_cache.json')
if len(sys.argv) > 1:
    PDF = sys.argv[1]
if len(sys.argv) > 2:
    OUTPUT = sys.argv[2]

# ── 翻译缓存 ──
cache = {}
if os.path.exists(CACHE_FILE):
    with open(CACHE_FILE, 'r', encoding='utf-8') as f:
        cache = json.load(f)

def api_translate(text):
    k = hashlib.md5(text.encode()).hexdigest()[:12]
    if k in cache:
        return cache[k]
    try:
        resp = httpx.post(
            f'{BASE_URL}/v1/chat/completions',
            headers={'Authorization': f'Bearer {API_KEY}', 'Content-Type': 'application/json'},
            json={'model': MODEL, 'temperature': 0.1, 'max_tokens': 4096,
                  'messages': [
                      {'role': 'system', 'content': '将以下英文翻译为简体中文。要求：1)专业术语准确 2)保留数字/单位/符号原样 3)中文流畅自然 4)只输出译文'},
                      {'role': 'user', 'content': text}
                  ]},
            timeout=60)
        data = resp.json()
        zh = data['choices'][0]['message']['content'].strip()
        cache[k] = zh
        with open(CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(cache, f, ensure_ascii=False)
        return zh
    except Exception as e:
        print(f'  API error: {e}')
        return text

# ── 提取PDF内容 ──
print('Extracting PDF...')
doc_pdf = pymupdf.open(PDF)

# 收集所有页面文本，检测结构
pages = []
for i in range(len(doc_pdf)):
    page = doc_pdf[i]
    blocks = page.get_text('dict')['blocks']
    lines = []
    for b in blocks:
        if b['type'] != 0:
            continue
        for l in b['lines']:
            text = ' '.join([s['text'] for s in l['spans']]).strip()
            if text:
                fs = l['spans'][0]['size']
                bold = bool(l['spans'][0].get('flags', 0) & 2)
                lines.append({'text': text, 'size': fs, 'bold': bold})
    pages.append(lines)

print(f'Pages: {len(pages)}')

# ── 分段翻译 ──
# 策略：将每页文本按段落分组（空行分隔），每段一起翻译
print('Translating...')
translated_pages = []
total_segments = 0

for pi, lines in enumerate(pages):
    if not lines:
        translated_pages.append([])
        continue

    # 按空行/标题分组段落
    groups = []
    current = []
    for l in lines:
        is_header = l['bold'] and l['size'] > 11
        if is_header and current:
            groups.append(current)
            current = []
        current.append(l)
    if current:
        groups.append(current)

    page_trans = []
    for gi, group in enumerate(groups):
        # 合并组内文本
        full_text = ' '.join([l['text'] for l in group])
        if len(full_text) < 3:
            page_trans.append({'text': full_text, 'bold': group[0]['bold'], 'size': group[0]['size']})
            continue

        # 翻译
        zh = api_translate(full_text)
        total_segments += 1
        if total_segments % 50 == 0:
            print(f'  Page {pi+1}/{len(pages)}, {total_segments} segments translated')

        page_trans.append({'text': zh, 'bold': group[0]['bold'], 'size': group[0]['size']})

    translated_pages.append(page_trans)

print(f'Translation done: {total_segments} segments')

# ── 生成DOCX ──
print('Generating DOCX...')
doc = Document()

# 页面设置
for sec in doc.sections:
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for pi, groups in enumerate(translated_pages):
    if pi > 0:
        doc.add_page_break()

    for g in groups:
        text = g['text']
        if not text.strip():
            doc.add_paragraph()
            continue

        para = doc.add_paragraph()
        run = para.add_run(text)

        if g['bold'] and g['size'] > 12:
            run.bold = True
            run.font.size = Pt(14)
        elif g['bold']:
            run.bold = True
            run.font.size = Pt(11)
        elif g['size'] < 9:
            run.font.size = Pt(9)
        else:
            run.font.size = Pt(10.5)

doc.save(OUTPUT)
print(f'\nSaved: {OUTPUT}')
