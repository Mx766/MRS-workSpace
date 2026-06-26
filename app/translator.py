"""
翻译引擎 — PDF(PyMuPDF原位替换) + Word(python-docx run级)
直接调 DeepSeek API，不依赖外部进程

v3: 重构 — 使用 core.llm_client (SHA256缓存 + 批量校验) + core.docx_engine (安全替换)
"""
import os, json, logging
import pymupdf

from core.llm_client import LlmClient, make_cache_key, SYSTEM_PROMPT
from core.docx_engine import collect_paragraphs, safe_replace_paragraph

try:
    from .validators import validate_translation
    from .glossary_enforcer import post_process_translation
    from .glossary import match_glossary
except ImportError:
    from validators import validate_translation
    from glossary_enforcer import post_process_translation
    from glossary import match_glossary

logger = logging.getLogger(__name__)


def _make_client(api_key, model, base_url):
    return LlmClient(api_key, model, base_url)


# ═══════════════════════════════
#  PDF 翻译
# ═══════════════════════════════
def translate_pdf(input_path, output_path, api_key, model='deepseek-chat',
                  base_url='', glossary_domain='', progress_callback=None):
    """翻译 PDF → 中文 PDF，原位保留格式"""
    doc = pymupdf.open(input_path)
    client = _make_client(api_key, model, base_url)

    # ── 注册中文字体 ──
    font_path = os.path.join(os.environ.get('WINDIR', 'C:/Windows'), 'Fonts', 'simhei.ttf')
    cn_font = 'china-ss'  # 默认回退
    try:
        # 在首页注册 SimHei
        doc[0].insert_font(fontname='SimHei', fontfile=font_path)
        cn_font = 'SimHei'
    except Exception:
        logger.info('SimHei not available, falling back to china-ss')

    def join_spans(spans):
        """智能合并spans：相邻且无间距的span直接拼接"""
        parts = []
        prev_x1 = None
        for s in spans:
            t = s['text']
            if not t.strip():
                continue
            x0 = s['bbox'][0]
            if prev_x1 is not None:
                gap = x0 - prev_x1
                char_w = s['size'] * 0.3
                if gap < char_w:
                    parts.append(t)
                else:
                    parts.append(' ' + t)
            else:
                parts.append(t)
            prev_x1 = s['bbox'][2]
        return ''.join(parts).strip()

    # 收集所有文本行
    page_lines = []
    for page_num, page in enumerate(doc):
        for b in page.get_text('dict')['blocks']:
            if b['type'] != 0:
                continue
            for l in b['lines']:
                text = join_spans(l['spans'])
                if text and len(text) > 1:
                    page_lines.append((page_num, l['bbox'], l['spans'], text))

    all_lines = set(t for _, _, _, t in page_lines)
    lines_list = sorted(all_lines, key=len, reverse=True)

    if progress_callback:
        progress_callback(0.1, f'提取到 {len(lines_list)} 条文本')

    # 加载缓存
    cache_path = input_path + '.transcache.json'
    cache = client.load_cache(cache_path)

    # 批量翻译
    cache = client.batch_translate(
        lines_list, SYSTEM_PROMPT, cache, glossary_domain,
        max_tokens=8192, batch_size=40, progress_callback=progress_callback
    )

    client.save_cache(cache_path, cache)

    if progress_callback:
        progress_callback(0.75, '翻译完成，正在渲染 PDF...')

    # 原位替换
    replaced = 0
    for page_num in range(len(doc)):
        page = doc[page_num]
        for b in page.get_text('dict')['blocks']:
            if b['type'] != 0:
                continue
            for l in b['lines']:
                line_text = join_spans(l['spans'])
                if not line_text or len(line_text) <= 1:
                    continue

                key = make_cache_key(line_text)
                zh = cache.get(key)
                if not zh or zh == line_text:
                    continue

                bbox = l['bbox']
                font_size = l['spans'][0]['size']
                color = l['spans'][0]['color']
                font_name = l['spans'][0].get('font', '')
                flags = l['spans'][0].get('flags', 0)
                is_bold = bool(flags & 2) or 'bold' in font_name.lower()

                # 白色矩形覆盖原文
                x0, y0, x1, y1 = bbox
                shape = page.new_shape()
                shape.draw_rect(pymupdf.Rect(x0 - 2, y0 - 1, x1 + 4, y1 + 2))
                shape.finish(fill=(1, 1, 1), color=None, width=0)
                shape.commit()

                # 字号适配
                ratio = len(zh) / max(len(line_text), 1)
                adj = font_size * 0.85 if ratio > 1.5 else font_size

                # 插入中文 — 使用已注册的字体
                pos = (x0, y1 - font_size * 0.15)
                page.insert_text(pos, zh, fontname=cn_font, fontsize=adj,
                                 color=pymupdf.sRGB_to_pdf(color))
                if is_bold:
                    page.insert_text((pos[0] + 0.3, pos[1]), zh,
                                     fontname=cn_font, fontsize=adj,
                                     color=pymupdf.sRGB_to_pdf(color))
                replaced += 1

    if progress_callback:
        progress_callback(0.9, f'已替换 {replaced} 处文本')

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path, deflate=True)
    doc.close()

    if progress_callback:
        progress_callback(1.0, f'完成！{replaced} 处文本已翻译')
    return output_path, replaced


# ═══════════════════════════════
#  Word 翻译（基础版）
# ═══════════════════════════════
def translate_docx(input_path, output_path, api_key, model='deepseek-chat',
                   base_url='', glossary_domain='', progress_callback=None):
    """翻译 Word → 中文 Word，保留 run 级格式（安全替换）"""
    from docx import Document
    doc = Document(input_path)
    client = _make_client(api_key, model, base_url)

    cache_path = input_path + '.transcache.json'
    cache = client.load_cache(cache_path)

    # 收集所有段落文本
    all_paras = collect_paragraphs(doc)
    all_texts = list(set(t for _, _, t in all_paras))

    if progress_callback:
        progress_callback(0.1, f'提取 {len(all_texts)} 条文本')

    # 批量翻译
    cache = client.batch_translate(
        all_texts, SYSTEM_PROMPT, cache, glossary_domain,
        max_tokens=8192, batch_size=40, progress_callback=progress_callback
    )

    client.save_cache(cache_path, cache)

    if progress_callback:
        progress_callback(0.75, '正在写入 Word...')

    # 安全替换
    from core.docx_engine import replace_all
    from core.llm_client import make_cache_key
    replaced = replace_all(doc, cache, make_cache_key)

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)

    if progress_callback:
        progress_callback(1.0, f'完成！{replaced} 处已翻译')
    return output_path, len(cache)


# ═══════════════════════════════════════
#  翻译 + 质量保障（推荐接口）
# ═══════════════════════════════════════
def translate_docx_with_quality(input_path, output_path, api_key, model='deepseek-chat',
                                 base_url='', glossary_domain='', progress_callback=None):
    """
    翻译 Word → 中文 Word，保留格式 + 术语强制 + 数字校验 + 质量报告

    返回: (output_path, quality_report)
    """
    from docx import Document
    doc = Document(input_path)
    client = _make_client(api_key, model, base_url)

    cache_path = input_path + '.transcache.json'
    cache = client.load_cache(cache_path)

    # ── 收集段落（保留顺序，用于上下文） ──
    ordered_texts = collect_paragraphs(doc)

    # 构建上下文映射
    all_texts_set = set()
    context_map = {}
    for i, (para, tbl, text) in enumerate(ordered_texts):
        all_texts_set.add(text)
        prev_text = ordered_texts[i - 1][2] if i > 0 else ''
        next_text = ordered_texts[i + 1][2] if i < len(ordered_texts) - 1 else ''
        context_map[text] = (prev_text, next_text)

    if progress_callback:
        progress_callback(0.1,
            f'提取 {len(all_texts_set)} 条，待翻译 {sum(1 for t in all_texts_set if make_cache_key(t) not in cache)} 条')

    # ── 批量翻译（带上下文） ──
    pending = [t for t in all_texts_set if make_cache_key(t) not in cache]

    if pending:
        BATCH = 30
        plist = sorted(pending, key=len, reverse=True)
        for i in range(0, len(plist), BATCH):
            batch = plist[i:i + BATCH]
            pct = 0.1 + 0.5 * (i / max(len(plist), 1))

            if progress_callback:
                progress_callback(pct,
                    f'翻译 {i+1}-{min(i+BATCH, len(plist))}/{len(plist)}')

            # 添加上下文标记
            batch_with_context = []
            for t in batch:
                prev_t, next_t = context_map.get(t, ('', ''))
                ctx = t
                if prev_t:
                    ctx = f'[上文: {prev_t[:200]}]\n{ctx}'
                if next_t:
                    ctx = f'{ctx}\n[下文: {next_t[:200]}]'
                batch_with_context.append(ctx)

            combined = '\n<<<SEP>>>\n'.join(batch_with_context)
            gprompt = client.build_glossary_prompt(' '.join(batch), glossary_domain)

            try:
                result = client.call_api([
                    {'role': 'system',
                     'content': f'{SYSTEM_PROMPT}\n批量翻译：逐条翻译，忽略[上文]/[下文]标记（仅作参考），每条用"<<<SEP>>>"分隔。{gprompt}'},
                    {'role': 'user', 'content': combined}
                ], max_tokens=8192)

                results = [r.strip() for r in result.split('<<<SEP>>>')]

                # 校验长度
                if len(results) != len(batch):
                    logger.warning(
                        f'Batch mismatch: {len(batch)} requested vs {len(results)} received. '
                        f'Retrying individually...'
                    )
                    for t in batch:
                        key = make_cache_key(t)
                        if key in cache:
                            continue
                        try:
                            single_result = client.call_api([
                                {'role': 'system', 'content': SYSTEM_PROMPT},
                                {'role': 'user', 'content': t}
                            ], max_tokens=4096)
                            if single_result and single_result != t:
                                cache[key] = single_result
                        except Exception as e:
                            logger.warning(f'Single retry failed: {e}')
                else:
                    for src, zh in zip(batch, results):
                        if zh and zh != src:
                            cache[make_cache_key(src)] = zh

            except Exception as e:
                logger.warning(f'Batch error: {e}')
                if progress_callback:
                    progress_callback(pct, f'API错误: {e} (继续中...)')

        client.save_cache(cache_path, cache)

    if progress_callback:
        progress_callback(0.60, '正在校验翻译质量...')

    # ── 质量保障 ──
    quality = {
        'total_segments': len(ordered_texts),
        'glossary_applied': 0,
        'glossary_violations': 0,
        'glossary_violation_details': [],
        'critical_items_total': 0,
        'critical_violations': 0,
        'critical_violation_details': [],
        'warnings': [],
    }

    # 收集所有术语（一次查询整篇文档）
    all_text = ' '.join(t for _, _, t in ordered_texts)
    all_glossary_matches = match_glossary(all_text, glossary_domain)

    for para, tbl, text in ordered_texts:
        key = make_cache_key(text)
        zh = cache.get(key)
        if not zh or zh == text:
            continue

        # 1) 术语强制
        segment_matches = match_glossary(text, glossary_domain)
        if segment_matches:
            post = post_process_translation(text, zh, segment_matches)
            if post['terms_applied'] > 0:
                zh = post['final_text']
                cache[key] = zh
                quality['glossary_applied'] += post['terms_applied']
                for src, tgt, cnt in post['terms_applied_list']:
                    quality['glossary_violation_details'].append({
                        'text': text[:80],
                        'source_term': src,
                        'expected': tgt,
                        'action': '强制替换'
                    })

            if post['violation_count'] > 0:
                quality['glossary_violations'] += post['violation_count']
                for v in post['term_violations']:
                    quality['warnings'].append({
                        'type': '术语未遵守',
                        'text': text[:80],
                        'detail': f'"{v["source"]}" 未翻译为 "{v["target"]}"'
                    })

        # 2) 数字/符号校验
        validation = validate_translation(text, zh)
        quality['critical_items_total'] += validation['total_items']
        if validation['violation_count'] > 0:
            quality['critical_violations'] += validation['violation_count']
            for v in validation['violations']:
                if v['severity'] == 'high':
                    quality['critical_violation_details'].append({
                        'text': text[:80],
                        'category': v['category'],
                        'value': v['value'],
                    })
                    quality['warnings'].append({
                        'type': f'关键项丢失({v["category"]})',
                        'text': text[:80],
                        'detail': f'"{v["value"]}" 在译文中未找到'
                    })

        # 3) 安全替换段落文本
        safe_replace_paragraph(para, zh)

    # 替换表格中的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if not t or len(t) <= 1:
                        continue
                    key = make_cache_key(t)
                    zh = cache.get(key)
                    if zh and zh != t:
                        safe_replace_paragraph(para, zh)

    # ── 计算总分 ──
    gs = 1.0 - (quality['glossary_violations'] /
                max(len(all_glossary_matches), 1)) if all_glossary_matches else 1.0
    cs = 1.0 - (quality['critical_violations'] /
                max(quality['critical_items_total'], 1))
    quality['overall_score'] = round(gs * 0.4 + cs * 0.6, 3)
    quality['glossary_score'] = round(gs, 3)
    quality['critical_score'] = round(cs, 3)

    if progress_callback:
        progress_callback(0.75,
            f'质量分: {quality["overall_score"]:.0%} | 正在写入...')

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)

    # 保存质量报告
    report_path = output_path.replace('.docx', '_quality_report.json')
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(quality, f, ensure_ascii=False, indent=2)

    client.save_cache(cache_path, cache)

    if progress_callback:
        progress_callback(1.0,
            f'完成！质量分: {quality["overall_score"]:.0%} | '
            f'术语: {quality["glossary_applied"]}处强制替换')

    return output_path, quality
