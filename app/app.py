"""
文档翻译工具 — Gradio Web UI
"""
import sys, io, os, time, uuid, tempfile, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import gradio as gr

# 添加 app 目录到路径
sys.path.insert(0, os.path.dirname(__file__))
from config import load_config, save_config, test_connection, SERVICE_PRESETS
from glossary import (search_terms, add_term, delete_term, update_term,
                       import_excel, export_csv, get_all_domains, get_stats,
                       export_glossary_for_pdf2zh, match_glossary)
from translator import translate_pdf, translate_docx, translate_docx_with_quality
from reviewer import review_action, execute_replace, extract_unknown_terms
from history import (add_record, get_records, delete_record, clean_old,
                      clean_all, get_storage_size, format_size)

OUTPUT_BASE = os.path.join(os.path.dirname(__file__), '..', 'output')

# ═══════════════════════════════════════
#  Tab 1: API 设置
# ═══════════════════════════════════════
def create_api_tab():
    cfg = load_config()
    with gr.Column():
        gr.Markdown('## API 设置')
        service = gr.Dropdown(
            choices=[(v['name'], k) for k, v in SERVICE_PRESETS.items()],
            value=cfg.get('service', 'deepseek'), label='服务商'
        )
        api_key = gr.Textbox(value=cfg.get('api_key', ''), label='API Key', type='password')
        model_dd = gr.Dropdown(
            choices=SERVICE_PRESETS.get(cfg.get('service', 'deepseek'), {}).get('models', []),
            value=cfg.get('model', 'deepseek-chat'), label='模型', allow_custom_value=True
        )
        base_url = gr.Textbox(value=cfg.get('base_url_override', ''), label='Base URL 覆盖（可选）')

        def on_service_change(svc):
            p = SERVICE_PRESETS.get(svc, SERVICE_PRESETS['custom'])
            return gr.update(choices=p['models'], value=p['default_model'])

        service.change(on_service_change, service, model_dd)

        test_btn = gr.Button('测试连接', variant='secondary')
        test_result = gr.Textbox(label='结果', interactive=False)

        def do_test(svc, key, mdl, url):
            if not key:
                return '请先填写 API Key'
            ok, msg = test_connection(svc, key, mdl, url)
            return '✅ ' + msg if ok else '❌ ' + msg

        test_btn.click(do_test, [service, api_key, model_dd, base_url], test_result)

        save_btn = gr.Button('保存配置', variant='primary')
        save_msg = gr.Textbox(label='', interactive=False, visible=False)

        def do_save(svc, key, mdl, url):
            save_config(svc, key, mdl, url)
            return gr.update(value='✅ 已保存', visible=True)

        save_btn.click(do_save, [service, api_key, model_dd, base_url], save_msg)

# ═══════════════════════════════════════
#  Tab 2: 术语库
# ═══════════════════════════════════════
def create_glossary_tab():
    with gr.Column():
        gr.Markdown('## 术语库管理')

        # 统计
        stats_text = gr.Markdown('')

        def refresh_stats():
            total, domains = get_stats()
            d_str = ' | '.join([f'{d}: {c}' for d, c in domains[:10]])
            return f'**总计**: {total} 条 | {d_str}'

        gr.Button('刷新统计', size='sm').click(refresh_stats, None, stats_text)

        with gr.Tabs():
            # 搜索/列表
            with gr.TabItem('浏览搜索'):
                with gr.Row():
                    search_box = gr.Textbox(placeholder='搜索术语...', scale=3, label='搜索')
                    domain_filter = gr.Dropdown(choices=['全部'], value='全部', label='领域', scale=1)
                search_btn = gr.Button('搜索', size='sm')
                term_table = gr.Dataframe(
                    headers=['ID', '英文', '中文', '领域', '备注'],
                    datatype=['number', 'str', 'str', 'str', 'str'],
                    interactive=False, label='术语列表'
                )
                page_info = gr.Markdown('')

                def do_search(query, domain, page=1):
                    rows, total = search_terms(query, domain, page)
                    return [[r['id'], r['source'], r['target'], r['domain'], r['note']] for r in rows], \
                           f'共 {total} 条，第 {page} 页'

                def refresh_domains():
                    return gr.update(choices=['全部'] + get_all_domains())

                search_btn.click(refresh_domains, None, domain_filter).then(
                    do_search, [search_box, domain_filter], [term_table, page_info]
                )

            # 添加
            with gr.TabItem('添加术语'):
                add_src = gr.Textbox(label='英文术语')
                add_tgt = gr.Textbox(label='中文翻译')
                add_domain = gr.Dropdown(choices=['通用', '化学', '医学', '法律', '机械', '电子'], value='通用', label='领域', allow_custom_value=True)
                add_note = gr.Textbox(label='备注')
                add_btn = gr.Button('添加', variant='primary')
                add_result = gr.Textbox(label='结果', interactive=False)
                add_btn.click(lambda s,t,d,n: add_term(s,t,d,n)[1], [add_src, add_tgt, add_domain, add_note], add_result)

            # 批量导入
            with gr.TabItem('批量导入'):
                gr.Markdown('支持 Excel (.xlsx) 格式，列顺序：英文术语 | 中文翻译 | 领域 | 备注')
                import_file = gr.File(label='上传 Excel 或 CSV', file_types=['.xlsx', '.csv'])
                import_btn = gr.Button('导入', variant='primary')
                import_result = gr.Textbox(label='结果', interactive=False)

                def do_import(file):
                    if not file: return '请先上传文件'
                    fp = file.name
                    if fp.endswith('.csv'):
                        added, skipped = import_csv(fp)
                    else:
                        added, skipped, _ = import_excel(fp)
                    return f'导入 {added} 条，跳过 {skipped} 条（重复或空行）'

                import_btn.click(do_import, import_file, import_result)

            # 导出
            with gr.TabItem('导出'):
                export_btn = gr.Button('导出为 CSV（pdf2zh 兼容格式）', variant='secondary')
                export_file = gr.File(label='下载')
                def do_export():
                    path = os.path.join(tempfile.gettempdir(), 'glossary_export.csv')
                    n = export_csv(path)
                    return path if n > 0 else None
                export_btn.click(do_export, None, export_file)

# ═══════════════════════════════════════
#  Tab 3: 文档翻译
# ═══════════════════════════════════════
def create_translate_tab():
    with gr.Column():
        gr.Markdown('## 文档翻译')

        upload_file = gr.File(label='上传文件（PDF 或 Word）', file_types=['.pdf', '.docx'])
        glossary_domain = gr.Dropdown(choices=['全部'] + get_all_domains(), value='全部', label='术语领域')
        translate_btn = gr.Button('开始翻译', variant='primary', size='lg')

        status_text = gr.Textbox(label='状态', interactive=False)
        pdf_output = gr.File(label='中文 PDF 下载', visible=False)
        docx_output = gr.File(label='中文 Word 下载', visible=False)

    def do_translate(file, domain, progress=gr.Progress()):
        if not file:
            return '请上传文件', gr.update(visible=False), gr.update(visible=False)

        cfg = load_config()
        if not cfg['api_key']:
            return '请先在 API 设置中配置密钥', gr.update(visible=False), gr.update(visible=False)

        progress(0.05, desc='准备中...')

        fp = file.name
        fname = os.path.basename(fp)
        ext = os.path.splitext(fp)[1].lower()
        doc_id = f'{int(time.time())}_{uuid.uuid4().hex[:8]}'
        out_dir = os.path.join(OUTPUT_BASE, doc_id)
        os.makedirs(out_dir, exist_ok=True)

        api_key = cfg['api_key']
        model = cfg['model']
        base_url = cfg.get('base_url_override', '')

        def progress_update(pct, msg):
            progress(pct, desc=msg)

        if ext == '.pdf':
            out_pdf = os.path.join(out_dir, fname.replace('.pdf', '_中文版.pdf'))
            try:
                result_path, count = translate_pdf(fp, out_pdf, api_key, model, base_url, domain,
                                                    progress_callback=progress_update)
                add_record(fname, 'pdf', mono_path=result_path)
                progress(1.0, desc='完成！')
                return f'完成！{count} 处文本已翻译', gr.update(value=result_path, visible=True, label='中文 PDF'), gr.update(visible=False)
            except Exception as e:
                return f'翻译失败: {e}', gr.update(visible=False), gr.update(visible=False)

        elif ext == '.docx':
            out_docx = os.path.join(out_dir, fname.replace('.docx', '_中文版.docx'))
            try:
                result_path, quality = translate_docx_with_quality(
                    fp, out_docx, api_key, model, base_url, domain,
                    progress_callback=progress_update)
                add_record(fname, 'docx', docx_path=result_path)
                progress(1.0, desc='完成！')

                # 构建质量报告信息
                score = quality['overall_score']
                ga = quality['glossary_applied']
                gv = quality['glossary_violations']
                cv = quality['critical_violations']
                warnings = quality.get('warnings', [])

                status = f'完成！质量评分: {score:.0%}\n'
                status += f'术语强制替换: {ga} 处 | 术语未遵守: {gv} 处 | 关键项丢失: {cv} 处\n'

                if warnings:
                    status += f'\n⚠️ 需人工复核 ({len(warnings)} 项):\n'
                    for w in warnings[:10]:  # 最多显示10条
                        status += f'  • [{w["type"]}] {w["detail"][:80]}\n'
                    if len(warnings) > 10:
                        status += f'  ... 还有 {len(warnings)-10} 项'

                quality_report_path = out_docx.replace('.docx', '_quality_report.json')
                status += f'\n详细报告: {quality_report_path}'

                return status, gr.update(visible=False), gr.update(value=result_path, visible=True, label='中文 Word')
            except Exception as e:
                import traceback
                return f'翻译失败: {e}\n{traceback.format_exc()}', gr.update(visible=False), gr.update(visible=False)

        else:
            return f'不支持的文件格式: {ext}', gr.update(visible=False), gr.update(visible=False)

    translate_btn.click(
        do_translate, [upload_file, glossary_domain],
        [status_text, pdf_output, docx_output]
    )

# ═══════════════════════════════════════
#  Tab 4: 对话审校
# ═══════════════════════════════════════
def create_review_tab():
    with gr.Column():
        gr.Markdown('## 对话式审校')
        gr.Markdown('加载译文，用自然语言描述修改需求。例："把第3页的消防员改成消防人员"、"全文的colour统一改成颜色"')

        review_file = gr.File(label='加载译文文件（PDF 或 DOCX）', file_types=['.pdf', '.docx'])
        page_display = gr.Textbox(label='当前页内容', lines=15, interactive=False, visible=False)

        with gr.Row():
            chat_input = gr.Textbox(label='审校指令', placeholder='例：第3页的"消防员"改成"消防人员"', scale=4)
            chat_btn = gr.Button('执行', scale=1, variant='primary')

        chat_history = gr.Chatbot(label='审校记录', height=300)

        def load_for_review(file):
            if not file: return '', [], gr.update()
            fp = file.name
            if fp.endswith('.pdf'):
                import pymupdf
                doc = pymupdf.open(fp)
                pages = [doc[i].get_text() for i in range(len(doc))]
                full = '\n---PAGE BREAK---\n'.join(pages)
                return full[:500], [], gr.update(value=pages[0][:1000] if pages else '', visible=True)
            elif fp.endswith('.docx'):
                from docx import Document
                doc = Document(fp)
                full = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                return full[:500], [], gr.update(value=full[:1000], visible=True)
            return '', [], gr.update()

        review_file.upload(load_for_review, review_file, [gr.State(), chat_history, page_display])

        def do_review(message, full_text, chat_hist):
            if not message or not full_text:
                return full_text, chat_hist + [['', '请先加载译文文件']], ''
            result = review_action(message, full_text)
            if result['action'] == 'replace':
                new_text, count = execute_replace(full_text, result['find'], result['replace'])
                return new_text, chat_hist + [[message, f'✅ {result["explanation"]}（{count}处替换）']], ''
            elif result['action'] == 'add_term':
                add_term(result.get('source', ''), result.get('target', ''))
                return full_text, chat_hist + [[message, f'✅ {result["explanation"]}（已加入术语库）']], ''
            else:
                return full_text, chat_hist + [[message, f'❓ {result.get("explanation", "无法解析")}']], ''

        chat_btn.click(do_review, [chat_input, gr.State(), chat_history], [gr.State(), chat_history, chat_input])

# ═══════════════════════════════════════
#  Tab 5: 术语标记
# ═══════════════════════════════════════
def create_terms_tab():
    with gr.Column():
        gr.Markdown('## 未匹配术语处理')
        gr.Markdown('加载翻译后的文件，自动检测术语库中缺失的专业词汇。')

        term_file = gr.File(label='加载翻译文件', file_types=['.pdf', '.docx'])
        detect_btn = gr.Button('检测未匹配术语', variant='primary')
        unknown_list = gr.Dataframe(headers=['序号', '疑似术语', '建议翻译'], interactive=True, label='待确认术语')
        confirm_btn = gr.Button('全部确认为术语并入库', variant='primary')
        term_result = gr.Textbox(label='结果', interactive=False)

        def detect_terms(file):
            if not file: return [], ''
            fp = file.name
            if fp.endswith('.pdf'):
                import pymupdf
                doc = pymupdf.open(fp)
                full = '\n'.join([doc[i].get_text() for i in range(min(len(doc), 10))])
            else:
                from docx import Document
                doc = Document(fp)
                full = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            terms = extract_unknown_terms(full, None)
            return [[i+1, t, ''] for i, t in enumerate(terms[:50])], f'检测到 {len(terms)} 个疑似未匹配术语'

        detect_btn.click(detect_terms, term_file, [unknown_list, term_result])

        def confirm_terms(data):
            if data is None or len(data) == 0:
                return '无术语需要确认'
            added = 0
            for row in data:
                if row[2] and row[2].strip():
                    ok, _ = add_term(row[1], row[2])
                    if ok: added += 1
            return f'✅ {added} 条术语已加入术语库'

        confirm_btn.click(confirm_terms, unknown_list, term_result)

# ═══════════════════════════════════════
#  Tab 6: 翻译历史
# ═══════════════════════════════════════
def create_history_tab():
    with gr.Column():
        gr.Markdown('## 翻译历史')
        storage_display = gr.Markdown('')

        def refresh_history():
            records = get_records(50)
            data = []
            for r in records:
                data.append([
                    r['id'], r['filename'], r['timestamp'],
                    '✅' if r.get('mono_path') or r.get('docx_path') else '❌'
                ])
            size = format_size(get_storage_size())
            return data, f'占用空间: {size}'

        history_table = gr.Dataframe(
            headers=['ID', '文件名', '时间', '状态'],
            interactive=False, label='翻译记录'
        )
        refresh_btn = gr.Button('刷新', size='sm')
        refresh_btn.click(refresh_history, None, [history_table, storage_display])

        with gr.Row():
            selected_record = gr.Textbox(label='选择记录 ID', scale=3)
            download_btn = gr.Button('下载文件', scale=1)
            delete_btn = gr.Button('删除记录', scale=1, variant='stop')

        download_output = gr.File(label='下载')

        def do_download(record_id):
            r = get_record(record_id)
            if not r: return None
            path = r.get('mono_path') or r.get('docx_path')
            return path if path and os.path.exists(path) else None

        download_btn.click(do_download, selected_record, download_output)

        def do_delete(record_id):
            delete_record(record_id)
            return refresh_history()

        delete_btn.click(do_delete, selected_record, [history_table, storage_display])

        with gr.Row():
            clean_days = gr.Slider(minimum=1, maximum=90, value=7, step=1, label='清理 N 天前的记录', scale=3)
            clean_btn = gr.Button('按时间清理', scale=1, variant='stop')
            clean_all_btn = gr.Button('清空全部', scale=1, variant='stop')

        clean_btn.click(lambda d: (clean_old(d), refresh_history())[1], clean_days, [history_table, storage_display])
        clean_all_btn.click(lambda: (clean_all(), refresh_history())[1], None, [history_table, storage_display])

# ═══════════════════════════════════════
#  Main App — 翻译为主，其他功能收进侧栏
# ═══════════════════════════════════════
def main():
    with gr.Blocks(title='文档翻译工具') as app:
        gr.Markdown('# 外文说明书翻译工具')

        # ═══ 首页：翻译 ═══
        create_translate_tab()

        # ═══ 其他功能：折叠区 ═══
        with gr.Accordion('更多功能', open=False):
            with gr.Tabs():
                with gr.TabItem('API 设置'):
                    create_api_tab()
                with gr.TabItem('术语库管理'):
                    create_glossary_tab()
                with gr.TabItem('对话审校'):
                    create_review_tab()
                with gr.TabItem('未匹配术语'):
                    create_terms_tab()
                with gr.TabItem('翻译历史'):
                    create_history_tab()

    app.launch(server_name='0.0.0.0', server_port=7860, share=False, theme=gr.themes.Soft())

if __name__ == '__main__':
    main()
