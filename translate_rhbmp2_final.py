"""
rhBMP-2 Acrobat DOCX → 中文 DOCX（段落级翻译，保留段落格式+表格）
"""
import sys, os, hashlib, json
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'app'))
from docx import Document
import httpx
from config import load_config

cfg = load_config()
API_KEY = cfg['api_key']
MODEL = cfg.get('model', 'deepseek-chat')
BASE_URL = cfg.get('base_url_override', '') or 'https://api.deepseek.com'

ROOT_DIR = os.path.dirname(__file__)
DOCX_IN = os.path.join(ROOT_DIR, 'output', 'toxicity', 'rhBMP-2 (DWP431) Repeat-Dose Toxicity_Rat_2 weeks.docx')
DOCX_OUT = os.path.join(ROOT_DIR, 'output', 'toxicity', 'rhBMP-2_中文版.docx')
CACHE = os.path.join(ROOT_DIR, 'output', 'toxicity', 'rhbmp2_transcache.json')
if len(sys.argv) > 1:
    DOCX_IN = sys.argv[1]
if len(sys.argv) > 2:
    DOCX_OUT = sys.argv[2]

cache = {}
if os.path.exists(CACHE):
    with open(CACHE, 'r', encoding='utf-8') as f:
        cache = json.load(f)

def translate_text(text):
    k = hashlib.md5(text.encode()).hexdigest()[:12]
    if k in cache:
        return cache[k]
    try:
        resp = httpx.post(
            f'{BASE_URL}/v1/chat/completions',
            headers={'Authorization': f'Bearer {API_KEY}', 'Content-Type': 'application/json'},
            json={'model': MODEL, 'temperature': 0.1, 'max_tokens': 4096,
                  'messages': [
                      {'role': 'system', 'content': '翻译为专业简体中文。保留数字/单位/符号/化学式原样。数字与字母间不加空格。只输出译文。'},
                      {'role': 'user', 'content': text}
                  ]},
            timeout=120)
        zh = resp.json()['choices'][0]['message']['content'].strip()
        cache[k] = zh
        return zh
    except Exception as e:
        return text

print(f'Loading: {DOCX_IN}')
doc = Document(DOCX_IN)

# ── 收集所有段落文本 ──
all_texts = []
for pi, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t and len(t) > 3:
        all_texts.append(('para', pi, t))

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for pi, para in enumerate(cell.paragraphs):
                t = para.text.strip()
                if t and len(t) > 3:
                    all_texts.append(('cell', ti, ri, ci, pi, t))

print(f'Paragraphs to translate: {len(all_texts)}')

# ── 翻译 ──
pending = [(item, item[-1]) for item in all_texts if hashlib.md5(item[-1].encode()).hexdigest()[:12] not in cache]
print(f'Pending: {len(pending)}')

if pending:
    BATCH = 20
    for i in range(0, len(pending), BATCH):
        batch = pending[i:i+BATCH]
        sep = '\n<<<SEP>>>\n'
        texts = [t for _, t in batch]
        combined = sep.join(texts)
        try:
            resp = httpx.post(
                f'{BASE_URL}/v1/chat/completions',
                headers={'Authorization': f'Bearer {API_KEY}', 'Content-Type': 'application/json'},
                json={'model': MODEL, 'temperature': 0.1, 'max_tokens': 8192,
                      'messages': [
                          {'role': 'system', 'content': f'逐条翻译{len(batch)}条文本为简体中文。保留数字/单位/符号。每条用"<<<SEP>>>"分隔。只输出译文。'},
                          {'role': 'user', 'content': combined}
                      ]},
                timeout=180)
            results = [r.strip() for r in resp.json()['choices'][0]['message']['content'].split('<<<SEP>>>')]
            for (item, src), zh in zip(batch, results):
                if zh and zh != src:
                    cache[hashlib.md5(src.encode()).hexdigest()[:12]] = zh
        except Exception as e:
            print(f'  Batch err: {e}')
        if (i // BATCH) % 10 == 0:
            print(f'  {min(i+BATCH, len(pending))}/{len(pending)}')
            with open(CACHE, 'w', encoding='utf-8') as f:
                json.dump(cache, f, ensure_ascii=False)

with open(CACHE, 'w', encoding='utf-8') as f:
    json.dump(cache, f, ensure_ascii=False)

# ── 替换段落文本 ──
print('Replacing text...')
replaced = 0
for item in all_texts:
    item_type = item[0]
    text = item[-1]
    k = hashlib.md5(text.encode()).hexdigest()[:12]
    zh = cache.get(k)
    if not zh or zh == text:
        continue

    if item_type == 'para':
        _, pi, _ = item
        para = doc.paragraphs[pi]
        # 保留第一个 run 的格式，设置整段中文
        if para.runs:
            first_run = para.runs[0]
            # 清除其他 runs
            for run in para.runs[1:]:
                run.text = ''
            first_run.text = zh
        replaced += 1
    elif item_type == 'cell':
        _, ti, ri, ci, pi, _ = item
        para = doc.tables[ti].rows[ri].cells[ci].paragraphs[pi]
        if para.runs:
            first_run = para.runs[0]
            for run in para.runs[1:]:
                run.text = ''
            first_run.text = zh
        replaced += 1

doc.save(DOCX_OUT)
import os as _os
print(f'Saved: {DOCX_OUT} ({_os.path.getsize(DOCX_OUT)/1024:.0f}KB, {replaced} paragraphs)')
