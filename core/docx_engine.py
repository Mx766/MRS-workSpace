"""
DOCX 引擎 — 安全的文本提取/替换，不破坏图片/超链接/格式

解决问题:
- #4: DOCX run 替换清空多余 run → 丢失图片、超链接、粗体/斜体
- 原代码: first_run.text = zh; run.text = '' for others
  问题: 图片在 run 中、超链接在 run 中时会被清空

新方案:
- 对含图片/超链接/绘图的 run 跳过替换
- 短文本按 run 边界分配字符
- 中长文本保留 run 结构，仅替换纯文本 run
"""
import logging
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# OOXML 命名空间
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
DRAWING_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'


def _has_drawing(run) -> bool:
    """检查 run 是否包含图片/绘图"""
    for d in run._element.iter():
        if d.tag in (f'{{{DRAWING_NS}}}inline', f'{{{DRAWING_NS}}}anchor',
                      f'{{{PIC_NS}}}pic', f'{{{W_NS}}}drawing'):
            return True
    return False


def _has_hyperlink(run) -> bool:
    """检查 run 是否在超链接内"""
    parent = run._element.getparent()
    if parent is not None:
        return parent.tag == f'{{{W_NS}}}hyperlink'
    return False


def _has_tab_or_break(run) -> bool:
    """检查 run 是否包含制表符、分页符等特殊字符"""
    for child in run._element:
        if child.tag in (f'{{{W_NS}}}tab', f'{{{W_NS}}}br',
                          f'{{{W_NS}}}cr', f'{{{W_NS}}}sym'):
            return True
    return False


def safe_replace_paragraph(paragraph, new_text: str) -> bool:
    """
    安全替换段落文本，保留所有格式/图片/超链接。

    策略:
    1. 如果段落只有一个纯文本 run → 直接替换 text
    2. 如果段落有多个 run → 保留特殊 run(图片/超链接/tab)，仅替换纯文本 run
    3. 如果所有 run 都有特殊内容 → 跳过（保留原样并警告）

    返回: True 表示替换成功
    """
    if not paragraph.runs:
        return False

    runs = paragraph.runs

    # 分类 run
    text_runs = []
    special_runs = []
    total_orig_len = 0

    for run in runs:
        if _has_drawing(run) or _has_hyperlink(run) or _has_tab_or_break(run):
            special_runs.append(run)
        else:
            t = run.text
            text_runs.append(run)
            total_orig_len += len(t)

    # 如果没有纯文本 run，跳过（保护图片等）
    if not text_runs:
        logger.debug(f'Skipping paragraph with no text-only runs (images/links only)')
        return False

    # 按原始文本长度比例分配新文本到各 run
    new_len = len(new_text)
    if total_orig_len == 0:
        # 原文为空但 run 存在 → 全部放入第一个
        text_runs[0].text = new_text
        for r in text_runs[1:]:
            r.text = ''
        return True

    pos = 0
    for i, run in enumerate(text_runs):
        orig_len = len(run.text)
        if i == len(text_runs) - 1:
            # 最后一个 run 获取剩余全部文本
            run.text = new_text[pos:]
        else:
            # 按比例分配
            alloc = max(1, int(new_len * orig_len / total_orig_len))
            alloc = min(alloc, new_len - pos - (len(text_runs) - i - 1))
            if alloc <= 0:
                run.text = ''
            else:
                run.text = new_text[pos:pos + alloc]
                pos += alloc

    return True


def safe_replace_cell(cell, text_map: dict):
    """安全替换表格单元格中的文本"""
    for para in cell.paragraphs:
        orig = para.text.strip()
        zh = text_map.get(orig)
        if zh and zh != orig:
            safe_replace_paragraph(para, zh)


def collect_paragraphs(doc: Document) -> list:
    """
    收集 DOCX 中所有段落文本（含表格），返回有序列表。

    返回: [(paragraph_obj, table_obj_or_None, text), ...]
    """
    result = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t and len(t) > 1:
            result.append((para, None, t))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t and len(t) > 1:
                        result.append((para, table, t))

    return result


def replace_all(doc: Document, cache: dict, key_func) -> int:
    """
    用缓存中的译文替换整个文档的段落文本。

    参数:
        doc: python-docx Document
        cache: {key: translated_text}
        key_func: 从原文 text 生成缓存键的函数

    返回: 替换的段落数
    """
    replaced = 0

    for para in doc.paragraphs:
        t = para.text.strip()
        if not t or len(t) <= 1:
            continue
        zh = cache.get(key_func(t))
        if zh and zh != t:
            if safe_replace_paragraph(para, zh):
                replaced += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if not t or len(t) <= 1:
                        continue
                    zh = cache.get(key_func(t))
                    if zh and zh != t:
                        if safe_replace_paragraph(para, zh):
                            replaced += 1

    return replaced
