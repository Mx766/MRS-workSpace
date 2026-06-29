#!/usr/bin/env python3
"""
修复 Phase 4 中过短/不完整的 target_quote：
从 DOCX 段落中反查完整句子，替换为可精确定位的原文。
"""
import json, sys, re
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

# Unicode normalization for common DOCX issues
CHAR_MAP = {
    '³': '3',  # ³ → 3
    '²': '2',  # ² → 2
    '°': '°',  # keeping degree
}

def normalize(text: str) -> str:
    """Normalize text for matching (superscript → regular)."""
    result = text
    for k, v in CHAR_MAP.items():
        result = result.replace(k, v)
    return result

def find_sentence_containing(para_text: str, fragment: str) -> str:
    """Find the full sentence in para_text that contains the fragment."""
    if not fragment.strip() or not para_text.strip():
        return fragment

    frag_norm = normalize(fragment.strip())
    text_norm = normalize(para_text)

    # Try exact match first
    pos = text_norm.find(frag_norm)
    if pos == -1:
        # Try fuzzy: find the longest common substring
        return fragment  # Can't find, keep original

    # Expand to sentence boundaries (。！？\n)
    start = pos
    end = pos + len(frag_norm)

    # Find sentence start
    for sep in '。！？\n':
        idx = text_norm.rfind(sep, 0, start)
        if idx > 0:
            start = idx + 1
            break

    # Find sentence end
    for sep in '。！？\n':
        idx = text_norm.find(sep, end)
        if idx > 0:
            end = idx + 1  # include the separator
            break

    sentence = para_text[start:end].strip()
    return sentence if sentence else fragment

def fix_issues(issues_path: str, docx_path: str, output_path: str):
    from docx import Document

    doc = Document(docx_path)
    issues = json.loads(Path(issues_path).read_text(encoding='utf-8'))

    fixed = 0
    for i, iss in enumerate(issues):
        if iss.get('type'):
            continue  # mechanical issues — no target_quote to fix

        tq = iss.get('target_quote', '')
        pi = iss.get('paragraph_index', 0)

        # Get paragraph text
        para_text = ''
        if 0 < pi <= len(doc.paragraphs):
            para_text = doc.paragraphs[pi - 1].text

        if not para_text:
            continue

        tq_norm = normalize(tq.strip())
        text_norm = normalize(para_text)

        # Skip if tq is already a complete sentence (ends with 。！？)
        if tq_norm and tq_norm[-1] in '。！？':
            if tq_norm in text_norm:
                continue  # Good — found in paragraph, long enough

        # Skip if tq is long enough (>= 20 chars) AND found in paragraph
        if len(tq_norm) >= 15 and tq_norm in text_norm:
            continue

        # Needs fix: expand to full sentence from paragraph
        new_tq = find_sentence_containing(para_text, tq)

        # If fragment-based search failed, try whole-paragraph fallback
        if new_tq == tq:
            # Try: search for the first 3 significant chars
            key = tq_norm[:5] if len(tq_norm) >= 5 else tq_norm
            if len(key) >= 3:
                new_tq = find_sentence_containing(para_text, key)

        if new_tq and new_tq != tq:
            # Quality gate: must be longer than original AND found in paragraph
            if len(new_tq) < len(tq.strip()):
                continue  # Reject — shorter than original
            if normalize(new_tq) not in text_norm:
                # Try harder: check if at least 80% of new_tq chars are in para
                overlap = sum(1 for c in new_tq if c in para_text)
                if overlap < len(new_tq) * 0.8:
                    continue  # Reject — not actually in paragraph

            old = iss['target_quote']
            iss['target_quote'] = new_tq
            iss['_tq_fixed'] = f'expanded from {repr(old)}'
            fixed += 1
            print(f'[{i}] pi={pi}: {repr(old)[:60]}')
            print(f'       →  {repr(new_tq)[:120]}')
            print()

    Path(output_path).write_text(json.dumps(issues, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f'Fixed {fixed} target_quotes')
    print(f'Output: {output_path}')

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--issues', required=True)
    parser.add_argument('--docx', required=True)
    parser.add_argument('--output', required=True)
    args = parser.parse_args()
    fix_issues(args.issues, args.docx, args.output)
