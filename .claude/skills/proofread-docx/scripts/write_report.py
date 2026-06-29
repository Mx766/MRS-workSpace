#!/usr/bin/env python3
"""
Phase 7: 生成 Word 格式校对报告
输入: 配对信息 + 各阶段疑点汇总 + 术语库变更
输出: Word (.docx) 格式的校对报告
"""

import argparse
import json
import os
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')
from datetime import datetime

try:
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError:
    pass


def generate_report(pair_info: dict, all_issues: list[dict],
                    glossary_changes: list[dict] = None,
                    mechanical_report: dict = None,
                    mech_summary: dict = None,
                    output_path: str = None) -> str:
    """
    生成 Word 格式校对报告。
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return json.dumps({'error': '需要 python-docx: pip install python-docx'})

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = '宋体'

    # ── 标题 ──
    title = doc.add_heading('翻译校对报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 一、基本信息 ──
    doc.add_heading('一、基本信息', level=1)

    # 兼容两种格式: 直接 {source:, target:} 或嵌套 {pairs: [{source:, target:}]}
    if 'pairs' in pair_info and pair_info['pairs']:
        src = pair_info['pairs'][0].get('source', {})
        tgt = pair_info['pairs'][0].get('target', {})
    else:
        src = pair_info.get('source', {})
        tgt = pair_info.get('target', {})

    source_file = src.get('filename', '')
    target_file = tgt.get('filename', '')
    target_lang = tgt.get('lang_name', '')
    source_format = src.get('format', '')

    info_data = [
        ('原文文件', f'{source_file}' + (f' ({source_format.upper()})' if source_format else '')),
        ('译文文件', f'{target_file}'),
        ('校对方向', f'原文语种 → {target_lang}' if target_lang else ''),
        ('校对日期', datetime.now().strftime('%Y-%m-%d')),
    ]

    if pair_info.get('glossary_files'):
        info_data.append(('使用术语库', ', '.join(pair_info['glossary_files'])))

    if pair_info.get('total_paragraphs'):
        info_data.append(('文档总段数', str(pair_info['total_paragraphs'])))

    if pair_info.get('chapter_count'):
        info_data.append(('校对章节数', str(pair_info['chapter_count'])))

    _add_info_table(doc, info_data)

    # ── 颜色图例 ──
    doc.add_paragraph()
    legend_heading = doc.add_paragraph('批注颜色说明：')
    legend_heading.runs[0].bold = True

    legend_items = [
        ('🔴 红色高亮 = 严重问题（必须修改）：漏译、术语前后不一致、相同段落翻译不一致、不通顺、数字错误', 'FF0000'),
        ('🟡 橙色高亮 = 中等问题（建议修改）：翻译腔、标点混用', 'FF8C00'),
        ('🟢 绿色高亮 = 低问题（格式优化）：格式问题', '228B22'),
    ]
    for text, color in legend_items:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.color.rgb = RGBColor(
                int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            )

    # ── 二、问题统计总览 ──
    doc.add_heading('二、问题统计总览', level=1)

    # 按维度和严重度分组统计
    stats = _calculate_stats(all_issues)

    # 统计表
    has_mech = mech_summary and mech_summary.get('total', 0) > 0
    nrows = len(stats['by_dimension']) + 2 + (1 if has_mech else 0)
    table = doc.add_table(rows=nrows, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['维度', '🔴 严重', '🟡 中等', '🟢 低', '合计']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    for row_idx, (dim, counts) in enumerate(stats['by_dimension'].items()):
        row = table.rows[row_idx + 1]
        row.cells[0].text = dim
        row.cells[1].text = str(counts.get('critical', 0))
        row.cells[2].text = str(counts.get('medium', 0))
        row.cells[3].text = str(counts.get('low', 0))
        row.cells[4].text = str(sum(counts.values()))

        # 颜色填充
        if counts.get('critical', 0) > 0:
            _set_cell_fill(row.cells[1], 'FFCCCC')
        if counts.get('medium', 0) > 0:
            _set_cell_fill(row.cells[2], 'FFF9C4')

    next_row = len(stats['by_dimension']) + 1

    # 机械检查补充行
    if has_mech:
        mech_row = table.rows[next_row]
        mech_row.cells[0].text = '机械检查（全文匹配参考）'
        for run in mech_row.cells[0].paragraphs[0].runs:
            run.bold = True
        mech_row.cells[1].text = '0'
        mech_row.cells[2].text = str(mech_summary['total'])
        mech_row.cells[3].text = '0'
        mech_row.cells[4].text = str(mech_summary['total'])
        _set_cell_fill(mech_row.cells[2], 'FFF9C4')
        next_row += 1

    # 合计行
    total_row = table.rows[next_row]
    total_row.cells[0].text = '合计'
    for run in total_row.cells[0].paragraphs[0].runs:
        run.bold = True
    total_crit = stats['total'].get('critical', 0)
    total_med = stats['total'].get('medium', 0)
    total_low = stats['total'].get('low', 0)
    if mech_summary:
        total_med += mech_summary.get('total', 0)
    total_row.cells[1].text = str(total_crit)
    total_row.cells[2].text = str(total_med)
    total_row.cells[3].text = str(total_low)
    total_row.cells[4].text = str(total_crit + total_med + total_low)

    # ── 三、🔴 严重问题清单（必须修改） ──
    critical_issues = [i for i in all_issues if i.get('severity') == 'critical']
    if critical_issues:
        doc.add_heading('三、🔴 严重问题清单（必须修改）', level=1)
        for idx, issue in enumerate(critical_issues):
            _add_issue_paragraph(doc, idx + 1, issue, 'critical')
    else:
        doc.add_heading('三、🔴 严重问题清单（必须修改）', level=1)
        doc.add_paragraph('✅ 未发现严重问题。')

    # ── 四、🟡 中等问题汇总（建议修改） ──
    medium_issues = [i for i in all_issues if i.get('severity') == 'medium']
    has_mech = mech_summary and mech_summary.get('total', 0) > 0
    if medium_issues or has_mech:
        doc.add_heading('四、🟡 中等问题汇总（建议修改）', level=1)

        # 4a. 机械检查参考（全文匹配，仅供参考）
        if has_mech:
            p = doc.add_paragraph(
                f'⚠️ 机械检查发现 {mech_summary["total"]} 处疑似问题'
                f'（跨格式 PDF→DOCX 全文匹配，无法定位到具体段落，仅供参考）：'
            )
            p.runs[0].bold = True
            for label, info in mech_summary.get('by_type', {}).items():
                examples = '、'.join(f'"{e}"' for e in info['examples'] if e)
                doc.add_paragraph(f'  • {label}: {info["count"]} 处（如 {examples}）')
            doc.add_paragraph('建议人工逐条复核确认。')
            doc.add_paragraph()

        # 4b. 语义审查问题（有上下文，逐条列出）
        if medium_issues:
            p = doc.add_paragraph(f'以下 {len(medium_issues)} 条问题经语义审查确认，附原文+译文上下文：')
            p.runs[0].bold = True
            for idx, issue in enumerate(medium_issues):
                _add_issue_paragraph(doc, idx + 1, issue, 'medium')
    else:
        doc.add_heading('四、🟡 中等问题汇总（建议修改）', level=1)
        doc.add_paragraph('✅ 未发现中等问题。')

    # ── 五、🟢 低问题汇总（格式优化） ──
    low_issues = [i for i in all_issues if i.get('severity') == 'low']
    if low_issues:
        doc.add_heading('五、🟢 低问题汇总（格式优化）', level=1)
        for idx, issue in enumerate(low_issues):
            _add_issue_paragraph(doc, idx + 1, issue, 'low')
    else:
        doc.add_heading('五、🟢 低问题汇总（格式优化）', level=1)
        doc.add_paragraph('✅ 未发现格式问题。')

    # ── 六、术语一致性报告 ──
    doc.add_heading('六、术语一致性报告', level=1)
    if mechanical_report:
        gl_vios = mechanical_report.get('glossary_violations', [])
        glossary_match_rate = _calc_glossary_rate(mechanical_report)
        doc.add_paragraph(f'术语库匹配率: {glossary_match_rate}')
        doc.add_paragraph(f'术语违规: {len(gl_vios)} 处')
    else:
        doc.add_paragraph('术语库匹配率: 无术语库')

    if glossary_changes:
        doc.add_paragraph(f'新增入库术语: {len(glossary_changes)} 条')

    # ── 七、范围表达统一性 ──
    if mechanical_report and mechanical_report.get('range_stats'):
        range_stats = mechanical_report['range_stats']
        doc.add_heading('七、范围表达统一性检查', level=1)
        non_zero = {k: v for k, v in range_stats.items() if v > 0}
        if len(non_zero) > 1:
            doc.add_paragraph(f'⚠️ 范围表达方式混用:')
            for form, count in non_zero.items():
                form_label = {'dash': '1-5(连字符)', 'tilde': '1～5(波浪号)', 'chinese': '1至5(中文)'}
                doc.add_paragraph(f'  • {form_label.get(form, form)}: {count} 处')
            doc.add_paragraph('建议全文统一为一种范围表达方式。')
        elif len(non_zero) == 1:
            form = list(non_zero.keys())[0]
            form_label = {'dash': '1-5(连字符)', 'tilde': '1～5(波浪号)', 'chinese': '1至5(中文)'}
            doc.add_paragraph(f'✅ 全文范围表达统一: {form_label.get(form, form)}')
        else:
            doc.add_paragraph('（文档中未使用范围表达）')

    # ── 八、待确认事项 ──
    doc.add_heading('八、待确认事项', level=1)
    low_confidence = [i for i in all_issues if i.get('confidence') == 'low']
    if low_confidence:
        doc.add_paragraph(f'以下 {len(low_confidence)} 条问题置信度较低，建议人工确认：')
        for idx, issue in enumerate(low_confidence):
            doc.add_paragraph(f'{idx+1}. [{issue.get("chapter", "")}] {issue.get("issue", "")[:100]}')
    else:
        doc.add_paragraph('无待确认事项。')

    # ── 九、附录：术语库变更记录 ──
    if glossary_changes:
        doc.add_heading('九、附录：术语库变更记录', level=1)
        table = doc.add_table(rows=len(glossary_changes) + 1, cols=5)
        table.style = 'Table Grid'
        for i, h in enumerate(['英文术语', '中文译法', '领域', '来源', '入库时间']):
            cell = table.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for row_idx, change in enumerate(glossary_changes):
            table.rows[row_idx + 1].cells[0].text = change.get('source_term', '')
            table.rows[row_idx + 1].cells[1].text = change.get('candidate_translation', '')
            table.rows[row_idx + 1].cells[2].text = change.get('domain', '')
            table.rows[row_idx + 1].cells[3].text = change.get('source', '')
            table.rows[row_idx + 1].cells[4].text = change.get('added_at', '')

    # ── 保存 ──
    output = output_path or '校对报告.docx'
    doc.save(output)
    return json.dumps({'status': 'ok', 'output': str(Path(output).resolve())}, ensure_ascii=False)


def _add_info_table(doc, data: list):
    """添加基本信息表格"""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True
    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12)


def _add_issue_paragraph(doc, idx: int, issue: dict, severity: str):
    """添加单条问题描述"""
    sev_colors = {
        'critical': ('FF0000', '🔴'),
        'medium': ('FF8C00', '🟡'),
        'low': ('228B22', '🟢'),
    }
    hex_color, icon = sev_colors.get(severity, ('000000', ''))

    chapter = issue.get('chapter', '')
    # Strip existing 'Ch' prefix if present to avoid 'ChCh' double prefix
    if isinstance(chapter, str) and chapter.startswith('Ch'):
        chapter = chapter[2:]
    para_idx = issue.get('paragraph_index', '')
    loc = f'[Ch{chapter} P{para_idx}]' if chapter and str(para_idx) else ''

    p = doc.add_paragraph()
    run = p.add_run(f'{idx}. {icon} {loc} {issue.get("check_item", "")}')
    run.bold = True
    run.font.color.rgb = RGBColor(
        int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    )

    # 动态构建行列表：空字段跳过对应行（铁律 D）
    lines = []
    sq = issue.get("source_quote", "").strip()
    tq = issue.get("target_quote", "").strip()
    is_text = issue.get("issue", "").strip()
    sug = issue.get("suggestion", "").strip()
    if sq:
        lines.append(('原文', sq, False))
    if tq:
        lines.append(('译文', tq, False))
    if is_text:
        lines.append(('问题', is_text, False))
    if sug:
        lines.append(('建议', sug, True))  # italic for suggestion
    # Fallback: if all fields empty, at least show suggestion/check
    if not lines:
        fallback = sug or issue.get("check", "").strip()
        if fallback:
            lines.append(('说明', fallback, True))

    for label, text, italic in lines:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        run = p2.add_run(f'{label}: {text}')
        if italic:
            run.italic = True


def _set_cell_fill(cell, color: str):
    """设置表格单元格底色"""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    except Exception:
        pass


def _calculate_stats(issues: list[dict]) -> dict:
    """计算问题统计"""
    by_dimension = {}
    total = {'critical': 0, 'medium': 0, 'low': 0}

    for issue in issues:
        dim = issue.get('dimension', '其他')
        sev = issue.get('severity', 'medium')
        if dim not in by_dimension:
            by_dimension[dim] = {'critical': 0, 'medium': 0, 'low': 0}
        by_dimension[dim][sev] = by_dimension[dim].get(sev, 0) + 1
        total[sev] = total.get(sev, 0) + 1

    return {
        'by_dimension': by_dimension,
        'total': total,
        'grand_total': sum(total.values()),
    }


def _calc_glossary_rate(mech_report: dict) -> str:
    """计算术语库匹配率"""
    gl_vios = mech_report.get('glossary_violations', [])
    if isinstance(gl_vios, list):
        return f'检测到 {len(gl_vios)} 处术语违规'
    return '无术语库数据'


def _summarize_mechanical(issues: list[dict]) -> dict:
    """汇总机械检查 pi=0 问题（跨格式全文匹配，仅供参考）。"""
    by_type = {}
    for issue in issues:
        t = issue.get('type', issue.get('check_item', '其他'))
        label = {
            'number_missing': '数字缺失',
            'decimal_mismatch': '小数位不匹配',
            'glossary_violation': '术语库违规',
            'symbol_missing': '符号缺失',
            'format_issue': '格式问题',
            'unit_issue': '单位问题',
        }.get(t, t)
        if label not in by_type:
            by_type[label] = {'count': 0, 'examples': []}
        by_type[label]['count'] += 1
        if len(by_type[label]['examples']) < 3:
            by_type[label]['examples'].append(
                issue.get('source_quote', '') or issue.get('source_term', '') or issue.get('source_value', '')
            )
    return {
        'total': len(issues),
        'by_type': by_type,
        'note': '以下问题由跨格式（PDF→DOCX）全文匹配产生，无法定位到具体段落，仅供参考。建议人工复核确认。',
    }


def _normalize_issues(issues: list[dict]) -> list[dict]:
    """将机械检查 issues 的字段名统一为 Phase 4 语义审查的字段名。
    
    机械检查格式 (Phase 3):
      - type (number_missing / glossary_violation / ...)
      - source_value (原文数值)
      - check (问题描述 + 建议合一)
      - paragraph_index
    
    语义审查格式 (Phase 4):
      - check_item (如 "1.4 错译")
      - source_quote (原文原句)
      - target_quote (译文原句)
      - issue (问题描述)
      - suggestion (修改建议)
      - chapter / paragraph_index / dimension / severity / confidence
    """
    type_labels = {
        'number_missing': '数字缺失',
        'decimal_mismatch': '小数位不匹配',
        'glossary_violation': '术语库违规',
        'symbol_missing': '符号缺失',
        'punctuation_mixed': '标点混用',
        'format_issue': '格式问题',
        'unit_issue': '单位问题',
    }
    for issue in issues:
        # If already has Phase 4 fields, skip
        if issue.get('check_item') and issue.get('source_quote'):
            continue
        
        issue_type = issue.get('type', '')
        if not issue_type:
            # Try to infer: has source_term → glossary_violation
            if issue.get('source_term'):
                issue_type = 'glossary_violation'
            elif issue.get('source_value'):
                issue_type = 'number_missing'
        
        # check_item
        if not issue.get('check_item'):
            issue['check_item'] = type_labels.get(issue_type, issue_type or '机械检查')
        
        # source_quote ← source_value or source_term
        if not issue.get('source_quote'):
            issue['source_quote'] = issue.get('source_value') or issue.get('source_term') or issue.get('check', '')[:200]
        
        # target_quote
        if not issue.get('target_quote'):
            issue['target_quote'] = issue.get('target_value') or issue.get('expected_target', '')
        
        # suggestion ← check
        if not issue.get('suggestion'):
            issue['suggestion'] = issue.get('check', '')
        
        # issue ← synthesize from check / type
        if not issue.get('issue'):
            check_text = issue.get('check', '')
            if '缺失' in check_text:
                issue['issue'] = f'原文数值 "{issue.get("source_value", "")}" 在译文中可能缺失。'
            elif '术语' in check_text or '应译为' in check_text:
                src = issue.get('source_term', '')
                exp = issue.get('expected_target', '')
                issue['issue'] = f'术语 "{src}" 术语库要求译为 "{exp}"，但译文中未找到匹配。'
            else:
                issue['issue'] = check_text[:200]
        
        # dimension (if missing)
        if not issue.get('dimension'):
            if 'number' in issue_type or 'decimal' in issue_type:
                issue['dimension'] = '三.数字符号单位'
            elif 'glossary' in issue_type:
                issue['dimension'] = '四.术语合规'
            elif 'symbol' in issue_type or 'punctuation' in issue_type:
                issue['dimension'] = '二.表达规范'
            elif 'format' in issue_type:
                issue['dimension'] = '五.格式排版'
            else:
                issue['dimension'] = '三.数字符号单位'
        
        # confidence (if missing)
        if not issue.get('confidence'):
            issue['confidence'] = 'medium' if issue.get('paragraph_index', 0) < 1 else 'high'
    
    return issues


def main():
    parser = argparse.ArgumentParser(description='生成校对报告')
    parser.add_argument('--pair-info', '-p', help='配对信息 JSON')
    parser.add_argument('--issues', '-i', nargs='+', default=[], help='疑点 JSON 文件列表（可多个，合并处理）')
    parser.add_argument('--glossary-changes', '-g', help='术语库变更 JSON')
    parser.add_argument('--mechanical-report', '-m', help='机械检查报告 JSON')
    parser.add_argument('--output', '-o', default='校对报告.docx', help='输出报告路径')
    args = parser.parse_args()

    # 加载配对信息
    pair_info = {}
    if args.pair_info:
        try:
            pair_info = json.loads(Path(args.pair_info).read_text(encoding='utf-8'))
        except Exception as e:
            print(f'警告: 无法加载配对信息 {args.pair_info}: {e}', file=sys.stderr)
            pair_info = {}

    # 合并所有 issues
    all_issues = []
    for issue_file in args.issues:
        try:
            data = json.loads(Path(issue_file).read_text(encoding='utf-8'))
            if isinstance(data, list):
                all_issues.extend(data)
            elif isinstance(data, dict):
                all_issues.extend(data.get('issues', data.get('findings', [])))
        except Exception as e:
            print(f'警告: 无法加载 issues 文件 {issue_file}: {e}', file=sys.stderr)

    # ── 格式归一化：机械检查 issues → Phase 4 issues 字段名 ──
    all_issues = _normalize_issues(all_issues)

    # ── 拆分：语义审查问题 vs 机械检查参考（pi=0 全文匹配）──
    review_issues = []       # Phase 4 语义审查 + pi>=1 问题（有真实上下文）
    mech_ref_issues = []     # pi=0 机械问题（跨格式全文匹配，仅供参考）
    for issue in all_issues:
        pi = issue.get('paragraph_index', 0)
        # 机械检查问题特征：有 type 字段且 pi=0
        if issue.get('type') and pi < 1:
            mech_ref_issues.append(issue)
        else:
            review_issues.append(issue)
    mech_summary = _summarize_mechanical(mech_ref_issues) if mech_ref_issues else None

    # 术语库变更
    glossary_changes = None
    if args.glossary_changes:
        try:
            raw = json.loads(Path(args.glossary_changes).read_text(encoding='utf-8'))
            # Normalize: accept both list and dict formats
            if isinstance(raw, list):
                glossary_changes = raw
            elif isinstance(raw, dict):
                glossary_changes = raw.get('new_terms', raw.get('terms', raw.get('changes', [])))
                if not isinstance(glossary_changes, list):
                    print(f'警告: glossary_changes JSON 格式不兼容（期望 list 或含 new_terms/terms/changes 的 dict），已设为空列表', file=sys.stderr)
                    glossary_changes = []
            else:
                print(f'警告: glossary_changes JSON 类型不兼容 ({type(raw).__name__})，已设为空列表', file=sys.stderr)
                glossary_changes = []
        except Exception as e:
            print(f'警告: 无法加载 glossary_changes {args.glossary_changes}: {e}', file=sys.stderr)
            glossary_changes = []

    # 机械检查报告
    mechanical_report = None
    if args.mechanical_report:
        try:
            mechanical_report = json.loads(Path(args.mechanical_report).read_text(encoding='utf-8'))
        except Exception:
            pass

    # Ensure output directory exists before generating report
    out_dir = os.path.dirname(os.path.abspath(args.output))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    result = generate_report(pair_info, review_issues, glossary_changes, mechanical_report, mech_summary, args.output)
    print(result)


if __name__ == '__main__':
    main()
