"""
MSDS PDF 翻译工具 — 提取 → 翻译 → 生成格式化的中英文对照 Word 文档
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json

# ============================================================
# 翻译内容 — 英文原文 → 中文翻译（逐段对应）
# ============================================================
TRANSLATIONS = {
    # Page 1
    "SAFETY DATA SHEET": "安全数据说明书（SDS）",
    "Revision Date": "修订日期",
    "product code": "产品代码",
    "Version / Revision": "版本/修订",
    "Issuing date": "签发日期",
    "Supersedes Version": "替代版本",
    "TOPAS® Cyclic Olefin Copolymers": "TOPAS® 环烯烃共聚物",
    "1: Identification": "第1节：化学品及企业标识",
    "Product Identifier": "产品标识符",
    "Identification of the substance/preparation": "物质/混合物标识",
    "Recommended uses and restrictions on use": "建议用途及使用限制",
    "Use of the Substance / Preparation": "物质/混合物的用途",
    "injection molding articles for optical industry, 3-D Printing, packaging Industry, medical articles.":
        "用于光学行业的注塑制品、3D打印、包装行业、医疗制品。",
    "Supplier information": "供应商信息",
    "Supplier": "供应商",
    "Product Information email:": "产品信息电子邮箱：",
    "Emergency telephone number": "紧急联系电话",
    "in USA, call 800 424 9300": "美国境内，拨打 800 424 9300",
    "outside USA, call +1.703.527.3887, collect calls accepted": "美国境外，拨打 +1.703.527.3887，可接听对方付费电话",
    "available 24/7": "全天候可用（24/7）",
    "2. Hazards identification": "第2节：危险性概述",
    "2.1. Classification of the substance or mixture": "2.1. 物质或混合物的分类",
    "This substance is not hazardous in accordance with paragraph (d) of §1910.1200 (GHS-US classification).***":
        "根据 §1910.1200(d)（美国GHS分类），本物质不具有危险性。***",
    "OSHA Specified Hazards": "OSHA 特定危险",
    "Not applicable.": "不适用。",
    "2.2. Label elements": "2.2. 标签要素",
    "Not required according to §1910.1200 (GHS-US labeling).": "根据 §1910.1200（美国GHS标签规范），无需标签。",
    "2.3. Other hazards": "2.3. 其他危害",
    "None known": "无已知危害",
    "3. Composition / Information on ingredients": "第3节：成分/组成信息",

    # Page 2
    "Chemical characterization": "化学特性",
    "contains ethylene-norbornene copolymer (CAS 26007-43-2)": "含有乙烯-降冰片烯共聚物（CAS 26007-43-2）",
    "Remarks": "备注",
    "The following specific grades of TOPAS are covered by this MSDS:": "本安全数据说明书涵盖以下 TOPAS 特定牌号：",
    "4. First aid measures": "第4节：急救措施",
    "Description of first aid measures": "急救措施描述",
    "General advice": "一般建议",
    "Remove/Take off immediately all contaminated clothing. Wash/Decontaminate removed clothing before reuse.":
        "立即脱去所有受污染的衣物。在重新使用前清洗/净化脱下的衣物。",
    "Inhalation": "吸入",
    "Aerate with fresh air. When symptoms persist or in all cases of doubt seek medical advice.":
        "移至空气新鲜处。如症状持续或有任何疑问，请就医。",
    "Eyes": "眼睛接触",
    "Resin particles, like other inert materials, are mechanically irritating to eyes. Rinse immediately with plenty of water, also under the eyelids, for at least 15 minutes. Obtain medical attention.":
        "树脂颗粒如同其他惰性材料，会对眼睛产生机械性刺激。立即用大量清水冲洗，同时翻开眼睑，至少冲洗15分钟。请就医。",
    "Skin": "皮肤接触",
    "Cool skin rapidly with cold water after contact with molten polymer. If polymer is stuck to skin, do not remove. Allow adhered polymer to come off naturally. Removal of adhered polymer may result in more tissue damages than if polymer is allowed to come off over time. When symptoms persist or in all cases of doubt seek medical advice.":
        "接触熔融聚合物后，用冷水迅速冷却皮肤。如果聚合物粘在皮肤上，不要强行去除。让粘附的聚合物自然脱落。强行去除粘附的聚合物可能比让聚合物随时间自然脱落造成更多的组织损伤。如症状持续或有任何疑问，请就医。",
    "Ingestion": "食入",
    "Do not induce vomiting without medical advice. Obtain medical attention.":
        "未经医生指导不要催吐。请就医。",
    "Most important symptoms and effects, both acute and delayed": "最重要的症状和影响（急性和迟发性）",
    "Main symptoms": "主要症状",
    "Advice for the rescuer": "急救人员须知",
    "Special note for doctor": "给医生的特别提示",
    "This product is essentially inert and non-toxic. Under conditions of thermal decomposition irritant gases may be formed. Exposed patients may need to have their arterial blood gases and carboxyhemoglobin levels checked":
        "本产品基本上是惰性且无毒的。在热分解条件下可能产生刺激性气体。暴露患者可能需要检查动脉血气和碳氧血红蛋白水平。",
    "5. Firefighting measures": "第5节：消防措施",
    "Extinguishing media": "灭火介质",
    "Suitable extinguishing media": "合适的灭火介质",
    "water spray, foam, dry chemical, carbon dioxide (CO2).": "水雾、泡沫、干粉、二氧化碳（CO₂）。",

    # Page 3
    "Unsuitable Extinguishing Media": "不合适的灭火介质",
    "Do not use a solid water stream as it may scatter and spread fire.": "不要使用直流水柱，因为它可能散射并扩散火势。",
    "Special hazards": "特殊危害",
    "Under conditions giving incomplete combustion, hazardous gases produced may consist of:":
        "在不完全燃烧的条件下，产生的有害气体可能包括：",
    "carbon monoxide (CO)": "一氧化碳（CO）",
    "carbon dioxide (CO2)": "二氧化碳（CO₂）",
    "Combustion gases of organic materials must in principle be graded as inhalation poisons":
        "有机材料的燃烧气体原则上必须被归类为吸入性毒物。",
    "Fire precautions and protective measures": "消防预防和保护措施",
    "Cool closed containers exposed to fire with water spray": "用水雾冷却暴露在火中的密闭容器。",
    "Keep people away from and upwind of fire": "人员远离火源并站在上风处。",
    "Dike and collect water used to fight fire": "围堵和收集消防用水。",
    "Special protective equipment for firefighters": "消防员特殊防护装备",
    "Fire fighter protection should include a self-contained breathing apparatus (NIOSH-approved or EN 133) and full fire-fighting turn out gear.":
        "消防员防护应包括自给式呼吸器（NIOSH批准或EN 133）和全套消防战斗服。",
    "6. Accidental release measures": "第6节：泄漏应急处理",
    "Personnel precautions, protective equipment and emergency procedures": "人员防护措施、防护装备和应急程序",
    "Personal precautions": "个人防护措施",
    "Avoid contact with skin and eyes. Do not breathe dust. Keep people away from and upwind of spill/leak.":
        "避免与皮肤和眼睛接触。不要吸入粉尘。人员远离泄漏/溢出区域并站在上风处。",
    "For emergency responders: Personal protection see section 8.": "对于应急响应人员：个人防护参见第8节。",
    "Environmental precautions": "环境保护措施",
    "Not readily biodegradable. Should not be released into the environment. Do not flush into surface water or sanitary sewer system.":
        "不易生物降解。不应释放到环境中。不要冲入地表水或卫生污水系统。",
    "Methods and material used for collection and disposal of leak": "泄漏的收集和处置方法及材料",
    "Methods for containment": "围堵方法",
    "Stop the flow of material, if possible without risk.": "如果无风险，停止物料流动。",
    "Methods for cleaning up": "清理方法",
    "Sweep up and shovel into suitable containers for disposal. Like most thermoplastic plastics the product can be recycled. Dispose of in accordance with local regulations.":
        "清扫并铲入适当的容器中待处置。如同大多数热塑性塑料，本产品可以回收。根据当地法规进行处置。",
    "Precautionary measures to prevent the occurrence of secondary disasters": "防止次生灾害发生的预防措施",
    "7. Handling and storage": "第7节：操作处置与储存",
    "Handling": "操作处置",

    # Page 4
    "Advice on safe handling": "安全操作建议",
    "Do not handle hot or molten material without appropriate protective equipment. Do not exceed recommended process temperatures to minimize release of decomposition products.":
        "没有适当的防护设备不要处理热或熔融材料。不要超过推荐的加工温度，以尽量减少分解产物的释放。",
    "Hygiene measures": "卫生措施",
    "Wash hands before breaks and immediately after handling the product. Take off all contaminated clothing immediately":
        "休息前及处理产品后立即洗手。立即脱去所有受污染的衣物。",
    "Advice on the protection of the environment": "环境保护建议",
    "See Section 8: Environmental exposure controls": "参见第8节：环境暴露控制",
    "Incompatible products": "不相容产品",
    "No special restrictions on storage with other products": "与其他产品共储无特殊限制。",
    "Storage": "储存",
    "Advice on protection against fire and explosion": "防火防爆建议",
    "Keep away from sources of ignition - No smoking. Take measures to prevent the build up of electrostatic charge. Dust can form an explosive mixture in air. Risks of ignition followed by flame propagation or secondary explosions shall be prevented by avoiding accumulation of dust, e.g. on floors and ledges.":
        "远离火源——禁止吸烟。采取措施防止静电荷积累。粉尘在空气中可能形成爆炸性混合物。应通过避免粉尘积聚（如在地面和台架上）来防止点燃随后火焰蔓延或二次爆炸的风险。",
    "Technical measures/Storage conditions": "技术措施/储存条件",
    "Keep away from direct sunlight. Keep containers tightly closed in a dry, cool and well-ventilated place.":
        "远离阳光直射。保持容器密闭，存放在干燥、阴凉、通风良好的地方。",
    "8. Exposure controls / Personal protection": "第8节：接触控制/个体防护",
    "Exposure limits United States of America": "美国暴露限值",
    "US OSHA Z-1": "美国 OSHA Z-1",
    "US NIOSH IDHL": "美国 NIOSH IDHL",
    "Appropriate Engineering controls": "适当的工程控制",
    "Ensure adequate ventilation. Provide for appropriate exhaust ventilation and dust collection at machinery.":
        "确保充分通风。在机械旁提供适当的排气通风和粉尘收集装置。",
    "Personal protective equipment": "个人防护设备",
    "General industrial hygiene practice": "一般工业卫生实践",
    "Avoid contact with skin, eyes and clothing. Do not breathe dust or mist. Ensure that eyewash stations and safety showers are close to the workstation location.":
        "避免与皮肤、眼睛和衣物接触。不要吸入粉尘或雾气。确保洗眼站和安全淋浴器靠近工位位置。",
    "Respiratory protection": "呼吸防护",
    "Based on workplace contaminant levels and working limits of the respirator, use a respirator approved by NIOSH":
        "根据工作场所污染物水平和呼吸器的工作限制，使用经NIOSH批准的呼吸器。",

    # Page 5
    "Hand protection": "手部防护",
    "Heat resistant gloves.": "耐热手套。",
    "Suitable material leather gloves": "适用材料：皮革手套。",
    "Eye protection": "眼部防护",
    "Tightly fitting safety goggles.": "紧密贴合的安全护目镜。",
    "Skin and body protection": "皮肤和身体防护",
    "Wear face-shield and protective suit for abnormal processing problems.": "在异常加工问题时佩戴面罩和防护服。",
    "Thermal Hazard": "热危险",
    "When handling hot material, use heat resistant gloves. Heat only in areas with appropriate exhaust ventilation.":
        "处理热材料时，使用耐热手套。仅在配备适当排气通风的区域加热。",
    "9. Physical and chemical properties": "第9节：理化特性",
    "9.1. Information on basic physical and chemical properties": "9.1. 基本物理和化学性质信息",
    "Appearance": "外观",
    "granules": "颗粒",
    "Colour": "颜色",
    "colourless": "无色",
    "Odour": "气味",
    "odourless": "无味",
    "Odour threshold": "气味阈值",
    "No data available": "无可用数据",
    "pH": "pH值",
    "Melting point/range": "熔点/范围",
    "approx. 75-185 °C": "约 75-185 °C",
    "Boiling point/range": "沸点/范围",
    "Flash point": "闪点",
    "Evaporation rate": "蒸发速率",
    "Flammability (solid, gas)": "可燃性（固体、气体）",
    "Lower explosion limit": "爆炸下限",
    "Upper explosion limit": "爆炸上限",
    "Vapour pressure": "蒸气压力",
    "< 0.001 mm Hg @25°C (77 F)": "< 0.001 mm Hg @25°C (77°F)",
    "Vapour density": "蒸气密度",
    "Relative density": "相对密度",
    "Bulk density": "堆积密度",
    "550 - 600 g/l": "550 - 600 g/L",
    "Method": "方法",
    "DIN 53466": "DIN 53466",
    "Water solubility": "水溶性",
    "insoluble": "不溶",
    "log Pow": "log Pow（辛醇/水分配系数）",
    "Autoignition temperature": "自燃温度",
    "Decomposition temperature": "分解温度",
    "Viscosity": "粘度",
    "Oxidizing properties": "氧化特性",
    "Does not apply, substance is not oxidising. There are no chemical groups associated with oxidizing properties":
        "不适用，本物质非氧化性物质。不存在与氧化性质相关的化学基团。",
    "9.2. Other information": "9.2. 其他信息",
    "VOC Content(%)": "VOC含量(%)",
    "< 0.5 % (wt/wt)": "< 0.5%（重量/重量）",

    # Page 6
    "10. Stability and reactivity": "第10节：稳定性和反应性",
    "Reactivity": "反应性",
    "The reactivity of the product corresponds to the typical reactivity shown by the substance group as described in any text book on organic chemistry.":
        "本产品的反应性与该类物质在有机化学教科书中所描述的典型反应性相符。",
    "Chemical stability": "化学稳定性",
    "Stable under normal conditions of handling, use and transportation.": "在正常操作、使用和运输条件下稳定。",
    "Possibility of hazardous reactions": "危险反应的可能性",
    "Hazardous polymerisation does not occur.": "不会发生危险聚合。",
    "Conditions to avoid": "应避免的条件",
    "Avoid temperatures above 350 °C / 662°F. Risk of decomposition.": "避免温度超过 350°C / 662°F。有分解风险。",
    "Incompatible materials": "不相容材料",
    "oxidizing agents.": "氧化剂。",
    "Hazardous decomposition products": "有害分解产物",
    "Thermal decomposition can lead to release of irritating gases and vapours.": "热分解可能导致释放刺激性气体和蒸气。",
    "11. Toxicological information": "第11节：毒理学信息",
    "Note": "注释",
    "No toxicology information is available. Handle in accordance with good industrial hygiene and safety practice.":
        "无可用毒理学信息。按照良好的工业卫生和安全实践进行操作。",
    "12. Ecological information": "第12节：生态学信息",
    "Ecotoxicity": "生态毒性",
    "Persistence/Degradability": "持久性/降解性",
    "Bioaccumulative potential": "生物蓄积潜力",

    # Page 7
    "Mobility in soil": "在土壤中的迁移性",
    "Other Adverse Effects": "其他不良效应",
    "No information on ecology is available. According to our experience and to the information provided to us, the product does not have any harmful effects if it is used and handled as specified.":
        "无可用生态学信息。根据我们的经验和向我们提供的信息，如果按照指定方式使用和处理，本产品不会产生任何有害影响。",
    "13. Disposal considerations": "第13节：废弃处置",
    "Product Information": "产品信息",
    "Where possible recycling is preferred to disposal or incineration. May be taken to waste disposal site or incineration plant, with household waste. Rules of the local authorities must be observed.":
        "在可行的情况下，回收优于处置或焚烧。可作为生活垃圾送往废物处置场或焚烧厂。必须遵守当地主管部门的规定。",
    "Uncleaned empty packaging": "未清洁的空包装",
    "Regulations concerning reuse or disposal of used packaging materials must be observed.": "必须遵守有关使用过的包装材料再利用或处置的规定。",
    "14. Transport information": "第14节：运输信息",
    "ICAO-TI / IATA-DGR": "国际民航组织-技术指令 / 国际航空运输协会-危险品规则",
    "Not restricted": "不受限制",
    "IMDG": "国际海运危险货物规则",
    "D.O.T. (49CFR)": "美国运输部 (49CFR)",
    "15. Regulatory information": "第15节：法规信息",
    "OSHA Regulatory Status": "OSHA 法规状态",
    "This material is non-hazardous as defined by the American OSHA Hazard Communication Standard (29CFR 1910.1200).":
        "根据美国 OSHA 危害通报标准 (29CFR 1910.1200)，本材料属于非危险品。",
    "Federal Regulations": "联邦法规",
    "This product complies with U.S. Toxic Substance Control Act (TSCA)": "本产品符合美国《有毒物质控制法》(TSCA)",
    "It is listed on the active inventory list via the 2018 reset rule": "根据2018年重置规则列入有效库存清单。",
    "International Inventories": "国际化学品名录",

    # Page 8
    "The products covered by this SDS are permitted under the following inventories:": "本安全数据说明书涵盖的产品被列入以下化学品名录：",
    "TSCA (US)": "TSCA（美国）",
    "AICS (AU)": "AICS（澳大利亚）",
    "KECI (KR)": "KECI（韩国）",
    "DSL (CA)": "DSL（加拿大）",
    "IECSC (CN)": "IECSC（中国）",
    "EINECS (EU)": "EINECS（欧盟）",
    "ENCS (JP)": "ENCS（日本）",
    "NZIoC (NZ)": "NZIoC（新西兰）",
    "TCSI (TW)": "TCSI（台湾地区）",
    "16. Other information": "第16节：其他信息",
    "Training advice": "培训建议",
    "For effective first-aid, special training / education is needed.": "为进行有效的急救，需要特殊培训/教育。",
    "Hazard Rating Systems": "危险等级系统",
    "NFPA (National Fire Protection Association)": "NFPA（美国国家消防协会）",
    "Health Hazard": "健康危害",
    "Fire Hazard": "火灾危险",
    "Reactivity": "反应性",
    "HMIS (Hazardous Material Information System)": "HMIS（危险物质信息系统）",
    "Flammability": "可燃性",
    "Physical Hazard": "物理危害",
    "Sources of key data used to compile the datasheet": "编制本数据表所用的关键数据来源",
    "Information contained in this safety data sheet is based on TOPAS owned data and public sources deemed valid or acceptable. The absence of data elements required by OSHA, ANSI or Annex II, Regulation 1907/2006/EC indicates, that no data meeting these requirements is available.":
        "本安全数据说明书中所含信息基于TOPAS自有数据和被认为有效或可接受的公开来源。缺少OSHA、ANSI或法规1907/2006/EC附件II所要求的数据要素，表明没有满足这些要求的可用数据。",
    "Further information for the safety data sheet": "安全数据说明书的更多信息",
    "For more information, consult the Technical Data Sheet (www.topas.com). Changes against the previous version are marked by ***.":
        "更多信息请参阅技术数据表 (www.topas.com)。与上一版本相比的变更以 *** 标记。",
    "Disclaimer": "免责声明",
    "The information contained herein is accurate to the best of our knowledge. We do not suggest or guarantee that any hazards listed herein are the only ones which exist. TOPAS Advanced Polymers, Inc. makes no warranty of any kind, express or implied, concerning the safe use of this material in your process or in combination with other substances. Effects can be aggravated by other materials and/or this material may aggravate or add to the effects of other materials. User has sole responsibility to determine the suitability of the materials for any use and the manner of use contemplated. User must meet all applicable safety and health standards":
        "本文件所含信息据我们所知是准确的。我们不表示或保证此处列出的危害是唯一存在的危害。TOPAS Advanced Polymers, Inc. 对在其工艺中或与其他物质组合安全使用本材料不作任何明示或暗示的保证。其他材料可能加重影响，和/或本材料可能加重或增加其他材料的影响。用户有唯一责任确定材料对任何用途的适用性及预期的使用方式。用户必须符合所有适用的安全和健康标准。",
}

def translate_line(line):
    """查找翻译，如果没找到则保留原文"""
    stripped = line.strip()
    if stripped in TRANSLATIONS:
        return TRANSLATIONS[stripped]
    # 尝试大小写不敏感匹配
    for en, zh in TRANSLATIONS.items():
        if en.lower().strip() == stripped.lower().strip():
            return zh
    return None


def create_bilingual_docx(json_path, output_path):
    """创建中英文对照的格式化Word文档"""
    with open(json_path, 'r', encoding='utf-8') as f:
        pages = json.load(f)

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ---- 封面/标题 ----
    title = doc.add_heading('安全数据说明书（中英文对照）', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('SAFETY DATA SHEET — TOPAS® Cyclic Olefin Copolymers')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('产品代码: TP01 | 版本: 8 | 修订日期: 2019-06-22 | 签发日期: 2019-08-05').font.size = Pt(9)

    doc.add_paragraph()  # 空行

    # ---- 逐页处理 ----
    for page_data in pages:
        page_num = page_data['page']
        text = page_data['text']
        lines = text.split('\n')

        # 页标题
        if page_num > 1:
            doc.add_page_break()

        page_header = doc.add_paragraph()
        run = page_header.add_run(f'第 {page_num} 页 / 共 8 页')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        page_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 解析每一行并生成翻译
        for line in lines:
            stripped = line.strip()

            # 跳过分隔线和页码标识
            if stripped.startswith('___') or stripped.startswith('---'):
                doc.add_paragraph().add_run('─' * 60).font.size = Pt(6)
                continue

            if stripped.endswith('/ 8 USA/ EN') or stripped.endswith('/ 8 USA/EN'):
                continue

            if not stripped:
                doc.add_paragraph()  # 空行
                continue

            # 跳过重复的页眉
            if stripped in ['SAFETY DATA SHEET', 'TOPAS® Cyclic Olefin Copolymers']:
                continue
            if stripped.startswith('Revision Date') or stripped.startswith('Version / Revision') or stripped.startswith('Supersedes'):
                continue

            zh = translate_line(stripped)

            # 检测是否为节标题（如 "1: Identification" 格式）
            is_section_header = False
            for prefix in ['1:', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.', '13.', '14.', '15.', '16.']:
                if stripped.startswith(prefix):
                    is_section_header = True
                    break

            para = doc.add_paragraph()

            if is_section_header:
                # 节标题：粗体大字
                run_en = para.add_run(f'{stripped}')
                run_en.bold = True
                run_en.font.size = Pt(12)
                run_en.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
                if zh:
                    para2 = doc.add_paragraph()
                    run_zh = para2.add_run(f'  → {zh}')
                    run_zh.bold = True
                    run_zh.font.size = Pt(11)
                    run_zh.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            elif zh:
                # 中英文对照
                run_en = para.add_run(f'{stripped}')
                run_en.font.size = Pt(9)
                run_en.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                para2 = doc.add_paragraph()
                run_zh = para2.add_run(f'  {zh}')
                run_zh.font.size = Pt(10)

            else:
                # 仅原文（无对应翻译）
                run_en = para.add_run(f'{stripped}')
                run_en.font.size = Pt(9)
                run_en.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # 表格处理（如果有）
        for table in page_data.get('tables', []):
            if table:
                doc.add_paragraph()
                tbl_title = doc.add_paragraph()
                run = tbl_title.add_run('[表格]')
                run.font.size = Pt(9)
                run.italic = True

                rows = len(table)
                cols = max(len(row) for row in table) if table else 0
                tbl = doc.add_table(rows=rows, cols=cols, style='Table Grid')
                for i, row in enumerate(table):
                    for j, cell_text in enumerate(row):
                        if j < cols:
                            zh_cell = translate_line(cell_text)
                            if zh_cell:
                                tbl.cell(i, j).text = f'{cell_text}\n{zh_cell}'
                            else:
                                tbl.cell(i, j).text = cell_text

    # ---- 保存 ----
    doc.save(output_path)
    print(f'✅ 翻译文档已保存到: {output_path}')


if __name__ == '__main__':
    json_path = r'd:\translation\msds_extracted.json'
    output_path = r'd:\translation\MSDS_TOPAS_中英文对照.docx'
    create_bilingual_docx(json_path, output_path)
