"""
生成干净的中文MSDS Word文档 — 标准SDS格式 + API翻译
"""
import sys, io, json, hashlib
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import pymupdf
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def text_key(text):
    return hashlib.md5(text.strip().encode()).hexdigest()[:12]

def load_cache(path):
    import os
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def get_zh(text, cache):
    k = text_key(text)
    zh = cache.get(k)
    return zh if zh and zh != text.strip() else None

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_para(doc, text, bold=False, size=10.5, color=None, align=None, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    return para

def gen_msds_docx(pdf_path, cache_path, output_path):
    """生成标准SDS格式的中文Word文档"""
    cache = load_cache(cache_path)
    doc = Document()

    # 页面设置 A4
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ====== 封面标题区 ======
    h = doc.add_heading('安全数据说明书', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.size = Pt(18)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 产品标识
    add_formatted_para(doc, 'TOPAS® 环烯烃共聚物', bold=True, size=14,
                       color=(0x1a, 0x56, 0x8e), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_para(doc, '产品代码：TP01　　修订日期：2019-06-22　　版本：8', size=9,
                       color=(0x88, 0x88, 0x88), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()  # 空行

    # ====== 读取PDF页面翻译内容 ======
    doc_pdf = pymupdf.open(pdf_path)
    page_blocks = {}  # page -> [(text, zh, font_size, is_bold), ...]

    for page_num in range(len(doc_pdf)):
        page = doc_pdf[page_num]
        blocks = page.get_text('dict')['blocks']
        items = []
        for block in blocks:
            if block['type'] != 0: continue
            for line in block['lines']:
                original = ' '.join([s['text'] for s in line['spans']]).strip()
                if not original or len(original) < 2: continue
                # 跳过水印和页眉
                if original in ['SAFETY DATA SHEET', 'TOPAS® Cyclic Olefin Copolymers']: continue
                if original.startswith(('Revision Date', 'Version / Revision', 'Supersedes',
                                        '________________________________________________________________')): continue
                if any(x in original for x in ['Downloaded for', 'For personal use', 'ClinicalKey',
                                               'All rights reserved', 'Elsevier']): continue
                if original.endswith('/ 8') or 'USA/ EN' in original: continue
                zh = get_zh(original, cache) or original
                fs = line['spans'][0]['size'] if line['spans'] else 10
                bold = bool(line['spans'][0].get('flags', 0) & 2) if line['spans'] else False
                items.append((original, zh, fs, bold))
        page_blocks[page_num] = items

    # ====== 逐节输出正文 ======
    # MSDS 有16个章节
    section_heads = [
        '第1节：化学品及企业标识', '第2节：危险性概述', '第3节：成分/组成信息',
        '第4节：急救措施', '第5节：消防措施', '第6节：泄漏应急处理',
        '第7节：操作处置与储存', '第8节：接触控制/个体防护', '第9节：理化特性',
        '第10节：稳定性和反应性', '第11节：毒理学信息', '第12节：生态学信息',
        '第13节：废弃处置', '第14节：运输信息', '第15节：法规信息', '第16节：其他信息'
    ]

    # 遍历所有页面的文本，按章节组织
    last_was_header = False
    current_para = ''

    for page_num in sorted(page_blocks.keys()):
        for original, zh, fs, bold in page_blocks[page_num]:
            # 检测章节标题
            is_section = False
            for i, sh in enumerate(section_heads):
                if zh.startswith(f'第{i+1}节') or (\
                   original.strip().startswith(f'{i+1}:') or original.strip().startswith(f'{i+1}.')):
                    is_section = True

            # 写之前的段落
            if current_para and (is_section or (bold and fs > 11)):
                add_formatted_para(doc, current_para, size=10.5)
                current_para = ''

            if is_section:
                doc.add_paragraph()  # 章节前空行
                add_formatted_para(doc, zh, bold=True, size=13, color=(0x1a, 0x56, 0x8e))
                last_was_header = True
            elif (bold or fs > 11) and len(zh) < 80:
                add_formatted_para(doc, zh, bold=True, size=11)
                last_was_header = True
            else:
                if zh != original:  # 已翻译
                    if last_was_header:
                        current_para = zh
                    elif current_para:
                        current_para += zh
                    else:
                        current_para = zh
                    last_was_header = False
                else:  # 原文保留（产品代码等）
                    if current_para:
                        add_formatted_para(doc, current_para, size=10.5)
                        current_para = ''
                    add_formatted_para(doc, original, size=9, color=(0x66, 0x66, 0x66))

    # 写最后一段
    if current_para:
        add_formatted_para(doc, current_para, size=10.5)

    # ====== 免责声明 ======
    doc.add_paragraph()
    add_formatted_para(doc, '免责声明', bold=True, size=11, color=(0x1a, 0x56, 0x8e))
    add_formatted_para(doc,
        '本文件所含信息据我们所知是准确的。我们不表示或保证此处列出的危害是唯一存在的危害。'
        'TOPAS Advanced Polymers, Inc. 对在其工艺中或与其他物质组合安全使用本材料不作任何明示或暗示的保证。'
        '用户有唯一责任确定材料对任何用途的适用性及预期使用方式。用户必须符合所有适用的安全和健康标准。',
        size=9, color=(0x77, 0x77, 0x77))

    doc.save(output_path)
    print(f'Saved: {output_path}')

if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('pdf')
    ap.add_argument('-c', '--cache')
    ap.add_argument('-o', '--output')
    args = ap.parse_args()

    gen_msds_docx(
        args.pdf,
        args.cache or args.pdf.rsplit('.', 1)[0] + '_cache.json',
        args.output or args.pdf.rsplit('.', 1)[0] + '_SDS_中文版.docx'
    )
