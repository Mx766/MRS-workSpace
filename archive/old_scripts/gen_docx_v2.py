"""
PDF → 格式化中文 DOCX v2 — 智能解析结构保留层级
"""
import sys, io, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import pymupdf
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

def load_dict(path):
    if not path: return {}
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)

def get_zh(text, d):
    t = text.strip()
    if t in d: return d[t]
    for en, zh in d.items():
        if len(en) > 15 and en in t:
            return zh
    return None

SECTION_PREFIXES = [
    '1:', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.',
    '10.', '11.', '12.', '13.', '14.', '15.', '16.',
]

def is_section_header(text):
    for p in SECTION_PREFIXES:
        if text.strip().startswith(p):
            return True
    return False

def is_sub_header(text):
    """二级标题：数字编号、短文本、首字母大写"""
    t = text.strip()
    if len(t) > 40: return False
    if t in ['SAFETY DATA SHEET', 'TOPAS® Cyclic Olefin Copolymers']: return False
    # 如 "2.1. Classification..." 或 "Product Identifier"
    if any(t.startswith(p) for p in ['2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.', '13.', '14.', '15.', '16.']):
        return '.' in t[:6]
    # 或全大写短标题
    if t.isupper() and len(t) > 3: return True
    return False

SKIP_LINES = ['________________________________________________________________', 'Downloaded for', 'For personal use',
              'All rights reserved', 'ClinicalKey', 'Elsevier', 'Clin Plastic Surg']

def should_skip(t):
    t = t.strip()
    if not t or len(t) < 2: return True
    for p in SKIP_LINES:
        if p.lower() in t.lower(): return True
    if t.endswith('/ 8') or t.endswith('USA/ EN'): return True
    return False

def gen_docx(pdf_path, output_path, dict_data, title=""):
    doc_pdf = pymupdf.open(pdf_path)
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    if title:
        h = doc.add_heading(title, level=0)
    else:
        h = doc.add_heading(os.path.basename(pdf_path).replace('.pdf',' 中文版'), level=0)

    current_para_text = ""
    last_was_header = False

    for page_num in range(len(doc_pdf)):
        page = doc_pdf[page_num]
        blocks = page.get_text('dict')['blocks']

        if page_num > 0:
            doc.add_page_break()

        for block in blocks:
            if block['type'] != 0: continue

            for line in block['lines']:
                spans = line['spans']
                original = ' '.join([s['text'] for s in spans]).strip()

                if should_skip(original): continue
                if original in ['SAFETY DATA SHEET', 'TOPAS® Cyclic Olefin Copolymers']: continue
                if original.startswith('Revision Date') or original.startswith('Version / Revision'): continue

                font_size = spans[0]['size'] if spans else 10
                zh = get_zh(original, dict_data)
                display_text = zh if zh else original

                # 检测标题层级
                if is_section_header(original):
                    if current_para_text:
                        doc.add_paragraph(current_para_text)
                        current_para_text = ""
                    h = doc.add_heading(display_text, level=2)
                    last_was_header = True
                    continue

                if is_sub_header(original) and not zh:
                    # 原文副标题（无翻译），保留原样
                    if current_para_text:
                        doc.add_paragraph(current_para_text)
                        current_para_text = ""
                    para = doc.add_paragraph()
                    run = para.add_run(original)
                    run.bold = True
                    run.font.size = Pt(11)
                    last_was_header = True
                    continue

                if is_sub_header(original) and zh:
                    # 已翻译的副标题
                    if current_para_text:
                        doc.add_paragraph(current_para_text)
                        current_para_text = ""
                    para = doc.add_paragraph()
                    run = para.add_run(display_text)
                    run.bold = True
                    run.font.size = Pt(11)
                    last_was_header = True
                    continue

                # 正文内容
                if last_was_header:
                    current_para_text = display_text
                    last_was_header = False
                else:
                    if current_para_text:
                        current_para_text += display_text
                    else:
                        current_para_text = display_text

        # 页面结束，输出当前段落
        if current_para_text:
            doc.add_paragraph(current_para_text)
            current_para_text = ""

    doc.save(output_path)
    print(f'Saved: {output_path}')

if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('pdf')
    ap.add_argument('-o', '--output')
    ap.add_argument('-d', '--dict')
    ap.add_argument('-t', '--title', default='')
    args = ap.parse_args()

    d = load_dict(args.dict)
    gen_docx(args.pdf, args.output or args.pdf.rsplit('.',1)[0]+'_中文版.docx', d, args.title)
