"""
PDF → 翻译DOCX — 提取PDF文本+表格+格式 → 生成纯中文Word
Word可直接另存为PDF，保持格式一致
"""
import sys, io, json, os, hashlib
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import pdfplumber
import pymupdf
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def load_translations(dict_path):
    """加载翻译字典"""
    if not dict_path or not os.path.exists(dict_path):
        return {}
    with open(dict_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    # 构建 key→zh 映射
    result = {}
    for en, zh in data.items():
        result[en] = zh
    return result

def get_zh(text, translations):
    """查找翻译"""
    t = text.strip()
    if not t:
        return None
    if t in translations:
        return translations[t]
    # 尝试模糊匹配
    for en, zh in translations.items():
        if len(en) > 15 and (en in t or t in en):
            return zh
    return None

SKIP_PATTERNS = [
    '________________________________________________________________',
    'Downloaded for',
    'For personal use only',
    'All rights reserved',
    'Copyright',
    'ClinicalKey',
    'Elsevier',
]

def should_skip(text):
    for p in SKIP_PATTERNS:
        if p.lower() in text.lower():
            return True
    return False

def convert_pdf_to_docx(pdf_path, output_path, translations=None):
    """将PDF翻译后输出为格式化Word文档"""
    if translations is None:
        translations = {}

    # 用 PyMuPDF 提取以保留更多格式信息
    doc_pdf = pymupdf.open(pdf_path)
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题
    heading = doc.add_heading('', level=0)
    run = heading.add_run(os.path.basename(pdf_path).replace('.pdf', ' 中文版'))
    run.font.size = Pt(14)

    stats = {'translated': 0, 'original': 0, 'skipped': 0}

    for page_num in range(len(doc_pdf)):
        page = doc_pdf[page_num]
        blocks = page.get_text('dict')['blocks']

        if page_num > 0:
            doc.add_page_break()

        for block in blocks:
            # 处理文本块
            if block['type'] == 0:
                for line in block['lines']:
                    # 收集行内文本
                    original_text = ' '.join([s['text'] for s in line['spans']]).strip()
                    if not original_text or should_skip(original_text):
                        stats['skipped'] += 1
                        continue

                    # 检测是否为分隔线
                    if len(original_text) < 3 and not original_text.isalnum():
                        continue

                    zh = get_zh(original_text, translations)

                    if zh:
                        # 只用中文
                        para = doc.add_paragraph()
                        run = para.add_run(zh)

                        # 尝试保留一些格式
                        first_span = line['spans'][0] if line['spans'] else None
                        if first_span:
                            if first_span['size'] > 12:
                                run.bold = True
                                run.font.size = Pt(first_span['size'] * 0.85)
                            elif first_span['size'] > 10:
                                run.font.size = Pt(11)
                            else:
                                run.font.size = Pt(10)

                        stats['translated'] += 1
                    else:
                        # 保留原文（产品代码、数字等）
                        para = doc.add_paragraph()
                        run = para.add_run(original_text)
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                        stats['original'] += 1

            # 处理表格
            elif block['type'] == 1:
                # PyMuPDF 表格块（如果检测到）
                pass

    doc.save(output_path)
    print(f'Translated: {stats["translated"]} segments')
    print(f'Kept original: {stats["original"]} segments')
    print(f'Skipped: {stats["skipped"]} lines')
    print(f'Saved: {output_path}')

if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('pdf', help='Input PDF file')
    ap.add_argument('-o', '--output', help='Output DOCX path')
    ap.add_argument('-d', '--dict', help='Translation dictionary JSON')
    args = ap.parse_args()

    if not args.output:
        args.output = args.pdf.rsplit('.', 1)[0] + '_中文版.docx'

    translations = load_translations(args.dict) if args.dict else {}
    print(f'Dictionary: {len(translations)} entries')
    convert_pdf_to_docx(args.pdf, args.output, translations)
