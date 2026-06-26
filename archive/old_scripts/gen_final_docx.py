"""
生成格式化中文 DOCX — 使用 API 翻译缓存 + 保持层级结构
"""
import sys, io, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import pymupdf, hashlib, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def text_key(text):
    return hashlib.md5(text.strip().encode()).hexdigest()[:12]

def load_cache(path):
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def get_zh(text, cache):
    k = text_key(text)
    zh = cache.get(k)
    return zh if zh and zh != text.strip() else None

SECTION_NUMS = [f'{i}{s}' for i in range(1,17) for s in [':','.']] + \
              [f'{i}.{j}' for i in range(1,17) for j in range(1,10)]

SKIP = ['Downloaded for', 'For personal use', 'All rights reserved', 'ClinicalKey',
        'Elsevier', 'Clin Plastic Surg', '________________________________________________________________']

def should_skip(t):
    t = t.strip()
    if not t: return True
    for p in SKIP:
        if p.lower() in t.lower(): return True
    if t.endswith('/ 8') or 'USA/ EN' in t: return True
    return False

def gen_docx(pdf_path, cache_path, output_path, title=''):
    doc_pdf = pymupdf.open(pdf_path)
    cache = load_cache(cache_path)
    doc = Document()

    # 页面设置
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    if title:
        h = doc.add_heading(title, level=0)

    for page_num in range(len(doc_pdf)):
        page = doc_pdf[page_num]
        blocks = page.get_text('dict')['blocks']

        if page_num > 0:
            doc.add_page_break()

        para_texts = []
        for block in blocks:
            if block['type'] != 0: continue
            for line in block['lines']:
                spans = line['spans']
                original = ' '.join([s['text'] for s in spans]).strip()
                if should_skip(original): continue

                # 跳过重复页眉
                if original in ['SAFETY DATA SHEET', 'TOPAS® Cyclic Olefin Copolymers']:
                    continue
                if original.startswith(('Revision Date', 'Version / Revision', 'Supersedes')):
                    continue

                zh = get_zh(original, cache)
                display = zh if zh else original

                font_size = spans[0]['size'] if spans else 10
                is_bold = bool(spans[0].get('flags', 0) & 2) if spans else False

                # 检测 section header
                is_sec = any(original.strip().startswith(p) for p in SECTION_NUMS)

                para = doc.add_paragraph()

                if is_sec:
                    run = para.add_run(display)
                    run.bold = True
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
                elif is_bold or (font_size > 11 and len(original) < 60):
                    run = para.add_run(display)
                    run.bold = True
                    run.font.size = Pt(11)
                elif zh:
                    run = para.add_run(display)
                    run.font.size = Pt(10.5)
                else:
                    run = para.add_run(display)
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(output_path)
    print(f'Saved: {output_path}')

if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('pdf')
    ap.add_argument('-c', '--cache')
    ap.add_argument('-o', '--output')
    ap.add_argument('-t', '--title', default='')
    args = ap.parse_args()

    if not args.cache:
        args.cache = args.pdf.rsplit('.', 1)[0] + '_cache.json'
    if not args.output:
        args.output = args.pdf.rsplit('.', 1)[0] + '_中文版.docx'

    gen_docx(args.pdf, args.cache, args.output, args.title)
