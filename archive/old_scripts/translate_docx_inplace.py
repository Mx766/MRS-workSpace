"""
DOCX 原位翻译 — 保留所有格式（字体/大小/颜色/粗体/斜体），只替换文字
"""
import sys, io, json, hashlib, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt

def text_key(text):
    return hashlib.md5(text.strip().encode()).hexdigest()[:12]

def load_cache(path):
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def translate_docx(input_path, cache_path, output_path):
    cache = load_cache(cache_path)
    doc = Document(input_path)

    translated = 0
    kept = 0

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text.strip()
            if not text or len(text) <= 1:
                continue

            # 尝试精确匹配翻译
            k = text_key(text)
            zh = cache.get(k)

            if zh and zh != text:
                run.text = zh
                translated += 1
            else:
                kept += 1

    # 也处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if not text or len(text) <= 1:
                            continue
                        k = text_key(text)
                        zh = cache.get(k)
                        if zh and zh != text:
                            run.text = zh
                            translated += 1
                        else:
                            kept += 1

    doc.save(output_path)
    print(f'Translated: {translated} runs, Kept: {kept} runs')
    print(f'Saved: {output_path}')

if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('-c', '--cache')
    ap.add_argument('-o', '--output')
    args = ap.parse_args()

    translate_docx(args.docx, args.cache, args.output or args.docx.replace('.docx', '_中文版.docx'))
