#!/usr/bin/env python3
"""Build annotated DOCX for proofread/1 — v2 (51 issues from updated checklist)"""
import sys, json, os, tempfile, zipfile, re as regex
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8')

NS_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
COMMENTS_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
CT_COMMENTS = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
XML_NS = 'http://www.w3.org/XML/1998/namespace'

with open('cache/issues_phase4_v2_1.json', 'r', encoding='utf-8') as f:
    issues = json.load(f).get('issues', [])

doc = Document('D:/translation/proofread/1/[1]-CN.docx')
hl_colors = {'critical': ('FFCCCC', 'red'), 'medium': ('FFF9C4', 'yellow'), 'low': ('C8E6C9', 'green')}
sev_order = {'critical': 0, 'medium': 1, 'low': 2}

issues_by_para = {}
for issue in issues:
    pi = issue.get('paragraph_index', 0)
    if pi < 1: continue
    issues_by_para.setdefault(pi, []).append(issue)

# STEP 1: Highlights via python-docx
for para_idx in sorted(issues_by_para.keys()):
    if para_idx > len(doc.paragraphs): continue
    para = doc.paragraphs[para_idx - 1]
    sorted_issues = sorted(issues_by_para[para_idx],
                           key=lambda x: sev_order.get(x.get('severity', 'medium'), 1))
    for issue in sorted_issues:
        sev = issue.get('severity', 'medium')
        fill, hl_name = hl_colors.get(sev, ('FFFF00', 'yellow'))
        target = issue.get('target_quote', '').strip()
        if target:
            for run in para.runs:
                if not run.text.strip(): continue
                key = target[:30].strip()
                if key and key in run.text:
                    rPr = run._r.get_or_add_rPr()
                    if rPr.find(qn('w:shd')) is not None: continue
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), fill)
                    rPr.append(shd)
                    hl = OxmlElement('w:highlight')
                    hl.set(qn('w:val'), hl_name)
                    rPr.append(hl)

# Save via python-docx
tmpfd, tmppath = tempfile.mkstemp(suffix='.docx')
os.close(tmpfd)
doc.save(tmppath)

# STEP 2: ZIP/lxml for comments
entries = {}
with zipfile.ZipFile(tmppath, 'r') as zf:
    for info in zf.infolist():
        entries[info.filename] = (info, zf.read(info.filename))

doc_xml = entries['word/document.xml'][1]
doc_tree = etree.fromstring(doc_xml)
body = doc_tree.find(f'{{{NS_w}}}body')
paragraphs = body.findall(f'{{{NS_w}}}p')

comments_tree = etree.Element(f'{{{NS_w}}}comments', nsmap={'w': NS_w})
comment_id = 0

for para_idx in sorted(issues_by_para.keys()):
    if para_idx > len(paragraphs): continue
    para_elem = paragraphs[para_idx - 1]
    sorted_issues = sorted(issues_by_para[para_idx],
                           key=lambda x: sev_order.get(x.get('severity', 'medium'), 1))

    para_comments = []
    for issue in sorted_issues:
        sev = issue.get('severity', 'medium')
        sev_label = {'critical': '🔴 严重（必须修改）', 'medium': '🟡 中等（建议修改）',
                     'low': '🟢 低（格式优化）'}.get(sev, '⚪ 未分级')
        para_comments.append((comment_id, sev, sev_label, issue))
        comment_id += 1

    # Insert crs in ID order
    pPr = para_elem.find(f'{{{NS_w}}}pPr')
    insert_after = pPr
    for cid, sev, sev_label, issue in para_comments:
        crs = etree.Element(f'{{{NS_w}}}commentRangeStart')
        crs.set(f'{{{NS_w}}}id', str(cid))
        if insert_after is not None:
            idx = list(para_elem).index(insert_after)
            para_elem.insert(idx + 1, crs)
        else:
            para_elem.insert(0, crs)
        insert_after = crs

    for cid, sev, sev_label, issue in para_comments:
        cre = etree.Element(f'{{{NS_w}}}commentRangeEnd')
        cre.set(f'{{{NS_w}}}id', str(cid))
        para_elem.append(cre)

    for cid, sev, sev_label, issue in para_comments:
        ref_run = etree.SubElement(para_elem, f'{{{NS_w}}}r')
        ref_rPr = etree.SubElement(ref_run, f'{{{NS_w}}}rPr')
        cr = etree.SubElement(ref_rPr, f'{{{NS_w}}}commentReference')
        cr.set(f'{{{NS_w}}}id', str(cid))

    for cid, sev, sev_label, issue in para_comments:
        date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
        cmt = etree.SubElement(comments_tree, f'{{{NS_w}}}comment')
        cmt.set(f'{{{NS_w}}}id', str(cid))
        cmt.set(f'{{{NS_w}}}author', 'TranslationReview')
        cmt.set(f'{{{NS_w}}}date', date_str)
        cmt.set(f'{{{NS_w}}}initials', 'TR')

        lines = [sev_label,
                 f'【{issue.get("dimension","")} — {issue.get("check_item","")}】', '',
                 f'原文: {issue.get("source_quote","")}', '',
                 f'问题: {issue.get("issue","")}', '',
                 f'建议: {issue.get("suggestion","")}']
        for i, line in enumerate(lines):
            pe = etree.SubElement(cmt, f'{{{NS_w}}}p')
            pp = etree.SubElement(pe, f'{{{NS_w}}}pPr')
            r_e = etree.SubElement(pe, f'{{{NS_w}}}r')
            if i == 0:
                rp = etree.SubElement(r_e, f'{{{NS_w}}}rPr')
                etree.SubElement(rp, f'{{{NS_w}}}b')
                sz = etree.SubElement(rp, f'{{{NS_w}}}sz')
                sz.set(f'{{{NS_w}}}val', '18')
                co = etree.SubElement(rp, f'{{{NS_w}}}color')
                co.set(f'{{{NS_w}}}val', {'critical':'FF0000','medium':'FF8C00','low':'228B22'}.get(sev,'000000'))
            te = etree.SubElement(r_e, f'{{{NS_w}}}t')
            te.set(f'{{{XML_NS}}}space', 'preserve')
            te.text = line

# Serialize
entries['word/document.xml'] = (entries['word/document.xml'][0],
    etree.tostring(doc_tree, xml_declaration=True, encoding='UTF-8', standalone=True))
entries['word/comments.xml'] = (None, etree.tostring(comments_tree, xml_declaration=True, encoding='UTF-8', standalone=True))

# Update rels
rels_path = 'word/_rels/document.xml.rels'
rels_data = entries[rels_path][1]
rels_tree = etree.fromstring(rels_data)
max_rid = 0
has_rel = False
for rel in rels_tree.iter():
    if rel.tag == f'{{{REL_NS}}}Relationship':
        rid = rel.get('Id', '')
        m = regex.match(r'rId(\d+)', rid)
        if m: max_rid = max(max_rid, int(m.group(1)))
        if rel.get('Type') == COMMENTS_NS: has_rel = True
if not has_rel:
    nr = etree.SubElement(rels_tree, f'{{{REL_NS}}}Relationship')
    nr.set('Id', f'rId{max_rid + 1}')
    nr.set('Type', COMMENTS_NS)
    nr.set('Target', 'comments.xml')
entries[rels_path] = (entries[rels_path][0], etree.tostring(rels_tree, xml_declaration=True, encoding='UTF-8', standalone=True))

# Update content types
ct_path = '[Content_Types].xml'
ct_data = entries[ct_path][1]
ct_tree = etree.fromstring(ct_data)
has_ct = False
for ov in ct_tree.iter():
    if ov.tag == f'{{{CT_NS}}}Override' and ov.get('PartName') == '/word/comments.xml':
        has_ct = True; break
if not has_ct:
    nct = etree.SubElement(ct_tree, f'{{{CT_NS}}}Override')
    nct.set('PartName', '/word/comments.xml')
    nct.set('ContentType', CT_COMMENTS)
entries[ct_path] = (entries[ct_path][0], etree.tostring(ct_tree, xml_declaration=True, encoding='UTF-8', standalone=True))

# Write
os.makedirs('D:/translation/proofread/1/v2', exist_ok=True)
output = 'D:/translation/proofread/1/v2/[1]-CN_校对稿.docx'
with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
    for name, (info, data) in entries.items():
        if info is not None:
            zf.writestr(info, data)
        else:
            zf.writestr(name, data)
os.unlink(tmppath)

# Verify
orig = Document('D:/translation/proofread/1/[1]-CN.docx')
anno = Document(output)
diffs = sum(1 for o, a in zip((p.text for p in orig.paragraphs), (p.text for p in anno.paragraphs)) if o != a)
print(f'OK: {comment_id} comments | {diffs} text diffs | saved to {output}')
