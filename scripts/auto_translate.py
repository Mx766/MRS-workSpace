"""
通用文档翻译工具 — 提取 → API翻译 → 生成格式化中英文对照Word
使用 Anthropic 兼容 API 进行翻译，支持断点续传
"""
import sys, io, json, os, time
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import httpx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ═══════════════════════════════════════
# API 配置 — 优先环境变量，回退到 config.py
# ═══════════════════════════════════════
API_KEY = os.environ.get('ANTHROPIC_AUTH_TOKEN', '')
BASE_URL = os.environ.get('ANTHROPIC_BASE_URL', '')
MODEL = os.environ.get('ANTHROPIC_MODEL', '')

if not API_KEY:
    try:
        sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'app'))
        from config import load_config
        cfg = load_config()
        API_KEY = cfg.get('api_key', '')
        BASE_URL = cfg.get('base_url_override', '') or 'https://api.deepseek.com'
        # Use OpenAI-compatible endpoint
        ANTHROPIC_API = False
    except Exception:
        BASE_URL = BASE_URL or 'https://api.deepseek.com/anthropic'
        MODEL = MODEL or 'deepseek-v4-pro'
        ANTHROPIC_API = True
else:
    BASE_URL = BASE_URL or 'https://api.deepseek.com/anthropic'
    MODEL = MODEL or 'deepseek-v4-pro'
    ANTHROPIC_API = True

def translate_via_api(text):
    """通过 Anthropic 兼容 API 翻译"""
    if not text.strip():
        return ""
    try:
        resp = httpx.post(
            f"{BASE_URL}/v1/messages",
            headers={
                "x-api-key": API_KEY,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json={
                "model": "deepseek-v4-pro",
                "max_tokens": 4096,
                "temperature": 0.1,
                "system": "你是一个专业翻译。将英文准确翻译为中文，保留专业术语、数字、单位、引用格式。只输出译文，不加解释。",
                "messages": [
                    {"role": "user", "content": f"将以下英文翻译为中文。只输出译文：\n\n{text}"}
                ]
            },
            timeout=60
        )
        data = resp.json()
        if 'content' in data:
            return data['content'][0]['text'].strip()
        else:
            print(f"  [API: {data.get('error', {}).get('message', 'unknown')}]")
            return text
    except Exception as e:
        print(f"  [API Error: {e}]")
        return text

# 需要跳过的行（水印、页眉等非翻译内容）
SKIP_PATTERNS = [
    "Downloaded for Anonymous User",
    "For personal use only",
    "All rights reserved",
    "Clin Plastic Surg",
    "Copyright",
    "Elsevier Inc",
    "ClinicalKey.com",
]

def should_skip(text):
    for pat in SKIP_PATTERNS:
        if pat.lower() in text.lower():
            return True
    return False

def translate_chunk(text):
    """调用 API 翻译"""
    return translate_via_api(text)

def split_text(text, max_chars=1500):
    """将文本拆分为适合翻译的段落块"""
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    chunks = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) < max_chars:
            current += " " + para if current else para
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)
    return chunks

def create_translated_docx(json_path, output_path, cache_path=None):
    """生成中英文对照的 Word 文档"""
    with open(json_path, 'r', encoding='utf-8') as f:
        pages = json.load(f)

    # 加载翻译缓存
    cache = {}
    if cache_path and os.path.exists(cache_path):
        with open(cache_path, 'r', encoding='utf-8') as f:
            cache = json.load(f)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    total_chunks = 0
    translated_chunks = 0

    for page_data in pages:
        page_num = page_data['page']
        text = page_data['text']
        if not text.strip():
            continue

        # 过滤不需要的行
        clean_lines = []
        for line in text.split('\n'):
            if not should_skip(line) and line.strip():
                clean_lines.append(line.strip())
        clean_text = '\n'.join(clean_lines)
        if not clean_text.strip():
            continue

        if page_num > 1:
            doc.add_page_break()

        # 页标记
        page_header = doc.add_paragraph()
        run = page_header.add_run(f'第 {page_num} 页 / 共 {len(pages)} 页')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        page_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 拆分并翻译
        chunks = split_text(clean_text, max_chars=1200)
        total_chunks += len(chunks)

        for chunk in chunks:
            # 检查缓存
            cache_key = chunk[:100]  # 用前100字符作key
            if cache_key in cache:
                zh = cache[cache_key]
                translated_chunks += 1
            else:
                print(f'  翻译中... ({len(chunk)} 字符)')
                zh = translate_chunk(chunk)
                cache[cache_key] = zh
                translated_chunks += 1
                # 增量保存缓存
                if cache_path:
                    with open(cache_path, 'w', encoding='utf-8') as f:
                        json.dump(cache, f, ensure_ascii=False)
                time.sleep(0.5)  # 避免API限速

            # 写入DOCX
            # 原文
            para_en = doc.add_paragraph()
            run_en = para_en.add_run(chunk)
            run_en.font.size = Pt(9)
            run_en.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # 译文
            para_zh = doc.add_paragraph()
            run_zh = para_zh.add_run(f'  {zh}')
            run_zh.font.size = Pt(10)

            # 空行分隔
            doc.add_paragraph()

    doc.save(output_path)
    print(f'\n✅ 总计: {total_chunks} 段, {translated_chunks} 段已翻译')
    print(f'✅ 保存到: {output_path}')


if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('json', help='Extracted JSON file')
    ap.add_argument('-o', '--output', help='Output DOCX path', default=None)
    ap.add_argument('-c', '--cache', help='Translation cache JSON', default=None)
    args = ap.parse_args()

    json_path = args.json
    if not args.output:
        base = os.path.splitext(os.path.basename(json_path))[0]
        args.output = os.path.join(os.path.dirname(json_path), f'{base}_翻译.docx')
    if not args.cache:
        args.cache = os.path.join(os.path.dirname(json_path), 'translation_cache.json')

    print(f'📄 源文件: {json_path}')
    print(f'📝 输出: {args.output}')
    print(f'💾 缓存: {args.cache}')

    create_translated_docx(json_path, args.output, args.cache)
