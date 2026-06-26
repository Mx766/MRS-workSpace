"""
翻译 [1].pdf — 微聚焦超声皮肤紧致医学论文
"""
import sys, io, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 读取提取的文本
ROOT = os.path.dirname(os.path.dirname(__file__))
json_path = os.path.join(ROOT, '[1]_extracted.json')
if len(sys.argv) > 1:
    json_path = sys.argv[1]
with open(json_path, 'r', encoding='utf-8') as f:
    pages = json.load(f)

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 标题
title = doc.add_heading('微聚焦超声用于皮肤紧致（中英文对照）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Microfocused Ultrasound for Skin Tightening — Karol A. Gutowski, MD')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()

# 翻译映射（医学论文关键内容）
translations = {
    "Microfocused Ultrasound for Skin Tightening": "微聚焦超声用于皮肤紧致",
    "Karol A. Gutowski, MD": "Karol A. Gutowski，医学博士",
    "KEYWORDS": "关键词",
    "KEY POINTS": "核心要点",
    "Noninvasive skin tightening is possible in one treatment session using Ulthera microfocused ultrasound.":
        "使用 Ulthera 微聚焦超声可在一次治疗中实现无创皮肤紧致。",
    "Noninvasive skintightening": "无创皮肤紧致",
    "Ulthera": "Ulthera（超声刀）",
    "Microfocused ultrasound": "微聚焦超声",
    "Thermalinjuryzones": "热损伤区",
    "INTRODUCTION": "引言",
    "DISCUSSION": "讨论",
    "SUMMARY": "总结",
    "REFERENCES": "参考文献",
    "Anatomy": "解剖学",
    "Treatment Technique": "治疗技术",
    "Clinical Results": "临床结果",
    "Complications": "并发症",
    "Patient Selection": "患者选择",

    # 正文段落翻译
    "The Ulthera system delivers transcutaneous MFU energy from an operator-controlled hand piece and a monitor allows direct visualization of where the energy will be delivered so that specific structures such as bone can be avoided.":
        "Ulthera 系统通过操作者控制的手柄经皮输送微聚焦超声能量，监视器可直接显示能量输送位置，从而避开骨骼等特定结构。",

    "The focal energy delivery in given": "聚焦能量输送在特定",
    "Noninvasive skin tightening is possible in one treatment session using Ulthera microfocused ultrasound. Results may not be seen until 3 to 6 months after treatment. Some patients may not see improvement.":
        "使用 Ulthera 微聚焦超声可在一次治疗中实现无创皮肤紧致。治疗效果可能需要 3 至 6 个月才能显现。部分患者可能看不到改善。",

    "Face": "面部",
    "Neck": "颈部",
    "30 to 60 minutes": "30至60分钟",
    "30 to 45 minutes": "30至45分钟",
    "patient comfort": "患者舒适度",
    "treatment surface": "治疗面",
    "protocol": "方案",
    "operator-controlled hand": "操作者控制手柄",
    "piece": "部件",
    "monitor allows direct visualiza-": "监视器允许直接可视化",
    "tion (Fig. 2) of where the energy will be delivered": "（图2）能量将被输送到何处",
    "so that specific structures such as bone can be": "以便可以避开骨头等特定结构",
    "avoided.": "避开。",
    "The focal energy delivery in given": "聚焦能量在给定的",
}

# 逐页处理
for page_data in pages:
    page_num = page_data['page']
    text = page_data['text']
    if not text.strip():
        continue

    lines = text.split('\n')

    if page_num > 1:
        doc.add_page_break()

    page_header = doc.add_paragraph()
    run = page_header.add_run(f'第 {page_num} 页 / 共 {len(pages)} 页')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    page_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # 查找翻译
        zh = None
        if stripped in translations:
            zh = translations[stripped]
        else:
            # 尝试部分匹配
            for en, z in translations.items():
                if en in stripped or stripped in en:
                    zh = z
                    break

        para = doc.add_paragraph()

        # 检测标题（全大写短文本）
        if len(stripped) < 60 and (stripped.isupper() or stripped[0].isdigit()):
            run_en = para.add_run(stripped)
            run_en.bold = True
            run_en.font.size = Pt(12)
            run_en.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
            if zh:
                para2 = doc.add_paragraph()
                run_zh = para2.add_run(f'  → {zh}')
                run_zh.bold = True
                run_zh.font.size = Pt(11)
        elif zh:
            run_en = para.add_run(stripped)
            run_en.font.size = Pt(9)
            run_en.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            para2 = doc.add_paragraph()
            run_zh = para2.add_run(f'  {zh}')
            run_zh.font.size = Pt(10)
        else:
            run_en = para.add_run(stripped)
            run_en.font.size = Pt(9)
            run_en.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

output = os.path.join(ROOT, 'output', '超声皮肤紧致_中英文对照.docx')
if len(sys.argv) > 2:
    output = sys.argv[2]
doc.save(output)
print(f'✅ 保存到: {output}')
