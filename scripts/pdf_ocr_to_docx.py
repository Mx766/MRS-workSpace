"""
完整流程: PDF → Xfyun OCR → 结构化文本 → DOCX
用法: python scripts/pdf_ocr_to_docx.py <pdf_path> [page_range]
示例: python scripts/pdf_ocr_to_docx.py "original/xxx.pdf"
      python scripts/pdf_ocr_to_docx.py "original/xxx.pdf" 0-5  # 只处理前5页
"""
import base64, hashlib, hmac, json, os, sys, time
from datetime import datetime, timezone
from urllib.parse import urlencode
from wsgiref.handlers import format_date_time
import httpx
import pymupdf
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ═══════════════════════════════════════
# 配置
# ═══════════════════════════════════════
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)

# 加载讯飞凭证
with open(os.path.join(PROJECT_DIR, 'xfyun_config.json')) as f:
    _cfg = json.load(f)
APPID = _cfg['appid']
API_KEY = _cfg['apikey']
API_SECRET = _cfg['apisecret']

HOST = "cn-huabei-1.xf-yun.com"
PATH = "/v1/private/se75ocrbm"
BASE_URL = f"https://{HOST}{PATH}"

# ═══════════════════════════════════════
# 鉴权
# ═══════════════════════════════════════
def build_request():
    """构建带鉴权的 POST 请求"""
    now = datetime.now(timezone.utc)
    date_str = format_date_time(now.timestamp())

    sig = f"host: {HOST}\ndate: {date_str}\nPOST {PATH} HTTP/1.1"
    sig_sha = hmac.new(API_SECRET.encode(), sig.encode(), hashlib.sha256).digest()
    sig_b64 = base64.b64encode(sig_sha).decode()
    auth = f'api_key="{API_KEY}", algorithm="hmac-sha256", headers="host date request-line", signature="{sig_b64}"'
    auth_b64 = base64.b64encode(auth.encode()).decode()

    url = f"{BASE_URL}?{urlencode({'authorization': auth_b64, 'host': HOST, 'date': date_str})}"
    headers = {"Date": date_str, "Content-Type": "application/json"}
    return url, headers

# ═══════════════════════════════════════
# OCR 调用
# ═══════════════════════════════════════
def ocr_page(pdf_path, page_num=0, dpi=200, retries=3):
    """PDF单页→图片→讯飞OCR→结构化JSON"""
    # 1. 渲染页面为图片
    doc = pymupdf.open(pdf_path)
    page = doc[page_num]
    mat = pymupdf.Matrix(dpi/72, dpi/72)
    pix = page.get_pixmap(matrix=mat)
    img_bytes = pix.tobytes("jpg")
    doc.close()
    img_b64 = base64.b64encode(img_bytes).decode()
    print(f"  第{page_num+1}页: {len(img_bytes)/1024:.0f}KB 图片 (dpi={dpi})")

    # 2. 调用OCR API
    url, headers = build_request()
    body = {
        "header": {"app_id": APPID, "status": 2},  # 2=一次性请求
        "parameter": {
            "ocr": {
                "result_option": "normal",
                "result_format": "json"
            }
        },
        "payload": {
            "image": {
                "encoding": "jpg",
                "image": img_b64,
                "status": 2
            }
        }
    }

    for attempt in range(retries):
        try:
            resp = httpx.post(url, headers=headers, json=body, timeout=120)
            if resp.status_code == 200:
                return resp.json()
            else:
                print(f"    HTTP {resp.status_code}: {resp.text[:200]}")
                if resp.status_code == 401:
                    return None  # 鉴权失败不重试
        except Exception as e:
            print(f"    重试 {attempt+1}/{retries}: {e}")
            time.sleep(2)
    return None

def extract_text_from_ocr(result):
    """从OCR结果中提取结构化文本"""
    if not result:
        return {"text": "", "blocks": []}

    payload = result.get('payload', {})
    header = result.get('header', {})

    # 检查错误
    if header.get('code', 0) != 0:
        print(f"    OCR错误: code={header.get('code')}, msg={header.get('message', '')}")
        return {"text": "", "blocks": []}

    # 尝试多种可能的输出格式
    output = None

    # 格式1: payload.ocr_output (JSON)
    if 'ocr_output' in payload:
        output = payload['ocr_output']

    # 格式2: payload.result.text
    elif 'result' in payload:
        result_data = payload['result']
        if isinstance(result_data, dict):
            output = result_data
        elif isinstance(result_data, str):
            return {"text": result_data, "blocks": []}

    # 格式3: payload.text (纯文本)
    elif 'text' in payload:
        return {"text": payload['text'], "blocks": []}

    # 格式4: 直接字符串
    elif isinstance(payload, str):
        return {"text": payload, "blocks": []}

    if output is None:
        print(f"    未知输出格式, payload keys: {list(payload.keys())}")
        return {"text": str(payload), "blocks": []}

    # 解析结构化输出
    if isinstance(output, str):
        return {"text": output, "blocks": []}
    elif isinstance(output, dict):
        # 常见字段: content, text, blocks, paragraphs
        text = output.get('content', '') or output.get('text', '') or json.dumps(output, ensure_ascii=False)
        blocks = output.get('blocks', []) or output.get('paragraphs', []) or output.get('result', [])
        return {"text": text, "blocks": blocks}

    return {"text": str(output), "blocks": []}

# ═══════════════════════════════════════
# DOCX 生成
# ═══════════════════════════════════════
def build_docx_from_pages(pages_ocr_data, output_path):
    """根据OCR结果生成DOCX"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    for page_data in pages_ocr_data:
        text = page_data.get('text', '')
        blocks = page_data.get('blocks', [])

        if blocks:
            # 有结构化block数据
            for block in blocks:
                if isinstance(block, dict):
                    block_type = block.get('type', 'text')
                    content = block.get('content', '') or block.get('text', '')
                    bbox = block.get('bbox', block.get('position', None))

                    if block_type in ('title', 'heading', 'header'):
                        para = doc.add_paragraph()
                        run = para.add_run(content)
                        run.bold = True
                        run.font.size = Pt(14)
                    elif block_type == 'table':
                        # 尝试重建表格
                        rows = block.get('rows', [])
                        if rows:
                            table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 1)
                            for ri, row_cells in enumerate(rows):
                                for ci, cell_text in enumerate(row_cells):
                                    if ci < len(table.rows[ri].cells):
                                        table.rows[ri].cells[ci].text = str(cell_text)
                    else:
                        para = doc.add_paragraph()
                        run = para.add_run(content)
                        run.font.size = Pt(10.5)
                else:
                    para = doc.add_paragraph()
                    para.add_run(str(block)).font.size = Pt(10.5)
        else:
            # 纯文本，按行拆分
            for line in text.split('\n'):
                line = line.strip()
                if line:
                    para = doc.add_paragraph()
                    para.add_run(line).font.size = Pt(10.5)

        # 分页
        if page_data != pages_ocr_data[-1]:
            doc.add_page_break()

    doc.save(output_path)
    print(f"保存: {output_path} ({os.path.getsize(output_path)/1024:.0f} KB)")

# ═══════════════════════════════════════
# 主流程
# ═══════════════════════════════════════
def main(pdf_path, start_page=0, end_page=None, dpi=200):
    print(f"PDF: {pdf_path}")
    print(f"大小: {os.path.getsize(pdf_path)/1024:.0f} KB")
    print(f"DPI: {dpi}")

    # 获取总页数
    doc = pymupdf.open(pdf_path)
    total_pages = len(doc)
    doc.close()

    if end_page is None or end_page > total_pages:
        end_page = min(total_pages, 5)  # 默认测试前5页

    print(f"处理范围: 第{start_page+1}-{end_page}页 (共{total_pages}页)")
    print()

    # 逐页OCR
    pages_data = []
    for pg in range(start_page, end_page):
        result = ocr_page(pdf_path, pg, dpi=dpi)
        extracted = extract_text_from_ocr(result)
        extracted['page'] = pg + 1
        pages_data.append(extracted)

        # 打印文本预览
        text_preview = extracted['text'][:200].replace('\n', ' | ')
        print(f"    → {len(extracted['text'])} 字符: {text_preview}...")
        print(f"    → {len(extracted['blocks'])} 个结构块")

        # 保存中间结果
        json_out = pdf_path.replace('.pdf', f'_ocr_pg{pg+1}.json')
        with open(json_out, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

    # 生成DOCX
    docx_out = pdf_path.replace('.pdf', '_ocr.docx')
    build_docx_from_pages(pages_data, docx_out)

    # 保存原始OCR（用于调试）
    raw_out = pdf_path.replace('.pdf', '_ocr_raw.json')
    with open(raw_out, 'w', encoding='utf-8') as f:
        json.dump(pages_data, f, ensure_ascii=False, indent=2)

    print(f"\n完成! 输出文件:")
    print(f"  DOCX: {docx_out}")
    print(f"  原始数据: {raw_out}")
    return docx_out

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        # 默认测试
        default_pdf = os.path.join(PROJECT_DIR, 'original',
            'CH2.4.1(4-2) 1. MSDS_Barrel_TOPAS MSDS-EN V8.00_0 (2).pdf')
        if os.path.exists(default_pdf):
            print(f"使用默认测试文件: {default_pdf}")
            main(default_pdf, dpi=150)
    else:
        pdf_path = sys.argv[1]
        start, end = 0, None
        if len(sys.argv) > 2:
            parts = sys.argv[2].split('-')
            start = int(parts[0])
            if len(parts) > 1:
                end = int(parts[1])
        main(pdf_path, start_page=start, end_page=end, dpi=200)
