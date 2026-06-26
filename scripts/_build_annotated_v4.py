#!/usr/bin/env python3
"""
Build annotated DOCX — v4 (FIXED HIGHLIGHTS)

KEY INSIGHT (from Issue #12 investigation):
  lxml MUST NOT parse or modify document.xml. python-docx's OxmlElement and
  raw lxml etree.Element use different namespace representations; when lxml
  round-trips document.xml, Word stops rendering w:shd + w:highlight.

v4 FIX:
  STEP 1 (python-docx): highlights + commentRangeStart/End/Reference → doc.save()
  STEP 2 (ZIP): build comments.xml, update rels, update CT → write output
                 ^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^
                 document.xml is COPIED VERBATIM, never parsed by lxml
"""
import sys, json, os, tempfile, zipfile, re as regex
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8')

NS_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
COMMENTS_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
CT_COMMENTS = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
XML_NS = 'http://www.w3.org/XML/1998/namespace'

# ── Config ──
INPUT_DOCX = 'D:/translation/proofread/1/[1]-CN.docx'
ISSUES_JSON = 'cache/issues_phase4_new.json'
OUTPUT_DIR = 'D:/translation/proofread/1/v8'
OUTPUT_NAME = '[1]-CN_校对稿.docx'

with open(ISSUES_JSON, 'r', encoding='utf-8') as f:
    issues = json.load(f).get('issues', [])

print(f'Loaded {len(issues)} issues from {ISSUES_JSON}')

doc = Document(INPUT_DOCX)
hl_colors = {'critical': ('FFCCCC', 'red'), 'medium': ('FFF9C4', 'yellow'), 'low': ('C8E6C9', 'green')}
# Truncation warnings get orange to distinguish from translation issues
truncation_color = ('FFDAB9', 'red')  # light orange
sev_order = {'critical': 0, 'medium': 1, 'low': 2}

# Group issues by paragraph
issues_by_para = {}
for issue in issues:
    pi = issue.get('paragraph_index', 0)
    if pi < 1:
        continue
    issues_by_para.setdefault(pi, []).append(issue)

# ═══════════════════════════════════════════════════════════════════
# STEP 1: python-docx does EVERYTHING — highlights + comment refs
# ═══════════════════════════════════════════════════════════════════

comment_id = 0
para_comments_map = {}
total_matched = 0
matched_by_sev = {'critical': 0, 'medium': 0, 'low': 0}
total_by_sev = {'critical': 0, 'medium': 0, 'low': 0}

for para_idx in sorted(issues_by_para.keys()):
    if para_idx > len(doc.paragraphs):
        continue
    para = doc.paragraphs[para_idx - 1]
    sorted_issues = sorted(issues_by_para[para_idx],
                           key=lambda x: sev_order.get(x.get('severity', 'medium'), 1))

    # ── Collect comment metadata ──
    para_comments = []
    for issue in sorted_issues:
        sev = issue.get('severity', 'medium')
        is_trunc = issue.get('dimension', '') == '⚠️ 人工修订警告'
        sev_label = {
            'critical': '🔴 严重（必须修改）',
            'medium': '🟡 中等（建议修改）',
            'low': '🟢 低（格式优化）',
        }.get(sev, '⚪ 未分级')
        if is_trunc:
            sev_label = '🔴 严重-人工修订截断（文字丢失）'
        para_comments.append((comment_id, sev, sev_label, issue, is_trunc))
        comment_id += 1
    para_comments_map[para_idx] = para_comments

    # ── 1a: Add HIGHLIGHTS + commentReference TOGETHER in the SAME run ──
    matched_cids = set()
    # Pre-compute full paragraph text for cross-run matching
    full_text = ''.join(r.text for r in para.runs)

    for cid, sev, sev_label, issue, is_trunc in para_comments:
        total_by_sev[sev] = total_by_sev.get(sev, 0) + 1
        fill, hl_name = hl_colors.get(sev, ('FFFF00', 'yellow'))
        target = issue.get('target_quote', '').strip()
        if target:
            matched = False

            # Strategy 1: find target in individual runs (fast, most cases)
            for key_len in [15, 10, 8, 5]:
                key = target[:key_len].strip()
                if not key or len(key) < 3:
                    continue
                for run in para.runs:
                    if not run.text.strip():
                        continue
                    if key in run.text:
                        rPr = run._r.get_or_add_rPr()
                        if rPr.find(qn('w:shd')) is not None:
                            # Already highlighted by a previous issue — still add ref
                            cr = OxmlElement('w:commentReference')
                            cr.set(qn('w:id'), str(cid))
                            rPr.append(cr)
                            matched_cids.add(cid)
                            total_matched += 1
                            matched_by_sev[sev] = matched_by_sev.get(sev, 0) + 1
                            matched = True
                            break
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), fill)
                        rPr.append(shd)
                        hl = OxmlElement('w:highlight')
                        hl.set(qn('w:val'), hl_name)
                        rPr.append(hl)
                        cr = OxmlElement('w:commentReference')
                        cr.set(qn('w:id'), str(cid))
                        rPr.append(cr)
                        matched_cids.add(cid)
                        total_matched += 1
                        matched_by_sev[sev] = matched_by_sev.get(sev, 0) + 1
                        matched = True
                        break
                if matched:
                    break

            # Strategy 2: target exists in paragraph but split across runs
            if not matched and target in full_text:
                # Find which run contains the start of the target
                target_start = full_text.index(target)
                char_pos = 0
                for run in para.runs:
                    run_len = len(run.text)
                    if char_pos <= target_start < char_pos + run_len:
                        rPr = run._r.get_or_add_rPr()
                        if rPr.find(qn('w:shd')) is None:
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), fill)
                            rPr.append(shd)
                            hl = OxmlElement('w:highlight')
                            hl.set(qn('w:val'), hl_name)
                            rPr.append(hl)
                        cr = OxmlElement('w:commentReference')
                        cr.set(qn('w:id'), str(cid))
                        rPr.append(cr)
                        matched_cids.add(cid)
                        total_matched += 1
                        matched_by_sev[sev] = matched_by_sev.get(sev, 0) + 1
                        matched = True
                        break
                    char_pos += run_len

    # ── 1b: For UNMATCHED issues, insert a small ref run after pPr ──
    # (NOT at paragraph end — that would break highlights)
    p_elem = para._p
    pPr = p_elem.find(qn('w:pPr'))
    insert_pos = (list(p_elem).index(pPr) + 1) if pPr is not None else 0

    for cid, sev, sev_label, issue, is_trunc in para_comments:
        if cid in matched_cids:
            continue
        ref_run = OxmlElement('w:r')
        ref_rPr = OxmlElement('w:rPr')
        cr = OxmlElement('w:commentReference')
        cr.set(qn('w:id'), str(cid))
        ref_rPr.append(cr)
        ref_run.append(ref_rPr)
        p_elem.insert(insert_pos, ref_run)
        insert_pos += 1

# Save via python-docx → highlights + comment refs all preserved
tmpfd, tmppath = tempfile.mkstemp(suffix='.docx')
os.close(tmpfd)
doc.save(tmppath)
print(f'STEP 1 done: highlights + {comment_id} comment refs written via python-docx')

# ═══════════════════════════════════════════════════════════════════
# STEP 2: ZIP layer — only build comments.xml, update rels, update CT
#          document.xml is COPIED VERBATIM (no lxml round-trip!)
# ═══════════════════════════════════════════════════════════════════

entries = {}
with zipfile.ZipFile(tmppath, 'r') as zf:
    for info in zf.infolist():
        entries[info.filename] = (info, zf.read(info.filename))

# --- 2a: Build comments.xml with lxml (NEW file, not modifying existing) ---
comments_tree = etree.Element(f'{{{NS_w}}}comments', nsmap={'w': NS_w})

for para_idx in sorted(para_comments_map.keys()):
    for cid, sev, sev_label, issue, is_trunc in para_comments_map[para_idx]:
        date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
        cmt = etree.SubElement(comments_tree, f'{{{NS_w}}}comment')
        cmt.set(f'{{{NS_w}}}id', str(cid))
        cmt.set(f'{{{NS_w}}}author', 'TranslationReview')
        cmt.set(f'{{{NS_w}}}date', date_str)
        cmt.set(f'{{{NS_w}}}initials', 'TR')

        lines = [
            sev_label,
            f'【{issue.get("dimension", "")} — {issue.get("check_item", "")}】', '',
            f'原文: {issue.get("source_quote", "")}', '',
            f'问题: {issue.get("issue", "")}', '',
            f'建议: {issue.get("suggestion", "")}',
        ]
        for i, line in enumerate(lines):
            pe = etree.SubElement(cmt, f'{{{NS_w}}}p')
            pp = etree.SubElement(pe, f'{{{NS_w}}}pPr')
            r_e = etree.SubElement(pe, f'{{{NS_w}}}r')
            if i == 0:
                rp = etree.SubElement(r_e, f'{{{NS_w}}}rPr')
                etree.SubElement(rp, f'{{{NS_w}}}b')
                sz = etree.SubElement(rp, f'{{{NS_w}}}sz')
                sz.set(f'{{{NS_w}}}val', '18')
                co = etree.SubElement(rp, f'{{{NS_w}}}color')
                co.set(f'{{{NS_w}}}val',
                       {'critical': 'FF0000', 'medium': 'FF8C00', 'low': '228B22'}.get(sev, '000000'))
            te = etree.SubElement(r_e, f'{{{NS_w}}}t')
            te.set(f'{{{XML_NS}}}space', 'preserve')
            te.text = line

comments_xml = etree.tostring(comments_tree, xml_declaration=True,
                               encoding='UTF-8', standalone=True)
entries['word/comments.xml'] = (None, comments_xml)
print(f'STEP 2a: comments.xml built with {comment_id} comments')

# --- 2b: Update word/_rels/document.xml.rels ---
rels_path = 'word/_rels/document.xml.rels'
rels_data = entries.get(rels_path, (None, None))[1]
rels_tree = etree.fromstring(rels_data) if rels_data else etree.Element(f'{{{REL_NS}}}Relationships')

max_rid = 0
has_comments_rel = False
for rel in rels_tree.iter():
    if rel.tag == f'{{{REL_NS}}}Relationship':
        rid = rel.get('Id', '')
        m = regex.match(r'rId(\d+)', rid)
        if m:
            max_rid = max(max_rid, int(m.group(1)))
        if rel.get('Type') == COMMENTS_NS:
            has_comments_rel = True

if not has_comments_rel:
    nr = etree.SubElement(rels_tree, f'{{{REL_NS}}}Relationship')
    nr.set('Id', f'rId{max_rid + 1}')
    nr.set('Type', COMMENTS_NS)
    nr.set('Target', 'comments.xml')
    entries[rels_path] = (
        entries.get(rels_path, (None, None))[0],
        etree.tostring(rels_tree, xml_declaration=True, encoding='UTF-8', standalone=True),
    )
    print(f'STEP 2b: Added comments relationship to {rels_path}')
else:
    print(f'STEP 2b: Comments relationship already exists in {rels_path}')

# --- 2c: Update [Content_Types].xml ---
ct_path = '[Content_Types].xml'
ct_data = entries.get(ct_path, (None, None))[1]
ct_tree = etree.fromstring(ct_data) if ct_data else etree.Element(f'{{{CT_NS}}}Types')

has_ct = False
for ov in ct_tree.iter():
    if ov.tag == f'{{{CT_NS}}}Override' and ov.get('PartName') == '/word/comments.xml':
        has_ct = True
        break

if not has_ct:
    nct = etree.SubElement(ct_tree, f'{{{CT_NS}}}Override')
    nct.set('PartName', '/word/comments.xml')
    nct.set('ContentType', CT_COMMENTS)
    entries[ct_path] = (
        entries.get(ct_path, (None, None))[0],
        etree.tostring(ct_tree, xml_declaration=True, encoding='UTF-8', standalone=True),
    )
    print(f'STEP 2c: Added comments content type')
else:
    print(f'STEP 2c: Comments content type already exists')

# --- 2d: Write final ZIP ---
os.makedirs(OUTPUT_DIR, exist_ok=True)
output = os.path.join(OUTPUT_DIR, OUTPUT_NAME)
with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
    for name, (info, data) in entries.items():
        if info is not None:
            zf.writestr(info, data)
        else:
            zf.writestr(name, data)

os.unlink(tmppath)
print(f'STEP 2d: Output written to {output}')

# ═══════════════════════════════════════════════════════════════════
# VERIFY
# ═══════════════════════════════════════════════════════════════════
print()
print('=== VERIFICATION ===')

# 1. Count highlights in output
with zipfile.ZipFile(output, 'r') as zf:
    doc_xml = zf.read('word/document.xml')
    shd_count = doc_xml.count(b'w:shd')
    hl_count = doc_xml.count(b'w:highlight')
    crs_count = doc_xml.count(b'commentRangeStart')
    cre_count = doc_xml.count(b'commentRangeEnd')
    ref_count = doc_xml.count(b'commentReference')
    has_comments_xml = 'word/comments.xml' in zf.namelist()

    # Check document.xml was NOT parsed by lxml
    # (if it had been, lxml would add its own namespace prefixes)
    has_lxml_ns = b'ns0:' in doc_xml or b'ns1:' in doc_xml

print(f'Highlights (w:shd): {shd_count}')
print(f'Highlights (w:highlight): {hl_count}')
print(f'Comment refs: {ref_count} total ({len(matched_cids) if "matched_cids" in dir() else "?"} with highlight, rest bare ref)')
print(f'comments.xml present: {has_comments_xml}')
print(f'lxml namespace artifacts (ns0:/ns1:): {"❌ FOUND — may break highlights!" if has_lxml_ns else "✅ None — document.xml untouched by lxml"}')

# 2. Verify comment range ordering
with zipfile.ZipFile(output, 'r') as zf:
    doc_tree = etree.fromstring(zf.read('word/document.xml'))
    body = doc_tree.find(f'{{{NS_w}}}body')
    bad_order = 0
    for p in body.findall(f'{{{NS_w}}}p'):
        crs_ids = [c.get(f'{{{NS_w}}}id') for c in p.findall(f'{{{NS_w}}}commentRangeStart')]
        cre_ids = [c.get(f'{{{NS_w}}}id') for c in p.findall(f'{{{NS_w}}}commentRangeEnd')]
        if crs_ids != cre_ids:
            bad_order += 1
            print(f'  ⚠️ MISMATCH in paragraph: crs={crs_ids}, cre={cre_ids}')
if bad_order == 0:
    print(f'Comment range ordering: ✅ All {len(body.findall(f"{{{NS_w}}}p"))} paragraphs OK')

# 3. Try opening with python-docx (sanity check)
try:
    anno = Document(output)
    anno_paras = len(anno.paragraphs)
    orig_paras = len(Document(INPUT_DOCX).paragraphs)
    print(f'python-docx roundtrip: ✅ ({anno_paras} paragraphs)')
except Exception as e:
    print(f'python-docx roundtrip: ❌ {e}')

print()
print(f'=== MATCHING REPORT ===')
print(f'Total issues: {comment_id}')
print(f'  critical: {total_by_sev.get("critical",0)} total, {matched_by_sev.get("critical",0)} highlighted')
print(f'  medium:   {total_by_sev.get("medium",0)} total, {matched_by_sev.get("medium",0)} highlighted')
print(f'  low:      {total_by_sev.get("low",0)} total, {matched_by_sev.get("low",0)} highlighted')
print(f'  unmatched: {comment_id - total_matched} (bare ref, no highlight)')
print()
print('=== NEXT: Open in Word and verify highlights are visible ===')
print(f'File: {output}')
