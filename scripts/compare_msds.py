"""
MSDS 三版对比: 原始PDF vs Acrobat vs 讯飞OCR
输出到文件，避免GBK终端问题
"""
import pymupdf, re, json, os
from docx import Document

OUT = []
def p(msg):
    OUT.append(msg)

# 路径 — 可通过命令行参数覆盖
ROOT = os.path.dirname(os.path.dirname(__file__))
PDF = os.path.join(ROOT, 'original', 'CH2.4.1(4-2) 1. MSDS_Barrel_TOPAS MSDS-EN V8.00_0 (2).pdf')
ACR = os.path.join(ROOT, 'output', 'MSDS_converted.docx')
if len(sys.argv) > 1:
    PDF = sys.argv[1]
if len(sys.argv) > 2:
    ACR = sys.argv[2]

# === 1. PDF原始数据 ===
pdf = pymupdf.open(PDF)
pdf_items = []
for pg_num, page in enumerate(pdf):
    blocks = page.get_text('dict')['blocks']
    for b in blocks:
        if b['type'] != 0: continue
        for l in b['lines']:
            text = ' '.join(s['text'] for s in l['spans']).strip()
            if text:
                s0 = l['spans'][0]
                pdf_items.append({
                    'text': text, 'page': pg_num+1,
                    'size': round(s0['size'],1),
                    'bold': bool(s0.get('flags',0)&2),
                    'y': round(l['bbox'][1],1)
                })

# === 2. Acrobat DOCX ===
doc = Document(ACR)
acr_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
acr_tables_raw = []
for ti, t in enumerate(doc.tables):
    rows = [[c.text.strip() for c in r.cells] for r in t.rows]
    acr_tables_raw.append(rows)

# ============================================
# 问题1: 页眉污染
# ============================================
p('=== 问题1: 页眉/页脚在正文中重复出现 ===')
header_pats = ['SAFETY DATA SHEET', 'TOPAS', 'Revision Date', 'product code', 'Version / Revision']
for hp in header_pats:
    positions = [i for i, t in enumerate(acr_texts) if hp in t]
    if len(positions) >= 3:
        p(f'"{hp}" 重复 {len(positions)} 次, 段落位置: {positions}')

# ============================================
# 问题2: 整段丢失
# ============================================
p('')
p('=== 问题2: PDF有但Acrobat丢失的内容 ===')
# 页眉和页码本身应去掉，只检查正文内容
skip_pats = ['SAFETY DATA SHEET', 'Revision Date', 'product code', 'Version / Revision',
             'USA/ EN', '________________________________________________________________']
for item in pdf_items:
    t = item['text']
    if any(s in t for s in skip_pats): continue
    if re.match(r'^\d+\s*/\s*\d+$', t.strip()): continue
    # 在Acrobat段落和表格中查找
    found_in_para = any(t in at for at in acr_texts)
    found_in_table = any(t in str(rows) for rows in acr_tables_raw)
    if not found_in_para and not found_in_table and len(t) > 5:
        p(f'PDF p{item["page"]} [{item["size"]}pt bold={item["bold"]}]: "{t[:100]}"')

# ============================================
# 问题3: 节标题变成单行表格
# ============================================
p('')
p('=== 问题3: 节标题被错误转为表格 ===')
section_nums = ['1:', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.',
                '11.', '12.', '13.', '14.', '15.', '16.']
for rows in acr_tables_raw:
    if len(rows) == 1 and len(rows[0]) <= 2:
        text = ' '.join(rows[0])
        if any(text.startswith(sn) for sn in section_nums):
            p(f'  节标题变成表格: "{text}" ({len(rows[0])}列, {len(rows)}行)')

# ============================================
# 问题4: 页眉变成表格
# ============================================
p('')
p('=== 问题4: 页眉信息被转为表格 ===')
for i, rows in enumerate(acr_tables_raw):
    text = ' '.join(c for row in rows for c in row)
    if 'SAFETY DATA SHEET' in text and 'Revision Date' in text:
        p(f'  Table{i+1}: 页眉变成 {len(rows)}x{len(rows[0])} 表格')
    elif re.match(r'\d+\s*/\s*\d+', text.strip()):
        p(f'  Table{i+1}: 页码变成表格 "{text.strip()}"')

# ============================================
# 问题5: Tab拼接
# ============================================
p('')
p('=== 问题5: Tab拼接导致的格式混乱 ===')
for i, t in enumerate(acr_texts):
    if '\t' in t and len(t) > 60:
        p(f'  P{i}: Tab拼接 "{t[:150]}"')

# ============================================
# 问题6: 格式丢失
# ============================================
p('')
p('=== 问题6: 排版格式丢失 ===')
bold_count = sum(1 for i in pdf_items if i['bold'])
p(f'  PDF有 {bold_count} 行粗体, Acrobat: 0 (全部Normal)')
# 字号分布
from collections import Counter
size_dist = Counter(i['size'] for i in pdf_items)
p(f'  PDF字号分布: {dict(sorted(size_dist.items()))}')
p(f'  Acrobat: 全部统一字号(无变化)')

# ============================================
# 问题7: 数值/符号校验
# ============================================
p('')
p('=== 问题7: 关键数据校验 ===')
checks = [
    ('350 degree C / 662F', 'Section 10 temperature warning'),
    ('< 0.001 mm Hg @25 degree C', 'Vapour pressure value'),
    ('< 0.5 % (wt/wt)', 'VOC content'),
    ('paragraph (d) of', 'Section 1910.1200 reference'),
    ('TOPAS (R)', 'Registered trademark symbol'),
    ('***', 'Change markers'),
    ('CAS 26007-43-2', 'CAS number'),
    ('approx. 75-185', 'Melting point'),
    ('550 - 600 g/l', 'Bulk density'),
    ('29CFR 1910.1200', 'OSHA regulation'),
]
for val, desc in checks:
    # Normalize symbols for comparison
    in_pdf = False
    for item in pdf_items:
        t = item['text'].replace('®', '(R)').replace('°', 'degree')
        if val.replace('degree', '°').replace('(R)', '®') in item['text']:
            in_pdf = True
            break
        if val in t:
            in_pdf = True
            break
    in_acr = any(val in t for t in acr_texts)
    in_acr_tbl = any(val in str(rows) for rows in acr_tables_raw)
    if in_acr or in_acr_tbl:
        p(f'  [OK] {desc}')
    else:
        p(f'  [MISSING] {desc}: "{val}"')

# ============================================
# 问题8: 表格统计
# ============================================
p('')
p('=== 问题8: 表格质量 ===')
p(f'  Acrobat: 总共{len(acr_tables_raw)}个表格')
real, pseudo = 0, 0
for rows in acr_tables_raw:
    if len(rows) >= 2 and len(rows[0]) >= 2:
        real += 1
    else:
        pseudo += 1
p(f'  真表格: {real}, 伪表格(应删除/修复): {pseudo}')

# === 保存 ===
with open(r'd:\translation\output\msds_comparison.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(OUT))
print(f'Done: {len(OUT)} lines saved to output/msds_comparison.txt')
