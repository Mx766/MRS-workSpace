"""
后处理有道 API 转换的 DOCX：
1. 修复表格边框 — 所有单元格补全边框
2. 清理多余空白段落 — 连续空段合并
"""
import sys, io, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from pathlib import Path


def fix_table_borders(doc):
    """给所有表格的所有单元格添加完整边框"""
    border_names = ['top', 'bottom', 'left', 'right']

    for table in doc.tables:
        # 确保 table 级别有边框
        tbl_pr = table._tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tbl_pr)

        tbl_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_borders is None:
            tbl_borders = OxmlElement('w:tblBorders')
            tbl_pr.append(tbl_borders)

        # Table-level defaults
        for bn in border_names:
            if tbl_borders.find(qn(f'w:{bn}')) is None:
                b = OxmlElement(f'w:{bn}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), '000000')
                tbl_borders.append(b)

        # Cell-level: 每个单元格都加边框
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tc_pr = tc.find(qn('w:tcPr'))
                if tc_pr is None:
                    tc_pr = OxmlElement('w:tcPr')
                    tc.insert(0, tc_pr)

                tc_borders = tc_pr.find(qn('w:tcBorders'))
                if tc_borders is None:
                    tc_borders = OxmlElement('w:tcBorders')
                    tc_pr.append(tc_borders)

                for bn in border_names:
                    if tc_borders.find(qn(f'w:{bn}')) is None:
                        b = OxmlElement(f'w:{bn}')
                        b.set(qn('w:val'), 'single')
                        b.set(qn('w:sz'), '4')
                        b.set(qn('w:space'), '0')
                        b.set(qn('w:color'), '000000')
                        tc_borders.append(b)


def clean_blank_paragraphs(doc, max_consecutive=1):
    """合并连续空段落：最多保留 max_consecutive 个"""
    body = doc.element.body
    paras = body.findall(qn('w:p'))

    to_remove = []
    consecutive_empty = 0

    for i, p in enumerate(paras):
        # 检查是否为空段落
        runs = p.findall(qn('w:r'))
        has_text = any(r.findall(qn('w:t')) for r in runs)
        has_image = any(r.findall('.//' + qn('wp:inline')) for r in runs) or \
                    any(r.findall('.//' + qn('wp:anchor')) for r in runs)
        has_field = any(r.findall(qn('w:fldChar')) for r in runs)

        is_empty = not has_text and not has_image and not has_field

        if is_empty:
            consecutive_empty += 1
            if consecutive_empty > max_consecutive:
                to_remove.append(p)
        else:
            consecutive_empty = 0

    for p in to_remove:
        body.remove(p)

    return len(to_remove)


def process_docx(input_path, output_path=None):
    """处理后处理 DOCX"""
    if output_path is None:
        p = Path(input_path)
        output_path = p.parent / f'{p.stem}_fixed{p.suffix}'

    print(f'Processing: {input_path}')
    doc = Document(input_path)

    # 1. 修复表格边框
    print(f'  Tables: {len(doc.tables)}')
    fix_table_borders(doc)
    print(f'  ✓ Table borders fixed')

    # 2. 清理空白段落
    removed = clean_blank_paragraphs(doc, max_consecutive=1)
    print(f'  ✓ Removed {removed} excessive blank paragraphs')

    # 3. 设置表格默认边距（让表格更紧凑）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.space_before = Pt(1)
                    pf.space_after = Pt(1)

    doc.save(str(output_path))
    print(f'  ✓ Saved: {output_path}')
    return str(output_path)


if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser(description='修复有道 API 转换的 DOCX')
    ap.add_argument('input', help='输入 DOCX')
    ap.add_argument('-o', '--output', default=None)
    args = ap.parse_args()
    process_docx(args.input, args.output)
