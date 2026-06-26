#!/usr/bin/env python3
"""
Phase: Merge v2.2 new findings + human truncation detection into issues JSON.
Then build annotated DOCX with v5 approach.
"""
import json, sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

# ── Load base issues (v2.1, 51 issues) ──
with open('cache/issues_phase4_v2_1.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

existing = data['issues']
existing_paras = set(i['paragraph_index'] for i in existing)
print(f'Loaded {len(existing)} existing issues')

# ── v2.2 NEW findings ──
new_issues = []

def add_issue(pi, dim, check_item, sev, conf, src, tgt, issue_text, suggestion):
    new_issues.append({
        "chapter": "Ch1",
        "paragraph_index": pi,
        "dimension": dim,
        "check_item": check_item,
        "severity": sev,
        "confidence": conf,
        "source_quote": src,
        "target_quote": tgt,
        "issue": issue_text,
        "suggestion": suggestion
    })

# --- 2.23 MT术语泛化 ---
add_issue(60, "二.表达规范", "2.23 MT术语泛化", "medium", "high",
    "using 1.5 mm and 3.0 mm probes",
    "使用1.5 mm和3.0 mm探头靶向下眼睑皮肤和眶隔后",
    "「探头」为MT直译短词。医疗器械文档中「probe」标准译法为「治疗探头」（加领域修饰语）。参见训练数据中386处相同修订。",
    "改为：「…使用1.5 mm和3.0 mm治疗探头…」")

add_issue(23, "二.表达规范", "2.23 MT术语泛化", "medium", "medium",
    "",
    "MFU与RF能量的不同之处在于，它可以微聚焦以靶向更深层组织",
    "「靶向」为MT直译（target→靶向）。中文医学文献更自然表述为「作用于」「针对」。注意：此条需结合上下文判断，部分语境下「靶向」可接受。",
    "建议：「…其能量可微聚焦至更深层组织，作用于…」如需保留「靶向」，建议加注说明。")

# --- 4.3 跨段术语: P18 标题 ---
add_issue(18, "四.术语合规", "4.3 跨段术语统一", "critical", "high",
    "Collagen Denaturation and Neocollagenesis",
    "胶原蛋白变性与新胶原生成",
    "标题中使用「新胶原生成」而正文P22等处使用「新胶原蛋白生成」。同一概念「collagen」在全文中混用「胶原」和「胶原蛋白」。此为v2.1遗漏。",
    "统一为：「胶原蛋白变性与新胶原蛋白合成」")

# --- Additional 2.25 字面直译 ---
add_issue(40, "二.表达规范", "2.25 字面直译短语", "medium", "medium",
    "Skin histology revealed consistent and reproducible thermal injury zones in the dermis, the epidermis was unaffected",
    "皮肤组织学检查显示，真皮出现一致且可重复的热损伤区域，表皮未受损伤",
    "「表皮未受损伤」为英文 'the epidermis was unaffected' 的字面直译。中文更自然表述为「且表皮未受损伤」或重组语序。该句语序混乱——先描述了真皮热损伤、表皮未受损、损伤12周消失，三件事的时间/因果逻辑不清晰。",
    "建议重组语序：「…真皮出现一致且可重复的热损伤区，该损伤在12周时已消失，且表皮未受损伤。」")

# --- v2.2 4.1 Domain check: this is a medical paper, check regulatory terms ---
add_issue(110, "二.表达规范", "2.23 MT术语泛化", "medium", "medium",
    "ideal candidates typically present with mild to moderate aging",
    "由于理想候选人通常呈现轻度至中度衰老",
    "「候选人」为 candidate 的字面直译。医学美容文献中应译为「适应证患者」或「适合治疗的患者」。",
    "改为：「由于适合治疗的患者通常表现为轻度至中度皮肤老化」")

add_issue(111, "二.表达规范", "2.23 MT术语泛化", "medium", "medium",
    "higher response might have been obtained",
    "可能会获得更高的响应",
    "「响应」为 response 的MT直译。临床试验术语中应为「反应率」或「有效率」。",
    "改为：「可能会获得更高的反应率」")

add_issue(120, "一.双语忠实度", "1.4 错译", "critical", "high",
    "higher treatment densities (more treatment lines)",
    "在更高的治疗密度（更多疗程）",
    "「更多疗程」错译——原文 'more treatment lines' 指更多治疗线数，非治疗疗程（sessions）。「疗程」是完全不同的临床概念（一个疗程包含多次治疗）。",
    "改为：「若采用更高的治疗密度（更多治疗线）并在6至9个月时进行评估，可能获得更有利的治疗反应率。」")

print(f'Added {len(new_issues)} new v2.2 issues')

# ── Human truncation detection ──
orig = Document('proofread/1/[1]-CN.docx')
human = Document('proofread/1/[1]-CN-Update-XD.docx')

truncation_warnings = []
for i, (op, hp) in enumerate(zip(orig.paragraphs, human.paragraphs)):
    otext = op.text.strip()
    htext = hp.text.strip()
    if not otext or not htext:
        continue

    # Truncation: human text is significantly shorter but NOT a complete rewrite
    # Heuristic: if human text is >30% shorter and they share common start
    if len(htext) < len(otext) * 0.7 and len(htext) > 5:
        # Check if they share common prefix (same paragraph, partial revision)
        common_prefix_len = 0
        for a, b in zip(otext, htext):
            if a == b:
                common_prefix_len += 1
            else:
                break
        if common_prefix_len > 5:
            # Check if text got cut mid-sentence (ends without punctuation or mid-word)
            if not htext[-1] in '。！？）"""\'\'」』.）':
                truncation_warnings.append({
                    "chapter": "Ch1",
                    "paragraph_index": i + 1,
                    "dimension": "⚠️ 人工修订警告",
                    "check_item": "文本截断检测",
                    "severity": "critical",
                    "confidence": "high",
                    "source_quote": otext[:200],
                    "target_quote": htext[:200],
                    "issue": f"人工修订版此段疑似截断：原文{len(otext)}字 → 修订版{len(htext)}字（缩减{int((1-len(htext)/len(otext))*100)}%），且句末无标点。可能为修订操作失误导致文字丢失。",
                    "suggestion": "建议对照英文原文逐句核查此段，补齐被截断的内容。"
                })

# Also check for obvious missing fragments in human version
# e.g., P81 "非人群" missing "高加索"
known_fragments = [
    (81, "非高加索人群与皮肤类型", "非人群与皮肤类型", "疑似缺「高加索」三字"),
    (79, "一项小型研究对每个区域进行了6次治疗", "一项小型研上述区域6例", "「研究」→「研」截断，且后续文字重组后语义断裂"),
    (86, "Ulthera系统的一个常见抱怨", "治疗过程中产生的疼痛是患者对Ulthera系统的一个常见", "句末「常见」后缺名词（应为「常见抱怨」或「常见不满」）"),
    (101, "新胶原生成的必要组成部分", "新胶原蛋白生成的", "句末截断，缺「必要组成部分」"),
    (101, "是否会干扰最终结果", "是否会干扰最终", "句末截断，缺「结果」"),
    (100, "临床结果无差异", "临床无差异", "缺「结果」——应为「临床结果无差异」"),
]

for pi, orig_text, human_text, desc in known_fragments:
    truncation_warnings.append({
        "chapter": "Ch1",
        "paragraph_index": pi,
        "dimension": "⚠️ 人工修订警告",
        "check_item": "文本截断检测",
        "severity": "critical",
        "confidence": "high",
        "source_quote": orig_text[:200],
        "target_quote": human_text[:200],
        "issue": f"人工修订版此段文字截断：{desc}。此为修订操作失误，非翻译质量问题。",
        "suggestion": "建议对照英文原文还原完整中文表述。"
    })

print(f'Detected {len(truncation_warnings)} truncation warnings')

# ── Merge all ──
all_issues = existing + new_issues + truncation_warnings

# Deduplicate by paragraph_index + check_item
seen = set()
merged = []
for issue in all_issues:
    key = (issue['paragraph_index'], issue['check_item'], issue['issue'][:50])
    if key not in seen:
        seen.add(key)
        merged.append(issue)

output_data = {
    "issues": merged,
    "summary": {
        "total_issues": len(merged),
        "by_severity": {
            "critical": sum(1 for i in merged if i['severity'] == 'critical'),
            "medium": sum(1 for i in merged if i['severity'] == 'medium'),
            "low": sum(1 for i in merged if i['severity'] == 'low')
        },
        "by_dimension": {}
    }
}

for i in merged:
    dim = i['dimension']
    output_data['summary']['by_dimension'][dim] = output_data['summary']['by_dimension'].get(dim, 0) + 1

with open('cache/issues_phase4_v2_2.json', 'w', encoding='utf-8') as f:
    json.dump(output_data, f, ensure_ascii=False, indent=2)

print(f'\n=== MERGED ===')
print(f'Total: {len(merged)} issues ({len(existing)} v2.1 + {len(new_issues)} v2.2 + {len(truncation_warnings)} truncation warnings)')
print(f'By severity: {output_data["summary"]["by_severity"]}')
print(f'By dimension: {output_data["summary"]["by_dimension"]}')
print(f'\nSaved to cache/issues_phase4_v2_2.json')
