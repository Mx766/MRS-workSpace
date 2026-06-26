#!/usr/bin/env python3
"""
Phase 2: 文档拆分
Word: 按标题层级（Heading 1/2）拆分为章节，检测跨页表格
Excel: 按 Sheet 拆分，单 Sheet 超 200 行按行分块
PDF: 提取文本后按页面分块
输出: 每章结构化 JSON
"""

import argparse
import json
import re
from pathlib import Path
from typing import Any


# ─── Word 拆分 ───────────────────────────────────────────

def detect_heading_level(para) -> int | None:
    """检测段落的标题级别。返回 1-9 或 None"""
    style_name = (para.style.name if para.style and para.style.name else '').lower()
    # 标准 Heading 样式
    heading_match = re.match(r'heading\s*(\d+)', style_name)
    if heading_match:
        return int(heading_match.group(1))
    # 中文样式名
    if '标题' in style_name or 'heading' in style_name:
        for i in range(1, 10):
            if str(i) in style_name:
                return i
        return 1  # 无编号的标题默认 H1
    # 大纲级别 (outline level)
    try:
        ol = para.paragraph_format.outline_level
        if ol is not None and 0 <= ol <= 8:
            return ol + 1  # outline_level 0 = Heading 1
    except Exception:
        pass
    return None


def is_empty_paragraph(para) -> bool:
    """判断段落是否为空"""
    text = para.text.strip() if para.text else ''
    return len(text) == 0


def extract_table_data(table) -> dict:
    """提取 Word 表格的结构化数据"""
    rows_data = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cells.append({
                'text': cell.text.strip(),
                'col_span': _get_cell_span(cell, 'grid_span'),
            })
        rows_data.append(cells)
    return {
        'rows': len(rows_data),
        'cols': max(len(r) for r in rows_data) if rows_data else 0,
        'data': rows_data,
    }


def _get_cell_span(cell, attr: str) -> int:
    """获取单元格合并数"""
    try:
        tc = cell._tc
        tcPr = tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
        if tcPr is not None:
            grid_span = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
            if grid_span is not None:
                return int(grid_span.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1'))
    except Exception:
        pass
    return 1


def split_docx_by_headings(doc_path: str, max_chapter_chars: int = 5000) -> dict:
    """按标题拆分 Word 文档"""
    try:
        from docx import Document
    except ImportError:
        return {'error': '需要 python-docx: pip install python-docx'}

    try:
        doc = Document(doc_path)
    except Exception as e:
        return {'error': f'无法打开文件: {e}'}

    chapters = []
    current_chapter = None
    para_idx = 0
    heading_boundary = 2  # 默认按 H1/H2 切分

    for para in doc.paragraphs:
        para_idx += 1
        text = para.text.strip() if para.text else ''
        heading_level = detect_heading_level(para)

        # 判断是否开始新章节
        should_split = (heading_level is not None and heading_level <= heading_boundary)

        if should_split and current_chapter is not None and current_chapter['paragraphs']:
            chapters.append(current_chapter)
            current_chapter = None

        if current_chapter is None:
            current_chapter = {
                'chapter_id': f'Ch{len(chapters) + 1}',
                'title': text if heading_level else f'段落 {para_idx} 起',
                'heading_level': heading_level,
                'paragraphs': [],
                'tables': [],
                'start_para_idx': para_idx,
            }

        current_chapter['paragraphs'].append({
            'index': para_idx,
            'text': text,
            'heading_level': heading_level,
            'style': para.style.name if para.style else '',
        })

    # 收尾最后一章
    if current_chapter and current_chapter['paragraphs']:
        chapters.append(current_chapter)

    # 处理表格（关联到所在章节）
    table_index = 0
    for table in doc.tables:
        table_index += 1
        table_data = extract_table_data(table)
        # 通过表格前一个段落定位章节
        # 简化: 通过 XML 位置大致关联
        try:
            prev_para = table._element.getprevious()
            # 遍历找最近的段落
            associated_chapter = None
            if prev_para is not None:
                for ch in chapters:
                    for p in ch['paragraphs']:
                        # 简化：表格关联到最后一章（可通过元素索引精确化）
                        associated_chapter = ch
        except Exception:
            associated_chapter = chapters[-1] if chapters else None

        if associated_chapter:
            associated_chapter['tables'].append({
                'table_index': table_index,
                **table_data
            })

    # 检测跨页表格标记
    cross_page_tables = detect_cross_page_tables(doc)

    # 修正章节编号（用原标题）
    for i, ch in enumerate(chapters):
        ch['chapter_number'] = i + 1
        # 如果标题仍是自动生成的，尝试取第一个有意义的标题段落
        if ch['title'].startswith('段落 '):
            for p in ch['paragraphs']:
                if p['heading_level'] and p['text']:
                    ch['title'] = p['text']
                    break

    return {
        'doc_path': doc_path,
        'total_paragraphs': para_idx,
        'total_tables': table_index,
        'chapter_count': len(chapters),
        'chapters': chapters,
        'cross_page_tables': cross_page_tables,
    }


def detect_cross_page_tables(doc) -> list[dict]:
    """检测可能因跨页被不合理拆分的表格"""
    issues = []
    try:
        tables = doc.tables
        for i, table in enumerate(tables):
            rows = len(table.rows)
            if rows <= 1:
                continue
            # 检查第一行是否为表头样式（加粗/底色）
            first_row_has_header_style = False
            try:
                for cell in table.rows[0].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.bold or run.font.highlight_color is not None:
                                first_row_has_header_style = True
                                break
            except Exception:
                pass

            # 检查表格是否在页面底部（通过最后一个单元格内容判断）
            last_row_text = ' '.join(c.text for c in table.rows[-1].cells).strip()
            first_row_text = ' '.join(c.text for c in table.rows[0].cells).strip()

            # 简单启发式: 行数少（1-3行）+ 第一行像表头 → 可能被拆分
            if rows <= 3 and first_row_has_header_style and last_row_text:
                issues.append({
                    'table_index': i + 1,
                    'rows': rows,
                    'warning': '表格行数较少且有表头样式，可能因跨页被不合理拆分',
                    'first_row': first_row_text[:100],
                    'last_row': last_row_text[:100],
                })

            # 检查是否有重复表头（连续出现的相似表头）
            for j in range(1, rows):
                prev_text = ' '.join(c.text for c in table.rows[j-1].cells).strip()
                curr_text = ' '.join(c.text for c in table.rows[j].cells).strip()
                if prev_text and curr_text and _text_similarity(prev_text, curr_text) > 0.8:
                    issues.append({
                        'table_index': i + 1,
                        'rows': rows,
                        'warning': f'第{j}行与第{j+1}行高度相似，可能是跨页表头重复',
                        'row_a': prev_text[:100],
                        'row_b': curr_text[:100],
                    })
                    break
    except Exception:
        pass

    return issues


def _text_similarity(a: str, b: str) -> float:
    """简单的文本相似度"""
    a_words = set(a.lower().split())
    b_words = set(b.lower().split())
    if not a_words or not b_words:
        return 0.0
    intersection = a_words & b_words
    return len(intersection) / min(len(a_words), len(b_words))


# ─── Excel 拆分 ──────────────────────────────────────────

def split_excel(filepath: str, max_rows_per_chunk: int = 200) -> dict:
    """按 Sheet 拆分 Excel，大 Sheet 按行分块"""
    try:
        import openpyxl
    except ImportError:
        return {'error': '需要 openpyxl: pip install openpyxl'}

    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    except Exception as e:
        return {'error': f'无法打开: {e}'}

    chunks = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        header_row = [str(c) if c is not None else '' for c in rows[0]]
        data_rows = rows[1:]
        total_data_rows = len(data_rows)

        if total_data_rows <= max_rows_per_chunk:
            chunks.append({
                'chunk_id': f'{sheet_name}',
                'sheet': sheet_name,
                'header': header_row,
                'rows': [list(r) for r in data_rows],
                'total_rows': total_data_rows,
                'row_range': f'1-{total_data_rows}',
            })
        else:
            # 分块
            for chunk_start in range(0, total_data_rows, max_rows_per_chunk):
                chunk_end = min(chunk_start + max_rows_per_chunk, total_data_rows)
                chunk_rows = data_rows[chunk_start:chunk_end]
                chunks.append({
                    'chunk_id': f'{sheet_name}_行{chunk_start + 1}-{chunk_end}',
                    'sheet': sheet_name,
                    'header': header_row,
                    'rows': [list(r) for r in chunk_rows],
                    'total_rows': len(chunk_rows),
                    'row_range': f'{chunk_start + 1}-{chunk_end}',
                })

    sheetnames = list(wb.sheetnames)
    wb.close()

    return {
        'filepath': filepath,
        'sheets': sheetnames,
        'total_chunks': len(chunks),
        'chunks': chunks,
    }


# ─── PDF 原文处理 ─────────────────────────────────────────

def extract_pdf_text(filepath: str) -> dict:
    """提取 PDF 文本，按页面分块"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {'error': '需要 PyMuPDF: pip install PyMuPDF'}

    try:
        doc = fitz.open(filepath)
    except Exception as e:
        return {'error': f'无法打开: {e}'}

    pages = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text('text')
        pages.append({
            'page_num': page_num + 1,
            'text': text.strip(),
            'char_count': len(text.strip()),
        })

    doc.close()

    return {
        'filepath': filepath,
        'total_pages': len(pages),
        'pages': pages,
    }


# ─── 主入口 ───────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='文档拆分')
    parser.add_argument('--input', '-i', required=True, help='输入文件路径')
    parser.add_argument('--type', '-t', choices=['docx', 'xlsx', 'pdf'], help='文件类型（自动检测）')
    parser.add_argument('--max-chars', type=int, default=5000, help='Word 章节最大字符数')
    parser.add_argument('--max-rows', type=int, default=200, help='Excel 每块最大行数')
    parser.add_argument('--output', '-o', help='输出 JSON 路径')
    args = parser.parse_args()

    filepath = args.input
    suffix = Path(filepath).suffix.lower()

    if args.type:
        file_type = args.type
    elif suffix == '.docx':
        file_type = 'docx'
    elif suffix in ('.xlsx', '.xls'):
        file_type = 'xlsx'
    elif suffix == '.pdf':
        file_type = 'pdf'
    else:
        print(json.dumps({'error': f'无法识别的文件类型: {suffix}'}, ensure_ascii=False))
        sys.exit(1)

    if file_type == 'docx':
        result = split_docx_by_headings(filepath, args.max_chars)
    elif file_type == 'xlsx':
        result = split_excel(filepath, args.max_rows)
    else:
        result = extract_pdf_text(filepath)

    output_json = json.dumps(result, ensure_ascii=False, indent=2)
    if args.output:
        Path(args.output).write_text(output_json, encoding='utf-8')
        print(f'已写入: {args.output}')
    else:
        sys.stdout.reconfigure(encoding='utf-8')
        print(output_json)


import sys
if __name__ == '__main__':
    main()
