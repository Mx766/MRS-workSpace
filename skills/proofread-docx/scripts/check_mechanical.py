#!/usr/bin/env python3
"""
Phase 3: 机械检查
数字一致性 / 符号与特殊字符 / 术语库匹配 / 格式合规 / 缩写单位检查
输入: 原文文件 + 译文文件路径（或已拆分的章节 JSON）
输出: 结构化疑点 JSON 报告
"""

import argparse
import json
import re
import sys
from pathlib import Path
from collections import Counter

sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')


# ═══════════════════════════════════════════════════════════
# 3.1 数字一致性检查
# ═══════════════════════════════════════════════════════════

# 数字匹配模式
NUMBER_PATTERNS = {
    'integer': re.compile(r'(?<![a-zA-Z0-9])\d+(?![a-zA-Z0-9])'),
    'decimal': re.compile(r'(?<![a-zA-Z0-9])\d+\.\d+(?![a-zA-Z0-9])'),
    'percentage': re.compile(r'(?<![a-zA-Z0-9])\d+(?:\.\d+)?\s*%'),
    'temperature': re.compile(r'(?<![a-zA-Z0-9])\d+(?:\.\d+)?\s*[°℃](?:[CFKcfk])?'),
    'range': re.compile(r'(?<![a-zA-Z0-9])\d+(?:\.\d+)?\s*[-–—至~]\s*\d+(?:\.\d+)?'),
    'scientific': re.compile(r'\d+(?:\.\d+)?\s*[×xX]\s*10\^?\d+'),
    'date_iso': re.compile(r'\d{4}[-/]\d{2}[-/]\d{2}'),
    'date_us': re.compile(r'\d{2}/\d{2}/\d{4}'),
}

# 范围表达模式（用于全局统一性检查）
RANGE_PATTERNS = {
    'dash': re.compile(r'\d+\s*[-–]\s*\d+'),        # 1-5, 1–5
    'tilde': re.compile(r'\d+\s*[～~]\s*\d+'),        # 1～5, 1~5
    'chinese': re.compile(r'\d+\s*至\s*\d+'),          # 1至5
}


def extract_numbers(text: str) -> dict:
    """从文本提取各类数字"""
    result = {}
    for name, pattern in NUMBER_PATTERNS.items():
        matches = pattern.findall(text)
        result[name] = [m.strip() if isinstance(m, str) else m for m in matches]
    return result


def extract_range_expressions(text: str) -> dict:
    """提取范围表达，统计各形式出现次数"""
    result = {}
    for name, pattern in RANGE_PATTERNS.items():
        matches = pattern.findall(text)
        result[name] = len(matches)
    return result


def _context_around(text: str, value: str, radius: int = 40) -> str:
    """在 text 中找到 value 的位置，提取前后各 radius 字符的上下文片段。
    用于机械检查中给译者提供定位信息——否则译者只看到"缺了某个符号"但不知道在哪。
    """
    idx = text.find(value)
    if idx == -1:
        return ''
    start = max(0, idx - radius)
    end = min(len(text), idx + len(value) + radius)
    snippet = text[start:end]
    # 标记前后截断
    prefix = '…' if start > 0 else ''
    suffix = '…' if end < len(text) else ''
    return prefix + snippet + suffix


# ══════════════════════════════════════════════════════════════
# v2.18: 句子计数 — 用于漏译检测（标记，非硬阻断）
# ══════════════════════════════════════════════════════════════

def count_sentences(text: str, lang: str = 'auto') -> int:
    """计算文本中的句子数。

    中文：基于 。！？… 和中文后的英文句号分割。
    英文：基于 .!? + 空格/结尾分割，过滤常见缩写。

    Args:
        text: 输入文本
        lang: 'zh' | 'en' | 'auto'（自动检测）

    Returns:
        句子数，空文本返回 0
    """
    if not text or not text.strip():
        return 0
    text = text.strip()

    # 自动检测语言
    if lang == 'auto':
        cn_chars = len(re.findall(r'[一-鿿]', text))
        lang = 'zh' if cn_chars > len(text) * 0.3 else 'en'

    if lang == 'zh':
        # 中文分句：。！？… 和中文后的英文句号
        sentences = re.split(r'[。！？…]|(?<=[一-鿿])\.(?!\s*[a-zA-Z])', text)
        return len([s for s in sentences if s.strip()])
    else:
        # 英文分句：.!? + 空格过滤常见缩写
        abbr_words = {'Mr', 'Mrs', 'Ms', 'Dr', 'Prof', 'etc', 'vs', 'Fig',
                      'Eq', 'et', 'al', 'Vol', 'No', 'pp', 'e.g', 'i.e'}
        # 先将缩写中的句号替换为占位符
        protected = text
        for abbr in sorted(abbr_words, key=len, reverse=True):
            protected = protected.replace(abbr + '.', abbr + '<DOT>')
        sentences = re.split(r'(?<=[.!?])\s+', protected)
        count = 0
        for s in sentences:
            s = s.strip()
            if s:
                count += 1
        return max(count, 1)  # 有文本则至少 1 句


def check_sentence_alignment(
    source_text: str, target_text: str, para_index: int = 0
) -> list[dict]:
    """检测段落级句子数量显著不匹配（可能漏译/断句错误）。

    这是启发式标记——不是硬错误。句子分割不完美（尤其是编号列表、
    技术缩写等），需要 AI 语义验证。

    阈值：差异 >= 3 句 或 比例 > 2:1

    Returns:
        标记列表（通常 0-1 条）
    """
    flags = []
    src_count = count_sentences(source_text)
    tgt_count = count_sentences(target_text)

    if src_count == 0 or tgt_count == 0:
        return flags

    diff = abs(src_count - tgt_count)
    ratio = max(src_count, tgt_count) / max(min(src_count, tgt_count), 1)

    if diff >= 3 or (diff >= 2 and ratio > 2.0):
        flags.append({
            'paragraph_index': para_index,
            'type': 'sentence_count_mismatch',
            'source_sentences': src_count,
            'target_sentences': tgt_count,
            'diff': diff,
            'ratio': round(ratio, 1),
            'severity': 'low',
            'check': (
                f'句子数差异显著: 原文 {src_count} 句 vs 译文 {tgt_count} 句 '
                f'(差值 {diff}, 比例 {ratio}:1) —— 可能存在漏译或断句错误，请AI逐句核对'
            ),
        })
    return flags


def check_numbers(source_text: str, target_text: str, para_index: int = 0) -> list[dict]:
    """对比原文-译文段落中的数字一致性"""
    issues = []
    src_nums = extract_numbers(source_text)
    tgt_nums = extract_numbers(target_text)

    # 整数和浮点数比对
    src_integers = set(src_nums.get('integer', []))
    tgt_integers = set(tgt_nums.get('integer', []))
    src_decimals = set(src_nums.get('decimal', []))
    tgt_decimals = set(tgt_nums.get('decimal', []))

    # 检查是否缺失（允许合理近似，如 1,000 → 1000）
    for num in src_integers:
        if num not in tgt_integers:
            # 容忍千分位差异
            alt = num.replace(',', '')
            if alt not in tgt_integers:
                ctx = _context_around(source_text, num)
                issues.append({
                    'paragraph_index': para_index,
                    'type': 'number_missing',
                    'source_value': num,
                    'source_quote': ctx,
                    'target_quote': target_text[:120] + ('…' if len(target_text) > 120 else ''),
                    'severity': 'medium',
                    'check': f'数字 "{num}" 在译文中可能缺失（原文上下文：{ctx}）',
                })

    for num in src_decimals:
        if num not in tgt_decimals:
            ctx = _context_around(source_text, num)
            issues.append({
                'paragraph_index': para_index,
                'type': 'decimal_mismatch',
                'source_value': num,
                'source_quote': ctx,
                'target_quote': target_text[:120] + ('…' if len(target_text) > 120 else ''),
                'severity': 'critical',
                'check': f'小数 "{num}" 在译文中不一致或缺失（原文上下文：{ctx}）',
            })

    # 百分比
    src_pct = [float(re.findall(r'\d+(?:\.\d+)?', m)[0]) for m in src_nums.get('percentage', [])]
    tgt_pct = [float(re.findall(r'\d+(?:\.\d+)?', m)[0]) for m in tgt_nums.get('percentage', [])]
    for i, val in enumerate(src_pct):
        if i < len(tgt_pct) and abs(val - tgt_pct[i]) > 0.01:
            issues.append({
                'paragraph_index': para_index,
                'type': 'percentage_mismatch',
                'source_value': f'{val}%',
                'target_value': f'{tgt_pct[i]}%',
                'severity': 'critical',
                'check': f'百分比不一致: {val}% → {tgt_pct[i]}%',
            })

    # 温度
    src_temps = src_nums.get('temperature', [])
    tgt_temps = tgt_nums.get('temperature', [])
    if len(src_temps) != len(tgt_temps):
        issues.append({
            'paragraph_index': para_index,
            'type': 'temperature_mismatch',
            'source_value': ', '.join(src_temps),
            'target_value': ', '.join(tgt_temps),
            'severity': 'critical',
            'check': f'温度值数量不一致: {len(src_temps)} vs {len(tgt_temps)}',
        })

    return issues


# ═══════════════════════════════════════════════════════════
# 3.2 符号与特殊字符检查
# ═══════════════════════════════════════════════════════════

SPECIAL_SYMBOLS = {
    '°': '度符号',
    '®': '注册商标',
    '™': '商标',
    '±': '正负号',
    '≤': '小于等于',
    '≥': '大于等于',
    '×': '乘号',
    'µ': '微(μ)',
    'α': 'alpha',
    'β': 'beta',
    'γ': 'gamma',
    'δ': 'delta',
}

# ── 符号等价替换：这些替代写法在技术文档中是合法等价形式，不应报缺失 ──
# 格式: {Unicode符号: [等效ASCII/替代写法]}
SYMBOL_EQUIVALENTS = {
    '≥': ['>='],
    '≤': ['<='],
    '≠': ['!='],
    '±': ['+/-'],
    '×': ['x', 'X'],
    '²': ['^2'],
    '³': ['^3'],
}

CAS_PATTERN = re.compile(r'CAS\s*\d{2,7}-\d{2}-\d{1}')
CHEM_FORMULA_PATTERN = re.compile(r'\b([A-Z][a-z]?\d*){2,}\b')
SUPERSCRIPT_PATTERN = re.compile(r'(m|cm|mm|km)\^?([²³⁴⁵⁶⁷⁸⁹]|\d+)')
PUNCTUATION_MIXED_CN = re.compile(r'[一-鿿][,.;:!?][一-鿿]')  # 中文中夹英文标点
PUNCTUATION_MIXED_EN = re.compile(r'[a-zA-Z][，。；：！？][a-zA-Z]')           # 英文中夹中文标点
PROPER_NOUN_PATTERNS = [
    re.compile(r'ISO\s+\d{4,6}'),
    re.compile(r'(?:IEC|IEEE|ASTM|ICH|NMPA|FDA|EMA|WHO)\s+\d+'),
    re.compile(r'§\s*\d+'),
    re.compile(r'(?:GB|GB/T|YY|YY/T)\s+\d{4,}'),
]


def check_symbols(source_text: str, target_text: str, para_index: int = 0,
                  direction: str = 'en→zh') -> list[dict]:
    """检查符号和特殊字符"""
    issues = []

    # 特殊符号检查（含等价替代宽容）
    for symbol, name in SPECIAL_SYMBOLS.items():
        in_source = symbol in source_text
        in_target = symbol in target_text
        if in_source and not in_target:
            # 检查是否有等价替代写法（如 ≥→>=, ≤→<=）
            equivalents = SYMBOL_EQUIVALENTS.get(symbol, [])
            has_equivalent = any(eq in target_text for eq in equivalents)
            if has_equivalent:
                continue  # 合法替代写法，不报缺失
            ctx = _context_around(source_text, symbol)
            issues.append({
                'paragraph_index': para_index,
                'type': 'symbol_missing',
                'symbol': symbol,
                'symbol_name': name,
                'source_quote': ctx,
                'target_quote': target_text[:120] + ('…' if len(target_text) > 120 else ''),
                'severity': 'medium',
                'check': f'符号 "{symbol}"({name}) 在原文中存在但译文中缺失（原文上下文：{ctx}）',
            })

    # CAS 号
    src_cas = CAS_PATTERN.findall(source_text)
    tgt_cas = CAS_PATTERN.findall(target_text)
    for cas in src_cas:
        if cas not in tgt_cas:
            ctx = _context_around(source_text, cas)
            issues.append({
                'paragraph_index': para_index,
                'type': 'cas_missing',
                'source_value': cas,
                'source_quote': ctx,
                'target_quote': target_text[:120] + ('…' if len(target_text) > 120 else ''),
                'severity': 'critical',
                'check': f'CAS号 "{cas}" 在译文中缺失或不一致（原文上下文：{ctx}）',
            })

    # 法规/标准编号
    for pattern in PROPER_NOUN_PATTERNS:
        src_matches = pattern.findall(source_text)
        tgt_matches = pattern.findall(target_text)
        for m in src_matches:
            if m not in tgt_matches:
                ctx = _context_around(source_text, m)
                issues.append({
                    'paragraph_index': para_index,
                    'type': 'regulation_missing',
                    'source_value': m,
                    'source_quote': ctx,
                    'target_quote': target_text[:120] + ('…' if len(target_text) > 120 else ''),
                    'severity': 'critical',
                    'check': f'法规/标准编号 "{m}" 在译文中缺失（原文上下文：{ctx}）',
                })

    # 标点混排检查
    if direction == 'en→zh':
        # 中文译文不应有英文标点夹在中文中
        mixed = PUNCTUATION_MIXED_CN.findall(target_text)
        if mixed:
            # 提取具体标点
            bad_puncts = set()
            for m in mixed:
                bad_puncts.add(m[1])
            issues.append({
                'paragraph_index': para_index,
                'type': 'punctuation_mixed',
                'detail': f'中文译文中使用了英文标点: {", ".join(bad_puncts)}',
                'severity': 'medium',
                'check': '标点混排：中文应使用中文标点（，。；：）',
            })
    elif direction == 'zh→en':
        mixed = PUNCTUATION_MIXED_EN.findall(target_text)
        if mixed:
            issues.append({
                'paragraph_index': para_index,
                'type': 'punctuation_mixed',
                'severity': 'medium',
                'check': '标点混排：英文应使用英文标点',
            })

    return issues


# ═══════════════════════════════════════════════════════════
# 3.3 术语库匹配
# ═══════════════════════════════════════════════════════════

def check_glossary(source_text: str, target_text: str, glossary: dict, para_index: int = 0,
                   domain_filter: str = None, strict_short_terms: bool = False,
                   client_glossary: dict = None) -> dict:
    """
    检查术语库合规性。
    规则: 若原文段落中出现术语库中的英文术语，则译文中必须包含规定的对应译法。

    glossary: {normalized_source: {source, target, domain}}
    domain_filter: 仅匹配该领域或"通用"的术语（None=全部匹配）
    strict_short_terms: 短词（≤3字符）是否保持原严重度（默认降级为low）
    client_glossary: {normalized_source: {source, target, note}} — 客户术语需求（最高优先级）
    返回: {violations, unknown_terms, client_violations}
    """
    violations = []
    client_violations = []
    unknown_terms = []

    source_lower = source_text.lower()
    target_lower = target_text.lower()

    # ── 1. 客户术语优先检查（Fix E: 客户要求指定词翻译成什么）──
    client_matched_sources = set()  # 客户术语已覆盖的源词
    if client_glossary:
        for key, cterm in client_glossary.items():
            src = cterm['source']
            required_tgt = cterm['target'].lower()
            if re.search(r'\b' + re.escape(src) + r'\b', source_lower, re.IGNORECASE):
                client_matched_sources.add(key)
                if required_tgt not in target_lower:
                    client_violations.append({
                        'paragraph_index': para_index,
                        'type': 'client_glossary_violation',
                        'source_term': src,
                        'required_target': cterm['target'],
                        'client_note': cterm.get('note', ''),
                        'severity': 'critical',
                        'check': f'【客户要求】术语 "{src}" 必须译为 "{cterm["target"]}"，但译文中未找到',
                    })

    # ── 2. 通用术语库检查（排除已被客户术语覆盖的源词）──
    if glossary:
        for key, term in glossary.items():
            src = term['source']
            expected_tgt = term['target'].lower()

            # 跳过已由客户术语覆盖的源词（避免重复检查）
            if key in client_matched_sources:
                continue

            # 领域过滤：传了 domain_filter 则只匹配该领域或"通用"术语
            if domain_filter:
                term_domain = term.get('domain', '')
                if term_domain not in (domain_filter, '通用', ''):
                    continue

            # 词边界匹配（\b）：杜绝子串误报（arm→harmonized, SE→USE）
            if re.search(r'\b' + re.escape(src) + r'\b', source_lower, re.IGNORECASE):
                if expected_tgt not in target_lower:
                    sev = 'critical'
                    # 短词（≤3字符）自动降级：占误报67%，领域错配风险极高
                    if len(src) <= 3 and not strict_short_terms:
                        sev = 'low'
                    violations.append({
                        'paragraph_index': para_index,
                        'type': 'glossary_violation',
                        'source_term': src,
                        'expected_target': term['target'],
                        'domain': term.get('domain', ''),
                        'severity': sev,
                        'check': f'术语 "{src}" 应译为 "{term["target"]}"，但译文中未找到',
                    })
                    if len(src) <= 3 and not strict_short_terms:
                        violations[-1]['_short_term_warning'] = True

    return {'violations': violations, 'unknown_terms': unknown_terms, 'client_violations': client_violations}


# ═══════════════════════════════════════════════════════════
# 3.4 格式合规检查
# ═══════════════════════════════════════════════════════════

def check_format(source_doc_path: str, target_doc_path: str) -> list[dict]:
    """检查原文和译文之间的格式一致性"""
    issues = []
    try:
        from docx import Document
    except ImportError:
        return [{'type': 'format_error', 'check': '需要 python-docx'}]

    try:
        src_doc = Document(source_doc_path)
        tgt_doc = Document(target_doc_path)
    except Exception as e:
        return [{'type': 'format_error', 'check': f'无法打开文件: {e}'}]

    # 段落数和样式对比
    src_paras = src_doc.paragraphs
    tgt_paras = tgt_doc.paragraphs

    # 标题层级对比
    src_headings = [(i, p.style.name) for i, p in enumerate(src_paras)
                    if p.style.name.startswith('Heading')]
    tgt_headings = [(i, p.style.name) for i, p in enumerate(tgt_paras)
                    if p.style.name.startswith('Heading')]

    if len(src_headings) != len(tgt_headings):
        issues.append({
            'type': 'heading_count',
            'source_count': len(src_headings),
            'target_count': len(tgt_headings),
            'severity': 'low',
            'check': f'标题数量不一致: 原文 {len(src_headings)} vs 译文 {len(tgt_headings)}',
        })

    # 表格对比
    src_tables = src_doc.tables
    tgt_tables = tgt_doc.tables
    if len(src_tables) != len(tgt_tables):
        issues.append({
            'type': 'table_count',
            'source_count': len(src_tables),
            'target_count': len(tgt_tables),
            'severity': 'low',
            'check': f'表格数量不一致: 原文 {len(src_tables)} vs 译文 {len(tgt_tables)}',
        })

    # 逐表格行列对比
    for i, (st, tt) in enumerate(zip(src_tables, tgt_tables)):
        src_rows = len(st.rows)
        tgt_rows = len(tt.rows)
        if src_rows != tgt_rows:
            issues.append({
                'type': 'table_rows',
                'table_index': i + 1,
                'source_rows': src_rows,
                'target_rows': tgt_rows,
                'severity': 'low',
                'check': f'表格{i+1}行数不一致: {src_rows} vs {tgt_rows}',
            })

        if st.rows and tt.rows:
            src_cols = len(st.rows[0].cells)
            tgt_cols = len(tt.rows[0].cells)
            if src_cols != tgt_cols:
                issues.append({
                    'type': 'table_cols',
                    'table_index': i + 1,
                    'source_cols': src_cols,
                    'target_cols': tgt_cols,
                    'severity': 'low',
                    'check': f'表格{i+1}列数不一致: {src_cols} vs {tgt_cols}',
                })

    # 加粗/斜体统计
    src_bold_count = 0
    tgt_bold_count = 0
    for p in src_paras:
        for r in p.runs:
            if r.bold:
                src_bold_count += 1
    for p in tgt_paras:
        for r in p.runs:
            if r.bold:
                tgt_bold_count += 1

    if src_bold_count > 0 and tgt_bold_count < src_bold_count * 0.5:
        issues.append({
            'type': 'bold_loss',
            'source_bold_runs': src_bold_count,
            'target_bold_runs': tgt_bold_count,
            'severity': 'low',
            'check': f'加粗格式大量丢失: 原文 {src_bold_count} → 译文 {tgt_bold_count}',
        })

    return issues


# ═══════════════════════════════════════════════════════════
# 3.5 缩写单位检查
# ═══════════════════════════════════════════════════════════

# 常见缩写单位（应保留不译）
ABBREVIATED_UNITS = re.compile(
    r'\d+(?:\.\d+)?\s*(?:'
    r'mg|kg|g|µg|mcg|ng|pg|'          # 质量
    r'mL|L|µL|dL|cL|'                  # 体积
    r'mm|cm|m|km|nm|µm|'               # 长度
    r'mmHg|Pa|kPa|hPa|atm|psi|bar|'    # 压力
    r'°C|°F|K|'                         # 温度
    r'ms|s|min|h|hr|hrs|'              # 时间
    r'Hz|kHz|MHz|GHz|'                 # 频率
    r'mol|mmol|nmol|'                  # 物质的量
    r'IU|CFU|ppb|ppm|ppmv|'            # 其他
    r'W|kW|MW|J|kJ|eV|'                # 能量
    r'V|mV|kV|A|mA|Ω|'                 # 电学
    r'M|mM|µM|nM'                      # 浓度
    r')',
    re.IGNORECASE
)

# 全写单位（应翻译）
FULL_UNIT_PAIRS = [
    ('milligram', '毫克'),
    ('kilogram', '千克'),
    ('gram', '克'),
    ('microgram', '微克'),
    ('milliliter', '毫升'),
    ('liter', '升'),
    ('millimeter', '毫米'),
    ('centimeter', '厘米'),
    ('meter', '米'),
    ('kilometer', '千米'),
    ('degree Celsius', '摄氏度'),
    ('degree centigrade', '摄氏度'),
    ('Fahrenheit', '华氏度'),
    ('second', '秒'),
    ('minute', '分钟'),
    ('hour', '小时'),
    ('mole', '摩尔'),
    ('millimole', '毫摩尔'),
]


def check_units(source_text: str, target_text: str, para_index: int = 0) -> list[dict]:
    """检查缩写单位保留和全写单位翻译"""
    issues = []

    # 检查全写单位是否翻译
    for en_unit, zh_unit in FULL_UNIT_PAIRS:
        pattern = re.compile(rf'\d+(?:\.\d+)?\s*{en_unit}s?\b', re.IGNORECASE)
        src_matches = pattern.findall(source_text)
        if src_matches:
            # 检查译文是否包含对应中文单位
            if zh_unit not in target_text:
                issues.append({
                    'paragraph_index': para_index,
                    'type': 'full_unit_untranslated',
                    'source_value': src_matches[0],
                    'expected': zh_unit,
                    'severity': 'low',
                    'check': f'全写单位 "{src_matches[0]}" 未翻译为 "{zh_unit}"',
                })

    # 检查缩写单位是否被错误翻译（如 5 mg/kg 译成了 5 毫克/千克）
    src_abbr = ABBREVIATED_UNITS.findall(source_text)
    for abbr in src_abbr:
        # 提取数字和单位
        parts = re.match(r'(\d+(?:\.\d+)?)\s*(.+)', abbr)
        if parts:
            num, unit = parts.groups()
            unit_lower = unit.strip().lower()
            # 简单检查：如果译文中出现了对应中文单位（如 mg → 毫克）
            zh_unit_map = {'mg': '毫克', 'kg': '千克', 'g': '克', 'ml': '毫升', 'l': '升',
                           'mm': '毫米', 'cm': '厘米', 'm': '米', 'km': '千米',
                           '°c': '℃', '°f': '℉'}
            zh_unit = zh_unit_map.get(unit_lower)
            if zh_unit and zh_unit in target_text and num in target_text:
                # 缩写被翻译了，标记
                pass  # 不强制报错，可能有合理场景（如全中文文档）

    return issues


# ═══════════════════════════════════════════════════════════
# 主入口
# ═══════════════════════════════════════════════════════════

def _formats_match(source_path: str, target_path: str) -> bool:
    """检测原文和译文是否为同格式（可做逐段对齐）"""
    src_sfx = Path(source_path).suffix.lower()
    tgt_sfx = Path(target_path).suffix.lower()
    # 同格式对：docx↔docx, xlsx↔xlsx
    return src_sfx == tgt_sfx


# ── 假阳性过滤模式 ──
_FP_DOI = re.compile(r'10\.\d{4,}/')
_FP_DATE = re.compile(
    r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}',
    re.IGNORECASE)
_FP_REF_RANGE = re.compile(r'\b\d{2,4}[-–]\d{1,4}\b')
_FP_REF_SECTION = re.compile(r'(References|REFERENCES|参考文献|Bibliography)', re.IGNORECASE)


def _is_false_positive_number(num_str: str, page_text: str) -> tuple[bool, str]:
    """判断一个数字是否为假阳性。返回 (skip, reason)。"""
    # Skip large page-number-like numbers (>500)
    try:
        if int(num_str) > 500:
            return (True, '>500')
    except ValueError:
        pass
    # DOI pattern in local context
    if _FP_DOI.search(page_text):
        return (True, 'DOI/reference identifier')
    # Date pattern
    if _FP_DATE.search(page_text):
        return (True, 'date component')
    # Reference page range (e.g. "164–71")
    try:
        int_val = int(num_str)
        if 1 <= int_val <= 2000:
            for m in _FP_REF_RANGE.finditer(page_text):
                parts = re.split(r'[-–]', m.group())
                if len(parts) == 2:
                    try:
                        if int_val in (int(parts[0]), int(parts[1])):
                            return (True, 'reference page number')
                    except ValueError:
                        pass
    except ValueError:
        pass
    return (False, '')


def _check_perpage_numbers(source_paras: list[dict], full_tgt: str) -> list[dict]:
    """跨格式模式：逐页数字一致性检查，记录 source_page，过滤假阳性"""
    issues = []
    tgt_nums = extract_numbers(full_tgt)
    tgt_ints = set(tgt_nums.get('integer', []))
    tgt_decimals = set(tgt_nums.get('decimal', []))
    tgt_pct_values = set()
    for m in tgt_nums.get('percentage', []):
        nums = re.findall(r'\d+(?:\.\d+)?', m)
        if nums:
            tgt_pct_values.add(float(nums[0]))

    for page_data in source_paras:
        page_text = page_data.get('text', '') if isinstance(page_data, dict) else str(page_data)
        page_num = page_data.get('page', 0) if isinstance(page_data, dict) else 0
        src_nums = extract_numbers(page_text)

        for num in src_nums.get('integer', []):
            if num not in tgt_ints:
                skip, reason = _is_false_positive_number(num, page_text)
                if skip:
                    continue
                issues.append({
                    'paragraph_index': 0,
                    'source_page': page_num,
                    'type': 'number_missing',
                    'source_value': num,
                    'severity': 'medium',
                    'check': f'数字 "{num}" 在译文中可能缺失（跨格式匹配，仅供参考）',
                })

        for num in src_nums.get('decimal', []):
            if num not in tgt_decimals:
                skip, reason = _is_false_positive_number(num, page_text)
                if skip:
                    continue
                issues.append({
                    'paragraph_index': 0,
                    'source_page': page_num,
                    'type': 'decimal_mismatch',
                    'source_value': num,
                    'severity': 'medium',
                    'check': f'小数 "{num}" 在译文中可能缺失（跨格式匹配，仅供参考）',
                })

        src_pct_values = set()
        for m in src_nums.get('percentage', []):
            nums = re.findall(r'\d+(?:\.\d+)?', m)
            if nums:
                src_pct_values.add(float(nums[0]))
        for val in src_pct_values:
            found = any(abs(val - tv) < 0.1 for tv in tgt_pct_values)
            if not found:
                issues.append({
                    'paragraph_index': 0,
                    'source_page': page_num,
                    'type': 'percentage_mismatch',
                    'source_value': f'{val}%',
                    'severity': 'medium',
                    'check': f'百分比 "{val}%" 在译文中可能缺失（跨格式匹配，仅供参考）',
                })

    return issues


def _check_perpage_glossary(source_paras: list[dict], full_tgt: str, glossary: dict,
                          domain_filter: str = None, strict_short_terms: bool = False) -> list[dict]:
    """跨格式模式：逐页术语合规检查，记录 source_page"""
    violations = []
    target_lower = full_tgt.lower()

    for page_data in source_paras:
        page_text = page_data.get('text', '') if isinstance(page_data, dict) else str(page_data)
        page_num = page_data.get('page', 0) if isinstance(page_data, dict) else 0
        page_lower = page_text.lower()

        for key, term in glossary.items():
            src = term['source']
            expected_tgt = term['target'].lower()

            # 领域过滤
            if domain_filter:
                term_domain = term.get('domain', '')
                if term_domain not in (domain_filter, '通用', ''):
                    continue

            # 词边界匹配
            if re.search(r'\b' + re.escape(src) + r'\b', page_lower, re.IGNORECASE):
                if expected_tgt not in target_lower:
                    sev = 'medium'
                    if len(src) <= 3 and not strict_short_terms:
                        sev = 'low'
                    violation = {
                        'paragraph_index': 0,
                        'source_page': page_num,
                        'type': 'glossary_violation',
                        'source_term': src,
                        'expected_target': term['target'],
                        'domain': term.get('domain', ''),
                        'severity': sev,
                        'check': f'术语 "{src}" 应译为 "{term["target"]}"，但译文中未找到（跨格式匹配，仅供参考）',
                    }
                    if len(src) <= 3 and not strict_short_terms:
                        violation['_short_term_warning'] = True
                    violations.append(violation)

    return violations


def check_all(source_path: str, target_path: str,
              glossary: dict = None, direction: str = 'en→zh',
              domain_filter: str = None, strict_short_terms: bool = False,
              client_glossary: dict = None) -> dict:
    """
    执行全部机械检查。
    source_path: 原文文件（.docx/.xlsx/.txt/.pdf）
    target_path: 译文文件（.docx/.xlsx/.txt）
    glossary: 预加载的术语表
    direction: 校对方向

    同格式（如 DOCX↔DOCX）: 逐段 1:1 对齐检查
    跨格式（如 PDF→DOCX）: 全文匹配模式，结果仅供参考
    """
    result = {
        'number_mismatches': [],
        'symbol_issues': [],
        'glossary_violations': [],
        'client_glossary_violations': [],
        'format_issues': [],
        'unit_issues': [],
        'sentence_flags': [],  # v2.18: 句子数不匹配标记（漏译检测）
        'range_stats': {},
        'alignment_mode': 'paragraph',  # default
        'alignment_warning': None,
    }

    source_paras = _extract_paragraphs(source_path)
    target_paras = _extract_paragraphs(target_path)

    if not source_paras or not target_paras:
        result['summary'] = {'total_issues': 0, 'by_severity': {'critical': 0, 'medium': 0, 'low': 0}}
        result['alignment_warning'] = '原文或译文中未提取到文本段落，无法执行机械检查'
        return result

    same_format = _formats_match(source_path, target_path)

    if same_format:
        # ── 同格式：逐段 1:1 对齐检查 ──
        all_range_stats = Counter()
        for i, (sp, tp) in enumerate(zip(source_paras, target_paras)):
            src_text = sp.get('text', '') if isinstance(sp, dict) else str(sp)
            tgt_text = tp.get('text', '') if isinstance(tp, dict) else str(tp)

            result['number_mismatches'].extend(check_numbers(src_text, tgt_text, i + 1))
            result['symbol_issues'].extend(check_symbols(src_text, tgt_text, i + 1, direction))
            result['unit_issues'].extend(check_units(src_text, tgt_text, i + 1))

            range_expr = extract_range_expressions(tgt_text)
            for k, v in range_expr.items():
                all_range_stats[k] += v

            if glossary or client_glossary:
                gl_result = check_glossary(src_text, tgt_text, glossary, i + 1,
                                           domain_filter, strict_short_terms,
                                           client_glossary)
                result['glossary_violations'].extend(gl_result['violations'])
                result['client_glossary_violations'].extend(gl_result['client_violations'])

            # v2.18: 句子数对齐检查（同格式才可靠）
            if sp.get('style', '') != 'TABLE_CELL' and tp.get('style', '') != 'TABLE_CELL':
                result['sentence_flags'].extend(
                    check_sentence_alignment(src_text, tgt_text, i + 1)
                )

        result['range_stats'] = dict(all_range_stats)
        result['range_consistency_warning'] = _check_range_consistency(all_range_stats)

    else:
        # ── 跨格式：逐页匹配模式 ──
        result['alignment_mode'] = 'cross_format_per_page'
        result['alignment_warning'] = (
            f'注意：原文为 {Path(source_path).suffix} 格式（{len(source_paras)} 个文本块），'
            f'译文为 {Path(target_path).suffix} 格式（{len(target_paras)} 个段落），'
            f'二者段落不对齐，使用逐页匹配模式。此模式下结果仅供参考，可能存在误报。'
        )

        full_src = '\n'.join(p.get('text', '') if isinstance(p, dict) else str(p) for p in source_paras)
        full_tgt = '\n'.join(p.get('text', '') if isinstance(p, dict) else str(p) for p in target_paras)

        # 数字检查（逐页匹配 + 假阳性过滤 + 记录 source_page）
        result['number_mismatches'].extend(_check_perpage_numbers(source_paras, full_tgt))

        # 术语检查（逐页匹配 + 记录 source_page）
        if glossary or client_glossary:
            if glossary:
                result['glossary_violations'].extend(_check_perpage_glossary(source_paras, full_tgt, glossary,
                                                                      domain_filter, strict_short_terms))
            if client_glossary:
                for pg_data in source_paras:
                    pg_text = pg_data.get('text', '') if isinstance(pg_data, dict) else str(pg_data)
                    gl_result = check_glossary(pg_text, full_tgt, None, 0,
                                               domain_filter, strict_short_terms,
                                               client_glossary)
                    for v in gl_result['client_violations']:
                        v['source_page'] = pg_data.get('page', 0) if isinstance(pg_data, dict) else 0
                    result['client_glossary_violations'].extend(gl_result['client_violations'])

        # ── 页面→段落映射 ──
        # 用滑动窗口相似度将 PDF 页码映射到 DOCX 段落
        if source_path.endswith('.pdf') and target_path.endswith('.docx'):
            page_map = _build_page_para_map(source_paras, target_paras)
            # 将 source_page 转为 paragraph_index
            for issue_list in [result['number_mismatches'], result['glossary_violations']]:
                for item in issue_list:
                    sp = item.get('source_page', 0)
                    if sp and sp in page_map:
                        paras = page_map[sp]
                        if paras:
                            # 轮询分配避免同页问题堆在同一段
                            key = f'_distrib_{sp}'
                            if key not in page_map:
                                page_map[key] = 0
                            idx = page_map[key] % len(paras)
                            page_map[key] += 1
                            item['paragraph_index'] = paras[idx]
                        else:
                            item['paragraph_index'] = 0
            result['page_map'] = {str(k): v for k, v in page_map.items() if isinstance(k, int)}

        # 符号检查 → 跨格式跳过（不可靠）
        # 范围表达统计
        all_range_stats = Counter()
        range_expr = extract_range_expressions(full_tgt)
        for k, v in range_expr.items():
            all_range_stats[k] += v
        result['range_stats'] = dict(all_range_stats)
        result['range_consistency_warning'] = _check_range_consistency(all_range_stats)

        # 跨格式时降低所有 issue severity
        for item in result['number_mismatches']:
            if item.get('severity') == 'critical':
                item['severity'] = 'medium'
        for item in result['glossary_violations']:
            if item.get('severity') == 'critical':
                item['severity'] = 'medium'

    # 格式检查（仅 Word）
    if source_path.endswith('.docx') and target_path.endswith('.docx'):
        result['format_issues'].extend(check_format(source_path, target_path))

    # 汇总
    all_severity_counts = {'critical': 0, 'medium': 0, 'low': 0}
    for key in ['number_mismatches', 'symbol_issues', 'glossary_violations',
                'client_glossary_violations', 'unit_issues', 'sentence_flags']:
        for item in result[key]:
            sev = item.get('severity', 'medium')
            if sev in all_severity_counts:
                all_severity_counts[sev] += 1
    all_severity_counts['low'] += len(result['format_issues'])

    result['summary'] = {
        'total_issues': sum(all_severity_counts.values()),
        'by_severity': all_severity_counts,
    }

    return result


def _build_page_para_map(source_paras: list[dict], target_paras: list[dict]) -> dict[int, list[int]]:
    """
    构建 PDF 页码 → DOCX 段落索引映射。
    使用滑动窗口相似度匹配每页文本到最佳段落范围。
    """
    from difflib import SequenceMatcher

    # Flatten target paragraphs to (1-based index, text)
    tgt_texts = []
    for i, p in enumerate(target_paras):
        text = p.get('text', '') if isinstance(p, dict) else str(p)
        if text.strip():
            tgt_texts.append((i + 1, text))

    n_tgt = len(tgt_texts)
    if n_tgt == 0:
        return {}

    # Window size: ~paragraphs per page
    window = max(2, n_tgt // max(1, len(source_paras)))
    page_map = {}

    for page_idx, page_data in enumerate(source_paras):
        page_text = page_data.get('text', '') if isinstance(page_data, dict) else str(page_data)
        page_num = page_data.get('page', page_idx + 1) if isinstance(page_data, dict) else page_idx + 1

        if not page_text.strip():
            # Inherit previous page's range
            prev = page_idx - 1
            while prev >= 0:
                prev_page_num = source_paras[prev].get('page', prev + 1) if isinstance(source_paras[prev], dict) else prev + 1
                if prev_page_num in page_map:
                    page_map[page_num] = page_map[prev_page_num]
                    break
                prev -= 1
            if page_num not in page_map:
                page_map[page_num] = []
            continue

        best_score = 0
        best_start = 0
        best_end = min(window, n_tgt)

        for start in range(0, n_tgt, max(1, window // 2)):
            end = min(start + window, n_tgt)
            combined = ' '.join(tgt_texts[i][1] for i in range(start, end))
            score = SequenceMatcher(None, page_text, combined).ratio()
            if score > best_score:
                best_score = score
                best_start = start
                best_end = end

        assigned = [tgt_texts[i][0] for i in range(best_start, best_end)]
        page_map[page_num] = assigned

    return page_map


def _check_range_consistency(stats: Counter) -> str | None:
    """检测范围表达是否混用多种格式"""
    non_zero = {k: v for k, v in stats.items() if v > 0}
    if len(non_zero) > 1:
        forms = ', '.join(f'"{k}"({v}次)' for k, v in non_zero.items())
        return f'范围表达混用多种格式: {forms}，建议全文统一'
    return None


def _extract_paragraphs(filepath: str) -> list[dict]:
    """从文件提取段落文本"""
    suffix = Path(filepath).suffix.lower()
    paras = []

    if suffix == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            for p in doc.paragraphs:
                paras.append({'text': p.text, 'style': p.style.name if p.style else ''})
            # v2.18: 提取表格单元格文本（标记为 TABLE_CELL 供下游区分）
            for ti, table in enumerate(doc.tables):
                for ri, row in enumerate(table.rows):
                    for ci, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            paras.append({
                                'text': cell_text,
                                'style': 'TABLE_CELL',
                                'table_index': ti + 1,
                                'row': ri + 1,
                                'col': ci + 1,
                            })
        except ImportError:
            pass
    elif suffix in ('.xlsx', '.xls'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows(values_only=True):
                    row_text = ' | '.join(str(c) if c is not None else '' for c in row)
                    if row_text.strip():
                        paras.append({'text': row_text, 'sheet': sheet_name})
            wb.close()
        except ImportError:
            pass
    elif suffix == '.txt':
        with open(filepath, 'r', encoding='utf-8') as f:
            text = f.read()
        for line in text.split('\n'):
            line = line.strip()
            if line:
                paras.append({'text': line})
    elif suffix == '.pdf':
        try:
            import fitz
            doc = fitz.open(filepath)
            for page in doc:
                paras.append({'text': page.get_text('text'), 'page': page.number + 1})
            doc.close()
        except ImportError:
            pass

    return paras


def main():
    parser = argparse.ArgumentParser(description='机械检查')
    parser.add_argument('--source', '-s', required=True, help='原文文件路径')
    parser.add_argument('--target', '-t', required=True, help='译文文件路径')
    parser.add_argument('--glossary', '-g', help='术语库 JSON 文件（由 load_glossary.py 生成）')
    parser.add_argument('--direction', '-d', default='en→zh', help='校对方向')
    parser.add_argument('--output', '-o', help='输出 JSON 路径')
    parser.add_argument('--domain', help='领域过滤：仅匹配该领域或"通用"的术语')
    parser.add_argument('--strict-short-terms', action='store_true',
                        help='短词（≤3字符）保持原严重度，不自动降级')
    parser.add_argument('--client-glossary', help='客户术语需求文件（_术语要求.xlsx 或 client_terms.json）')
    args = parser.parse_args()

    glossary = None
    if args.glossary:
        try:
            gl_data = json.loads(Path(args.glossary).read_text(encoding='utf-8'))
            glossary = gl_data.get('terms', {})
        except Exception as e:
            print(f'警告: 无法加载术语库: {e}', file=sys.stderr)

    client_glossary = None
    if args.client_glossary:
        try:
            from load_glossary import load_client_glossary
            cl_data = load_client_glossary(args.client_glossary)
            client_glossary = cl_data.get('terms', {})
            print(f'加载客户术语要求: {cl_data.get("count", 0)} 条 → {args.client_glossary}', file=sys.stderr)
        except Exception as e:
            print(f'警告: 无法加载客户术语文件: {e}', file=sys.stderr)

    result = check_all(args.source, args.target, glossary, args.direction,
                       args.domain, args.strict_short_terms, client_glossary)

    output_json = json.dumps(result, ensure_ascii=False, indent=2)
    if args.output:
        Path(args.output).write_text(output_json, encoding='utf-8')
        print(f'已写入: {args.output}')
    else:
        sys.stdout.reconfigure(encoding='utf-8')
        print(output_json)


import sys
if __name__ == '__main__':
    main()
