#!/usr/bin/env python3
"""
Phase 6: 批注写入（v6 — v5 + WPS 兼容：commentRangeStart/End 包裹每个批注 run）

关键修复:
  - commentReference 放在高亮 run 的 rPr 内部（紧挨 w:shd + w:highlight）
  - commentRangeStart/End 加回（WPS 兼容性所需），commentReference 保留在 rPr 内（Word 高亮 fix）
  - document.xml 只用 python-docx OxmlElement 操作，禁止 lxml 触碰
  - lxml 仅用于构建 comments.xml / rels / CT（ZIP 层注入）
  - 跨 run 匹配：target_quote 跨多段 run 时，通过全文拼接定位起始 run
  - 多长度 key 匹配: [15, 10, 8, 5] 递降
  - v5: pi=0 的全部问题（机械检查 + 语义审查无定位）均作为文档级裸引用插入第一段，
    段落级 comment_id 优先分配，批注带 [全文匹配，仅供参考] 前缀
"""

import argparse, json, shutil, zipfile, io, os, re, sys, tempfile
from pathlib import Path
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

# ── Namespaces (lxml, used ONLY for comments.xml/rels/CT — never document.xml) ──
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
COMMENTS_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
CT_COMMENTS = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
XML_NS = 'http://www.w3.org/XML/1998/namespace'
DOC_RELS = 'word/_rels/document.xml.rels'
CT_PATH = '[Content_Types].xml'
COMMENTS_PATH = 'word/comments.xml'

HL_COLORS = {'critical': ('FFCCCC', 'red'), 'medium': ('FFF9C4', 'yellow'), 'low': ('C8E6C9', 'green')}
SEV_ORDER = {'critical': 0, 'medium': 1, 'low': 2}


def _build_table_cell_map(split_target_path: str) -> dict:
    """从 split_target.json 构建 paragraph_index → (table_index, row, col) 映射。

    扫描所有章节的 tables 数组，为每个非空 TABLE_CELL 条目分配一个连续的
    paragraph_index（从 len(doc.paragraphs)+1 开始），返回映射字典。
    """
    if not split_target_path or not os.path.exists(split_target_path):
        return {}

    try:
        with open(split_target_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return {}

    # 先统计总段落数（用于确定 table cell 索引起点）
    # split_target.json 的 paragraphs 索引与 doc.paragraphs 对齐
    total_paras = 0
    if 'chapters' in data:
        for ch in data['chapters']:
            total_paras += len(ch.get('paragraphs', []))

    cell_map = {}
    cell_idx = total_paras + 1
    if 'chapters' in data:
        for ch in data['chapters']:
            for table in ch.get('tables', []):
                ti = table.get('table_index', 0)
                rows_data = table.get('data', table.get('rows', []))
                for ri, row in enumerate(rows_data):
                    if not isinstance(row, list):
                        continue
                    for ci, cell in enumerate(row):
                        cell_text = ''
                        if isinstance(cell, dict):
                            cell_text = cell.get('text', '')
                        elif isinstance(cell, str):
                            cell_text = cell
                        if cell_text.strip():
                            cell_map[cell_idx] = {
                                'table_index': ti,
                                'row': ri + 1,
                                'col': ci + 1,
                            }
                            cell_idx += 1
    return cell_map


def write_word_comments(filepath: str, issues: list[dict],
                        output_path: str = None, verbose: bool = False,
                        table_cell_map: dict = None) -> dict:
    """主入口：写高亮 + 批注到译文 DOCX。"""
    try:
        from lxml import etree
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        return {'status': 'error', 'message': f'缺少依赖: {e}'}

    def _add_cmt_range(para_elem, run_elem, cid):
        """Insert w:commentRangeStart before and w:commentRangeEnd after run_elem in para_elem.

        Required for WPS compatibility — WPS cannot locate comments without range markers,
        even though Word renders them fine with just commentReference in rPr.
        """
        children = list(para_elem)
        try:
            pos = children.index(run_elem)
        except ValueError:
            return
        rs = OxmlElement('w:commentRangeStart')
        rs.set(qn('w:id'), str(cid))
        para_elem.insert(pos, rs)
        # Re-find run position (shifted by the start marker insertion)
        children = list(para_elem)
        pos = children.index(run_elem)
        re = OxmlElement('w:commentRangeEnd')
        re.set(qn('w:id'), str(cid))
        para_elem.insert(pos + 1, re)

    output = output_path or filepath.replace('.docx', '_校对稿.docx')
    os.makedirs(os.path.dirname(output) or '.', exist_ok=True)

    # ─── 按段落分组、排序 ───
    # v5: pi=0 的全部走文档级路径（机械检查 + 语义审查无定位的均在此）
    #     段落级问题先分配 comment_id，文档级问题后分配
    issues_by_para = {}
    doc_level_issues = []
    mech_count = 0
    for issue in issues:
        pi = issue.get('paragraph_index', 0)
        if pi < 1:
            doc_level_issues.append(issue)
            if issue.get('type'):
                mech_count += 1
        else:
            issues_by_para.setdefault(pi, []).append(issue)

    # ─── 合并 pi=0 机械问题：按 type 归为综合批注 ───
    # 96 条散落在第一段太乱，按数字缺失/小数位/术语违规三类各合并为一条
    type_label_cn = {
        'number_missing': '数字缺失', 'decimal_mismatch': '小数位不匹配',
        'glossary_violation': '术语库违规',
    }
    mech_groups = {}  # type → list of issues
    other_doc = []    # non-mechanical pi=0 issues
    for issue in doc_level_issues:
        typ = issue.get('type', '')
        if typ and typ in type_label_cn:
            mech_groups.setdefault(typ, []).append(issue)
        else:
            other_doc.append(issue)

    # Build consolidated mechanical issues
    consolidated_mech = []
    for typ, group in mech_groups.items():
        label = type_label_cn.get(typ, typ)
        svals = []
        for iss in group:
            sv = iss.get('source_value') or iss.get('source_term', '')
            if sv:
                svals.append(sv)
        sev = group[0].get('severity', 'medium')
        dim = group[0].get('dimension', '')
        total = len(group)

        # Build rich text body
        body_lines = [
            f'以下 {total} 处{label}在原文中存在但译文中未找到对应项',
            f'（跨格式 PDF→DOCX 全文匹配，仅供参考，建议人工逐条复核）：',
            '',
        ]
        # Wrap values in groups of 10 per line
        for i in range(0, len(svals), 10):
            chunk = svals[i:i+10]
            body_lines.append('、'.join(chunk))

        consolidated_mech.append({
            'paragraph_index': 0,
            'type': typ,
            '_consolidated': True,
            '_count': total,
            '_values': svals,
            '_body': '\n'.join(body_lines),
            'severity': sev,
            'dimension': dim,
            'source_value': f'{label}（{total}处）',
            'check': '\n'.join(body_lines),
        })

    # Replace doc_level_issues: other (non-mech) + consolidated mech
    doc_level_issues = other_doc + consolidated_mech
    if verbose and consolidated_mech:
        print(f'  Consolidated {mech_count} mechanical issues into {len(consolidated_mech)} comments')

    # ══════════════════════════════════════════════════
    # STEP 1: python-docx does ALL document.xml work
    #         (highlights + commentReference in rPr)
    # ══════════════════════════════════════════════════
    doc = Document(filepath)
    comment_id = 0
    para_comments = {}
    total_matched = 0
    total_fuzzy = 0
    total_neighbor = 0
    total_global = 0
    total_bare = 0

    # v2.24: Pre-build global full-text index for document-wide fallback search
    _all_para_texts = []
    for i, p in enumerate(doc.paragraphs):
        ft = ''.join(r.text for r in p.runs)
        _all_para_texts.append((i + 1, ft))

    # v2.24: Resolve pi=-1 issues by searching document for target_quote
    # Many issues have pi=-1 because the Agent couldn't determine paragraph_index,
    # but the target_quote text IS present somewhere in the document.
    # This pre-fix resolves them before the matching pipeline runs.
    resolved_from_doc_level = 0
    still_doc_level = []
    for issue in doc_level_issues:
        tq = issue.get('target_quote', '').strip()
        is_mech = bool(issue.get('type'))
        if tq and not is_mech:
            found = False
            for gpi, gtext in _all_para_texts:
                if tq in gtext:
                    issues_by_para.setdefault(gpi, []).append(issue)
                    resolved_from_doc_level += 1
                    found = True
                    break
            if not found:
                still_doc_level.append(issue)
        else:
            still_doc_level.append(issue)
    doc_level_issues = still_doc_level
    if resolved_from_doc_level > 0:
        print(f'[write_comments] v2.24: Resolved {resolved_from_doc_level} pi=-1 issues → found in paragraphs')
        # Re-count mechanical issues in the remaining doc_level_issues
        mech_count = sum(1 for i in doc_level_issues if i.get('type'))

    # v2.23: 匹配失败诊断计数器
    diag = {
        "empty_target": 0,
        "para_out_of_range": 0,
        "para_empty": 0,
        "target_not_in_para": 0,
        "matched_exact": 0,
        "matched_fuzzy": 0,
        "matched_neighbor": 0,
        "matched_global": 0,
    }

    # ── 1a: Paragraph-level issues (pi >= 1) ──
    # Process first so they get lower comment_ids
    for para_idx in sorted(issues_by_para.keys()):
        is_table_cell = False
        if para_idx > len(doc.paragraphs):
            # v2.22: 检查是否为表格单元格索引
            cell_info = (table_cell_map or {}).get(para_idx)
            if cell_info:
                try:
                    ti = cell_info['table_index']
                    ri = cell_info['row']
                    ci = cell_info['col']
                    cell = doc.tables[ti - 1].rows[ri - 1].cells[ci - 1]
                    para = cell.paragraphs[0]
                    is_table_cell = True
                except (IndexError, AttributeError):
                    continue
            else:
                continue
        else:
            para = doc.paragraphs[para_idx - 1]

        sorted_issues = sorted(issues_by_para[para_idx],
                               key=lambda x: SEV_ORDER.get(x.get('severity', 'medium'), 1))

        para_comments[para_idx] = []

        # Pre-compute full paragraph text for cross-run matching
        full_text = ''.join(r.text for r in para.runs)

        for issue in sorted_issues:
            sev = issue.get('severity', 'medium')
            fill_color, hl_name = HL_COLORS.get(sev, ('FFFF00', 'yellow'))
            target = issue.get('target_quote', '').strip()

            para_comments[para_idx].append((comment_id, sev, issue))
            matched = False

            if target:
                # Strategy 1: find target in individual runs
                for key_len in [15, 10, 8, 5]:
                    key = target[:key_len].strip()
                    if not key or len(key) < 3:
                        continue
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        if key in run.text:
                            rPr = run._r.get_or_add_rPr()
                            # Already highlighted by a previous issue
                            if rPr.find(qn('w:shd')) is not None:
                                cr = OxmlElement('w:commentReference')
                                cr.set(qn('w:id'), str(comment_id))
                                rPr.append(cr)
                            else:
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), fill_color)
                                rPr.append(shd)
                                hl = OxmlElement('w:highlight')
                                hl.set(qn('w:val'), hl_name)
                                rPr.append(hl)
                                cr = OxmlElement('w:commentReference')
                                cr.set(qn('w:id'), str(comment_id))
                                rPr.append(cr)
                            _add_cmt_range(para._p, run._r, comment_id)
                            total_matched += 1
                            diag["matched_exact"] += 1
                            matched = True
                            break
                    if matched:
                        break

                # Strategy 2: target exists in paragraph but split across runs
                if not matched and target in full_text:
                    target_start = full_text.index(target)
                    char_pos = 0
                    for run in para.runs:
                        run_len = len(run.text)
                        if char_pos <= target_start < char_pos + run_len:
                            rPr = run._r.get_or_add_rPr()
                            if rPr.find(qn('w:shd')) is None:
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), fill_color)
                                rPr.append(shd)
                                hl = OxmlElement('w:highlight')
                                hl.set(qn('w:val'), hl_name)
                                rPr.append(hl)
                            cr = OxmlElement('w:commentReference')
                            cr.set(qn('w:id'), str(comment_id))
                            rPr.append(cr)
                            _add_cmt_range(para._p, run._r, comment_id)
                            total_matched += 1
                            diag["matched_exact"] += 1
                            matched = True
                            break
                        char_pos += run_len

            if not matched and target:
                # ── v2.23 Strategy 3: fuzzy matching ──
                # 3a: 去空白后精确匹配
                target_nosp = re.sub(r'\s+', '', target)
                full_nosp = re.sub(r'\s+', '', full_text)
                if len(target_nosp) >= 4 and target_nosp in full_nosp:
                    # 在原始 full_text 中定位
                    pos = full_nosp.index(target_nosp)
                    # 找到对应的 run
                    char_count = 0
                    for run in para.runs:
                        run_text_nosp = re.sub(r'\s+', '', run.text)
                        run_start = char_count
                        run_end = char_count + len(run_text_nosp)
                        if run_start <= pos < run_end:
                            rPr = run._r.get_or_add_rPr()
                            if rPr.find(qn('w:shd')) is None:
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), fill_color)
                                rPr.append(shd)
                                hl = OxmlElement('w:highlight')
                                hl.set(qn('w:val'), hl_name)
                                rPr.append(hl)
                            cr = OxmlElement('w:commentReference')
                            cr.set(qn('w:id'), str(comment_id))
                            rPr.append(cr)
                            _add_cmt_range(para._p, run._r, comment_id)
                            total_fuzzy += 1
                            matched = True
                            diag["matched_fuzzy"] += 1
                            break
                        char_count = run_end

                # 3b: 最长中文片段匹配
                if not matched and target:
                    import unicodedata
                    # 提取 target 中最长的纯中文字段 (≥4字)
                    chunks = re.findall(r'[一-鿿]{4,}', target)
                    for chunk in sorted(chunks, key=len, reverse=True):
                        if chunk in full_text:
                            # 同 Strategy 2 的高亮逻辑
                            target_start = full_text.index(chunk)
                            char_pos = 0
                            for run in para.runs:
                                run_len = len(run.text)
                                if char_pos <= target_start < char_pos + run_len:
                                    rPr = run._r.get_or_add_rPr()
                                    if rPr.find(qn('w:shd')) is None:
                                        shd = OxmlElement('w:shd')
                                        shd.set(qn('w:val'), 'clear')
                                        shd.set(qn('w:color'), 'auto')
                                        shd.set(qn('w:fill'), fill_color)
                                        rPr.append(shd)
                                        hl = OxmlElement('w:highlight')
                                        hl.set(qn('w:val'), hl_name)
                                        rPr.append(hl)
                                    cr = OxmlElement('w:commentReference')
                                    cr.set(qn('w:id'), str(comment_id))
                                    rPr.append(cr)
                                    _add_cmt_range(para._p, run._r, comment_id)
                                    total_fuzzy += 1
                                    matched = True
                                    diag["matched_fuzzy"] += 1
                                    break
                                char_pos += run_len
                        if matched:
                            break

            # ── v2.23: 邻段搜索 (当本段为空或 target 不在本段) ──
            if not matched and target:
                # 尝试上下 ±5 段
                for offset in list(range(1, 6)) + list(range(-1, -6, -1)):
                    nb_idx = para_idx + offset
                    if 1 <= nb_idx <= len(doc.paragraphs):
                        nb_para = doc.paragraphs[nb_idx - 1]
                        nb_text = ''.join(r.text for r in nb_para.runs)
                        if target in nb_text or (len(target) >= 4 and target[:10] in nb_text):
                            # 标记找到的邻段为匹配位置
                            for key_len in [15, 10, 8, 5]:
                                key = target[:key_len].strip()
                                if not key or len(key) < 3:
                                    continue
                                for run in nb_para.runs:
                                    if not run.text.strip():
                                        continue
                                    if key in run.text:
                                        rPr = run._r.get_or_add_rPr()
                                        if rPr.find(qn('w:shd')) is None:
                                            shd = OxmlElement('w:shd')
                                            shd.set(qn('w:val'), 'clear')
                                            shd.set(qn('w:color'), 'auto')
                                            shd.set(qn('w:fill'), fill_color)
                                            rPr.append(shd)
                                            hl = OxmlElement('w:highlight')
                                            hl.set(qn('w:val'), hl_name)
                                            rPr.append(hl)
                                        cr = OxmlElement('w:commentReference')
                                        cr.set(qn('w:id'), str(comment_id))
                                        rPr.append(cr)
                                        _add_cmt_range(nb_para._p, run._r, comment_id)
                                        total_neighbor += 1
                                        matched = True
                                        diag["matched_neighbor"] += 1
                                        break
                                if matched:
                                    break
                            break
                if matched:
                    break

            # ── v2.24: Document-wide search as ultimate fallback ──
            # When paragraph_index is wrong (Agent used item_id instead of split_index),
            # the claimed paragraph and its ±5 neighbors won't contain the target text.
            # Search the ENTIRE document — target_quote text IS correct and somewhere in the DOCX.
            if not matched and target:
                for gpi, gtext in _all_para_texts:
                    # Skip the paragraph and neighbors we already checked (avoid redundant work)
                    if abs(gpi - para_idx) <= 5:
                        continue
                    if target in gtext:
                        nb_para = doc.paragraphs[gpi - 1]
                        for key_len in [15, 10, 8, 5]:
                            key = target[:key_len].strip()
                            if not key or len(key) < 3:
                                continue
                            for run in nb_para.runs:
                                if not run.text.strip():
                                    continue
                                if key in run.text:
                                    rPr = run._r.get_or_add_rPr()
                                    if rPr.find(qn('w:shd')) is None:
                                        shd = OxmlElement('w:shd')
                                        shd.set(qn('w:val'), 'clear')
                                        shd.set(qn('w:color'), 'auto')
                                        shd.set(qn('w:fill'), fill_color)
                                        rPr.append(shd)
                                        hl = OxmlElement('w:highlight')
                                        hl.set(qn('w:val'), hl_name)
                                        rPr.append(hl)
                                    cr = OxmlElement('w:commentReference')
                                    cr.set(qn('w:id'), str(comment_id))
                                    rPr.append(cr)
                                    _add_cmt_range(nb_para._p, run._r, comment_id)
                                    total_global += 1
                                    diag["matched_global"] += 1
                                    matched = True
                                    break
                            if matched:
                                break
                        if matched:
                            break

            if not matched:
                # Bare ref: insert after pPr (NOT at paragraph end — avoids highlight suppression)
                if not target:
                    diag["empty_target"] += 1
                elif not full_text.strip():
                    diag["para_empty"] += 1
                else:
                    diag["target_not_in_para"] += 1
                total_bare += 1
                p_elem = para._p
                pPr = p_elem.find(qn('w:pPr'))
                insert_pos = (list(p_elem).index(pPr) + 1) if pPr is not None else 0
                ref_run = OxmlElement('w:r')
                ref_rPr = OxmlElement('w:rPr')
                cr = OxmlElement('w:commentReference')
                cr.set(qn('w:id'), str(comment_id))
                ref_rPr.append(cr)
                ref_run.append(ref_rPr)
                p_elem.insert(insert_pos, ref_run)
                _add_cmt_range(p_elem, ref_run, comment_id)

            comment_id += 1

    # ── 1b: Document-level issues (pi=0) ──
    # Insert bare refs at LAST paragraph (appendix-style, not polluting the beginning)
    doc_level_count = 0
    if doc_level_issues and len(doc.paragraphs) > 0:
        para_comments[-1] = []  # special key -1 for doc-level (appendix)
        plast = doc.paragraphs[-1]
        pelem = plast._p
        pPr = pelem.find(qn('w:pPr'))
        insert_pos = (list(pelem).index(pPr) + 1) if pPr is not None else 0

        for issue in doc_level_issues:
            sev = issue.get('severity', 'medium')
            para_comments[-1].append((comment_id, sev, issue))

            # Bare ref (no highlight — can't locate exact run for pi=0)
            ref_run = OxmlElement('w:r')
            ref_rPr = OxmlElement('w:rPr')
            cr = OxmlElement('w:commentReference')
            cr.set(qn('w:id'), str(comment_id))
            ref_rPr.append(cr)
            ref_run.append(ref_rPr)
            pelem.insert(insert_pos, ref_run)
            _add_cmt_range(pelem, ref_run, comment_id)

            comment_id += 1
            doc_level_count += 1

    total_issues = comment_id

    # Save via python-docx → highlights + comment refs preserved
    tmpfd, tmppath = tempfile.mkstemp(suffix='.docx')
    os.close(tmpfd)
    doc.save(tmppath)

    if verbose:
        print(f'STEP 1: highlights + {comment_id} comment refs written via python-docx')
        if mech_count > 0:
            print(f'  Mechanical (pi=0, doc-level): {mech_count} issues')
        if doc_level_count > 0:
            print(f'  Document-level (pi=0): {doc_level_count} comments at last paragraph')

    # ── v2.23: 匹配诊断 ──
    total_matched_all = total_matched + total_fuzzy + total_neighbor + total_global
    total_with_target = total_issues - diag["empty_target"]
    if total_with_target > 0:
        match_rate = total_matched_all / total_with_target * 100
    else:
        match_rate = 0
    print(f'[write_comments] 匹配统计: exact={total_matched}, fuzzy={total_fuzzy}, neighbor={total_neighbor}, global={total_global}, bare={total_bare}')
    print(f'[write_comments] 匹配率: {total_matched_all}/{total_with_target} = {match_rate:.0f}%')
    if total_bare > total_with_target * 0.5:
        print(f'[write_comments] ⚠ 匹配率仅{match_rate:.0f}%——建议检查paragraph_index是否与DOCX对齐，'
              f'target_quote是否从split_target.json逐字复制')
    if verbose:
        print(f'[write_comments] 失败原因: empty_target={diag["empty_target"]}, '
              f'para_empty={diag["para_empty"]}, target_not_found={diag["target_not_in_para"]}')
        if total_global > 0:
            print(f'[write_comments] ⚠ global search rescued {total_global} issues — paragraph_index was significantly wrong')

    # ══════════════════════════════════════════════════
    # STEP 2: ZIP layer — build comments.xml,
    #          update rels, update CT
    #          document.xml COPIED VERBATIM
    # ══════════════════════════════════════════════════
    from lxml import etree as lxe

    entries = {}
    with zipfile.ZipFile(tmppath, 'r') as zf:
        for info in zf.infolist():
            entries[info.filename] = (info, zf.read(info.filename))

    # --- 2a: Build comments.xml ---
    comments_tree = lxe.Element(f'{{{NS_W}}}comments', nsmap={'w': NS_W})

    for para_idx in sorted(para_comments.keys()):
        for cid, sev, issue in para_comments[para_idx]:
            is_doc_level = (para_idx == -1)
            is_consolidated = issue.get('_consolidated', False)
            is_mechanical = bool(issue.get('type'))  # Phase 3 mechanical issue
            date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
            cmt = lxe.SubElement(comments_tree, f'{{{NS_W}}}comment')
            cmt.set(f'{{{NS_W}}}id', str(cid))
            cmt.set(f'{{{NS_W}}}author', 'TranslationReview')
            cmt.set(f'{{{NS_W}}}date', date_str)
            cmt.set(f'{{{NS_W}}}initials', 'TR')

            sev_label = {
                'critical': '🔴 严重（必须修改）',
                'medium': '🟡 中等（建议修改）',
                'low': '🟢 低（格式优化）',
            }.get(sev, '⚪ 未分级')

            # Mechanical issues (Phase 3 cross-format) get a warning prefix
            doc_prefix = '⚠️ [全文匹配，仅供参考] ' if (is_doc_level or is_mechanical) else ''

            # Normalize field access: fall back to mechanical-issue field names
            def _get(issue, *keys):
                for k in keys:
                    v = issue.get(k, '')
                    if v:
                        return v
                return ''
            src_quote = _get(issue, 'source_quote', 'source_value', 'source_term')
            tgt_quote = _get(issue, 'target_quote', 'expected_target')
            iss_text = _get(issue, 'issue')
            sug_text = _get(issue, 'suggestion', 'check')
            dim_text = _get(issue, 'dimension')
            chk_text = _get(issue, 'check_item', 'type')
            # Map raw type names to Chinese labels
            type_map = {
                'number_missing': '数字缺失', 'decimal_mismatch': '小数位不匹配',
                'glossary_violation': '术语库违规', 'symbol_missing': '符号缺失',
                'punctuation_mixed': '标点混用', 'format_issue': '格式问题',
                'unit_issue': '单位问题',
            }
            if chk_text in type_map:
                chk_text = type_map[chk_text]

            if is_consolidated:
                # Consolidated mechanical issue: header + pre-built body
                count = issue.get('_count', 0)
                lines = [
                    doc_prefix + sev_label,
                    f'【{dim_text} — {chk_text}（{count}处）】', '',
                ]
                for bline in issue.get('_body', '').split('\n'):
                    lines.append(bline)
            elif is_doc_level or is_mechanical:
                # Single mechanical issue (doc-level or PDF-page-located):
                # check field contains the full description
                check_desc = sug_text
                lines = [doc_prefix + sev_label, f'【{dim_text} — {chk_text}】']
                if src_quote:
                    lines.extend(['', f'原文: {src_quote}'])
                lines.extend(['', check_desc])
            else:
                # Paragraph-level semantic: build dynamically, skip empty fields
                lines = [sev_label, f'【{dim_text} — {chk_text}】']
                if src_quote:
                    lines.extend(['', f'原文: {src_quote}'])
                if tgt_quote:
                    lines.extend(['', f'译文: {tgt_quote}'])
                if iss_text:
                    lines.extend(['', f'问题: {iss_text}'])
                if sug_text:
                    lines.extend(['', f'建议: {sug_text}'])
            for i, line in enumerate(lines):
                pe = lxe.SubElement(cmt, f'{{{NS_W}}}p')
                pp = lxe.SubElement(pe, f'{{{NS_W}}}pPr')
                r_e = lxe.SubElement(pe, f'{{{NS_W}}}r')
                if i == 0:
                    rp = lxe.SubElement(r_e, f'{{{NS_W}}}rPr')
                    lxe.SubElement(rp, f'{{{NS_W}}}b')
                    sz = lxe.SubElement(rp, f'{{{NS_W}}}sz')
                    sz.set(f'{{{NS_W}}}val', '18')
                    co = lxe.SubElement(rp, f'{{{NS_W}}}color')
                    co.set(f'{{{NS_W}}}val',
                           {'critical': 'FF0000', 'medium': 'FF8C00', 'low': '228B22'}.get(sev, '000000'))
                te = lxe.SubElement(r_e, f'{{{NS_W}}}t')
                te.set(f'{{{XML_NS}}}space', 'preserve')
                te.text = line

    comments_xml = lxe.tostring(comments_tree, xml_declaration=True,
                                 encoding='UTF-8', standalone=True)
    entries[COMMENTS_PATH] = (None, comments_xml)

    # --- 2b: Update word/_rels/document.xml.rels ---
    rels_raw = entries.get(DOC_RELS, (None, None))[1]
    rels_tree = lxe.fromstring(rels_raw) if rels_raw else lxe.Element(f'{{{REL_NS}}}Relationships')

    max_rid = 0
    has_rel = False
    for rel in rels_tree.iter():
        if rel.tag == f'{{{REL_NS}}}Relationship':
            rid = rel.get('Id', '')
            m = re.match(r'rId(\d+)', rid)
            if m:
                max_rid = max(max_rid, int(m.group(1)))
            if rel.get('Type') == COMMENTS_NS:
                has_rel = True

    if not has_rel:
        nr = lxe.SubElement(rels_tree, f'{{{REL_NS}}}Relationship')
        nr.set('Id', f'rId{max_rid + 1}')
        nr.set('Type', COMMENTS_NS)
        nr.set('Target', 'comments.xml')
        entries[DOC_RELS] = (
            entries.get(DOC_RELS, (None, None))[0],
            lxe.tostring(rels_tree, xml_declaration=True, encoding='UTF-8', standalone=True),
        )

    # --- 2c: Update [Content_Types].xml ---
    ct_raw = entries.get(CT_PATH, (None, None))[1]
    ct_tree = lxe.fromstring(ct_raw) if ct_raw else lxe.Element(f'{{{CT_NS}}}Types')

    has_ct = False
    for ov in ct_tree.iter():
        if ov.tag == f'{{{CT_NS}}}Override' and ov.get('PartName') == '/word/comments.xml':
            has_ct = True
            break

    if not has_ct:
        nct = lxe.SubElement(ct_tree, f'{{{CT_NS}}}Override')
        nct.set('PartName', '/word/comments.xml')
        nct.set('ContentType', CT_COMMENTS)
        entries[CT_PATH] = (
            entries.get(CT_PATH, (None, None))[0],
            lxe.tostring(ct_tree, xml_declaration=True, encoding='UTF-8', standalone=True),
        )

    # --- 2d: Write final ZIP ---
    with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
        for name, (info, data) in entries.items():
            if info is not None:
                zf.writestr(info, data)
            else:
                zf.writestr(name, data)

    os.unlink(tmppath)

    # ── Verify ──
    with zipfile.ZipFile(output, 'r') as zf:
        doc_xml = zf.read('word/document.xml')
        shd_count = doc_xml.count(b'w:shd')
        ref_count = doc_xml.count(b'commentReference')
        has_lxml_ns = b'ns0:' in doc_xml or b'ns1:' in doc_xml

    unmatched = total_issues - total_matched_all

    if verbose:
        print(f'STEP 2: comments.xml + rels + CT injected via ZIP')
        print(f'  Highlights (w:shd): {shd_count}')
        print(f'  Comment refs: {ref_count}')
        print(f'  lxml artifacts: {"NONE — OK" if not has_lxml_ns else "FOUND — BAD!"}')
        print(f'  Matched: {total_matched_all}/{total_issues} highlighted')
        print(f'  Unmatched: {unmatched} (bare ref)')
        print(f'  Output: {output}')

    return {
        'status': 'ok',
        'output': str(Path(output).resolve()),
        'total_comments': total_issues,
        'highlights': shd_count,
        'matched': total_matched_all,
        'unmatched': unmatched,
        'has_lxml_artifacts': has_lxml_ns,
    }


def main():
    parser = argparse.ArgumentParser(description='Phase 6: 批注写入（v3 — 修复高亮渲染 bug）')
    parser.add_argument('--input', '-i', required=True, help='译文文件路径')
    parser.add_argument('--issues', '-j', required=True, help='疑点 JSON 文件路径')
    parser.add_argument('--output', '-o', help='输出路径（默认: 原文名_校对稿.docx）')
    parser.add_argument('--split-target', help='split_target.json 路径（v2.22: 用于表格单元格定位）')
    parser.add_argument('--verbose', '-v', action='store_true', help='详细输出')
    args = parser.parse_args()

    data = json.loads(Path(args.issues).read_text(encoding='utf-8'))
    if isinstance(data, list):
        issues = data
    elif isinstance(data, dict):
        issues = data.get('issues', data.get('findings', []))
    else:
        issues = []

    if not issues:
        print(json.dumps({'status': 'error', 'message': 'issues 列表为空'}, ensure_ascii=False))
        sys.exit(1)

    # v2.22: 构建表格单元格索引映射（用于 phase 4 表格问题标注）
    table_cell_map = _build_table_cell_map(args.split_target) if args.split_target else {}

    result = write_word_comments(args.input, issues, args.output,
                                 verbose=args.verbose, table_cell_map=table_cell_map)
    print(json.dumps(result, ensure_ascii=False))


if __name__ == '__main__':
    main()
